# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 1: Importaciones y Configuración Base
# ==================================================================================
#
# Este artefacto contiene:
# 1. Todas las importaciones de bibliotecas necesarias para la aplicación
# 2. Configuración base de la aplicación Streamlit
# 3. Configuración de logging
# 4. Definición de la versión de la aplicación
# 5. Configuración de rutas para assets
# 6. Manejo seguro de errores de importación
#
# Es fundamental para establecer el entorno de ejecución y garantizar que todas
# las dependencias estén correctamente importadas para su uso en artefactos posteriores.
# ==================================================================================

# Bibliotecas estándar de Python
import os
import sys
import json
import time
import uuid
import base64
import hashlib
import logging
import traceback
import re
import csv
import copy
from io import BytesIO, StringIO
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Tuple, Union, Any
from urllib.parse import urlparse

# Configuración de logging mejorada
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("SpanishFactorIA")

# Bloque try/except para manejar bibliotecas faltantes y proporcionar mensajes útiles
try:
    # Bibliotecas de terceros - UI y visualización
    import streamlit as st
    import pandas as pd
    import numpy as np
    import matplotlib.pyplot as plt
    import altair as alt
    from PIL import Image
    import qrcode

    # Para análisis estadístico
    from scipy import stats

    # Bibliotecas de terceros - Procesamiento de documentos
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    import markdown

    # Bibliotecas de terceros - Servicios cloud y APIs
    import requests

except ImportError as e:
    module_name = str(e).split("'")[1]
    logger.error(
        f"Error al importar el módulo {module_name}. Por favor, instálalo con: pip install {module_name}")
    if 'streamlit' in str(e):
        print(f"Error crítico: No se pudo importar streamlit. Instálalo con: pip install streamlit")
        sys.exit(1)
    # Para otras bibliotecas, continuamos pero registramos el error

# Verificar versiones de bibliotecas críticas para compatibilidad
try:
    import pkg_resources
    streamlit_version = pkg_resources.get_distribution("streamlit").version
    logger.info(f"Versión de streamlit: {streamlit_version}")
except Exception as e:
    logger.warning(
        f"No se pudieron verificar las versiones de las bibliotecas: {e}")

# Configuración de versión y constantes globales
APP_VERSION = "3.2.0"  # Versión actualizada
APP_NAME = "Textocorrector ELE"
ORGANIZATION = "Spanish FactorIA"

# Configuración de Streamlit
st.set_page_config(
    page_title=f"{APP_NAME} - {ORGANIZATION}",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Definir rutas para assets
ASSETS_PATH = "assets"
LOGO_PATH = os.path.join(ASSETS_PATH, "Spanish_FactorIA_Logo.png")

# Verificar existencia de directorios necesarios y crearlos si no existen
try:
    if not os.path.exists(ASSETS_PATH):
        os.makedirs(ASSETS_PATH)
        logger.warning(f"Directorio de assets creado: {ASSETS_PATH}")
except Exception as e:
    logger.error(f"Error al verificar/crear directorios: {e}")

# Función de utilidad para manejar excepciones de manera consistente


def handle_exception(func_name, exception, show_user=True):
    """
    Función de utilidad para manejar excepciones de manera consistente.

    Args:
        func_name: Nombre de la función donde ocurrió el error
        exception: La excepción capturada
        show_user: Si se debe mostrar un mensaje al usuario

    Returns:
        None
    """
    error_msg = f"Error en {func_name}: {str(exception)}"
    logger.error(error_msg)
    logger.error(traceback.format_exc())

    if show_user and 'st' in globals():
        st.error(f"⚠️ {error_msg}")
        with st.expander("Detalles técnicos", expanded=False):
            st.code(traceback.format_exc())
            st.info(
                "Si el problema persiste, contacta con el administrador del sistema.")

    return None


# Inicializar caché y estado de sesión
if 'st' in globals():
    try:
        # Limpiar caché al inicio para evitar problemas con datos anteriores
        st.cache_data.clear()
    except Exception as e:
        logger.warning(f"No se pudo limpiar la caché: {e}")


def load_custom_css():
    """
    Carga estilos CSS personalizados para la aplicación,
    incluyendo la fuente Poppins y mejoras en la interfaz.
    """
    st.markdown("""
    <style>
    /* Importar fuente Poppins de Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap');
    
    /* Aplicar Poppins a todos los elementos */
    html, body, [class*="st-"] {
        font-family: 'Poppins', sans-serif;
    }
    
    /* Bordes redondeados y sombras para los elementos interactivos */
    div.stButton > button:first-child {
        border-radius: 16px;
        box-shadow: 0px 4px 8px rgba(0,0,0,0.15);
        transition: all 0.3s ease-in-out;
    }
    
    /* Animación suave en botones al pasar el ratón */
    div.stButton > button:first-child:hover {
        background-color: #33c8bd; /* variante oscura del primaryColor */
        transform: translateY(-2px);
    }
    
    /* Estilo en los widgets del formulario */
    div[data-baseweb="input"], div[data-baseweb="textarea"], div[data-baseweb="select"] {
        border-radius: 16px;
        box-shadow: 0px 4px 6px rgba(0,0,0,0.1);
    }
    
    /* Estilo para las tarjetas y contenedores */
    .css-1r6slb0, .css-12w0qpk {
        border-radius: 16px;
        border: 1px solid #e0e0e0;
    }
    
    /* Mejoras en pestañas */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0px 0px;
        padding: 10px 16px;
        background-color: transparent;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #3FF2E1;
        color: #1F1F1F;
    }
    
    /* Personalización de los expanders */
    .streamlit-expanderHeader {
        border-radius: 8px;
        font-weight: 500;
    }
    
    /* Personalización de las métricas */
    [data-testid="stMetricValue"] {
        font-weight: 600;
        color: #1F1F1F;
    }
    
    /* Estilos para el pie de página personalizado */
    .footer {
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        background-color: #f8f9fa;
        padding: 10px 0;
        text-align: center;
        font-size: 0.8rem;
        color: #666666;
        border-top: 1px solid #eeeeee;
        z-index: 999;
    }
    
    .footer-content {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        gap: 5px;
    }
    
    .footer-logo {
        max-height: 50px; /* Triplicado de 30px a 90px */
        margin-bottom: 5px;
    }
    
    /* Añadir padding al final del contenido principal para evitar que el footer oculte contenido */
    .main .block-container {
        padding-bottom: 70px; /* Aumentado para dar espacio al logo más grande */
    }
    </style>
    """, unsafe_allow_html=True)


def add_footer_local():
    """
    Añade un pie de página minimalista con el logo local (tamaño triple), copyright y texto de Spanish FactorIA.
    """
    # Cargar el logo desde assets si existe
    logo_path = os.path.join(ASSETS_PATH, "Spanish_FactorIA_Logo.png")
    logo_base64 = ""

    try:
        if os.path.exists(logo_path):
            with open(logo_path, "rb") as img_file:
                logo_bytes = img_file.read()
                logo_base64 = base64.b64encode(logo_bytes).decode("utf-8")
    except Exception as e:
        logger.error(f"Error al cargar logo para footer: {e}")

    if logo_base64:
        logo_img = f'<img src="data:image/png;base64,{logo_base64}" alt="Spanish FactorIA Logo" class="footer-logo">'
    else:
        # Texto alternativo si no hay logo, con tamaño aumentado
        logo_img = '<span style="font-weight: bold; font-size: 2rem;">Spanish FactorIA</span>'

    footer = f"""
    <div class="footer">
        <div class="footer-content">
            {logo_img}
            <p>© 2025 Spanish FactorIA | Herramientas para estudiantes de español</p>
        </div>
    </div>
    """
    st.markdown(footer, unsafe_allow_html=True)


# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 2: Inicialización de Session State y Estados Internos
# ==================================================================================
#
# Este artefacto contiene:
# 1. Funciones para inicializar variables en el session_state de Streamlit
# 2. Funciones seguras para acceder y modificar valores en el session_state
# 3. Generación de UID para usuarios y manejo de estados de sesión
# 4. Inicialización del sidebar con información básica y branding
#
# Es crucial para garantizar que la aplicación mantenga su estado entre recargas
# y para proporcionar una forma segura y consistente de acceder a los datos de sesión.
# ==================================================================================

def generate_user_uid(email: str = None) -> str:
    """
    Genera un identificador único para el usuario basado en su correo o un UUID aleatorio.

    Args:
        email: Correo electrónico del usuario (opcional)

    Returns:
        str: UID del usuario
    """
    try:
        if not email:
            return str(uuid.uuid4())

        # Crear un hash determinista a partir del email para usarlo como UID
        email_hash = hashlib.sha256(email.lower().strip().encode()).hexdigest()
        return f"user_{email_hash[:16]}"
    except Exception:
        # Método alternativo si hay problemas con hashlib
        if email:
            safe_email = email.replace('@', '_at_').replace('.', '_dot_')
            return f"user_{safe_email}"
        return f"user_{str(uuid.uuid4())}"


def init_session_state():
    """
    Inicializa variables de session_state con valores predeterminados seguros
    para evitar KeyError durante la ejecución. Incluye mejor manejo de errores.
    """
    try:
        default_values = {
            # Información de usuario y autenticación
            "usuario_actual": "",
            "email_usuario": "",
            "uid_usuario": "",
            "nivel_estudiante": "intermedio",
            "fecha_inicio_sesion": datetime.now().isoformat(),
            "is_authenticated": False,

            # Estados de navegación - Actualizados para la nueva estructura
            "active_tab": 0,  # Índice de la pestaña activa
            "mostrar_exportacion": False,
            "mostrar_login": True,  # Control de la página de bienvenida

            # Estados de corrección
            "consigna_actual": "",
            "usar_consigna_como_texto": False,
            "texto_correccion_corregir": "",
            "info_adicional_corregir": "",
            "ultimo_texto": "",
            "correction_result": None,
            "last_correction_time": None,
            "last_correction_data": None,

            # Estados para examen
            "examen_result": None,
            "inicio_simulacro": None,
            "duracion_simulacro": None,
            "tarea_simulacro": None,
            "simulacro_respuesta_texto": "",

            # Estados para herramientas
            "ultima_imagen_url": "",
            "ultima_descripcion": "",
            "ultimo_texto_transcrito": "",
            "imagen_generada_state": False,
            "imagen_url_state": None,
            "descripcion_state": None,
            "tema_imagen_state": None,
            "descripcion_estudiante_state": "",
            "mostrar_correccion_imagen": False,
            "mostrar_correccion_transcripcion": False,

            # Estados para APIs
            "api_error_count": 0,
            "api_last_error_time": None,
            "circuit_breaker_open": False,

            # Estado de inicialización
            "app_initialized": False,
            "openai_available": False,  # Nueva bandera para OpenAI
            "firebase_available": False,

            # Flags de navegación explícitos para herramientas
            "active_tools_tab": 0,

            # Estados para modo degradado
            "modo_degradado": False,
            "servicios_disponibles": {
                "openai": False,
                "firebase": False,
                "elevenlabs": False,
                "dalle": False
            },

            # Nuevos estados para el perfil de estudiante
            "metricas_modelos": {},  # Para almacenar métricas de rendimiento de modelos
            "priorizar_costo": True,  # Por defecto prioriza modelos económicos
            "mostrar_perfil": False,  # Control para mostrar vista del perfil
            "historial_correcciones": None,  # Caché de historial para perfil
            "ultima_area_mejora": None,  # Última área de mejora identificada
        }

        # Inicializar valores por defecto solo si no existen
        for key, default_value in default_values.items():
            if key not in st.session_state:
                st.session_state[key] = default_value

        # Generar ID único para esta sesión si no existe
        if "session_id" not in st.session_state:
            st.session_state.session_id = str(uuid.uuid4())

        # Generar UID de usuario si existe email pero no UID
        if st.session_state.get("email_usuario") and not st.session_state.get("uid_usuario"):
            st.session_state.uid_usuario = generate_user_uid(
                st.session_state.email_usuario)

    except Exception as e:
        # Asegurar valores mínimos incluso si hay error
        print(f"Error en init_session_state: {str(e)}")
        if "is_authenticated" not in st.session_state:
            st.session_state["is_authenticated"] = False
        if "mostrar_login" not in st.session_state:
            st.session_state["mostrar_login"] = True
        if "session_id" not in st.session_state:
            st.session_state["session_id"] = str(uuid.uuid4())


def get_session_var(key: str, default: Any = None) -> Any:
    """
    Obtiene una variable de session_state de forma segura.

    Args:
        key: Clave de la variable en session_state
        default: Valor por defecto si la clave no existe

    Returns:
        Valor almacenado o valor por defecto
    """
    try:
        return st.session_state.get(key, default)
    except Exception:
        # En caso de cualquier error, retornar el valor por defecto
        return default


def set_session_var(key: str, value: Any) -> None:
    """
    Establece una variable en session_state de forma segura.

    Args:
        key: Clave de la variable en session_state
        value: Valor a almacenar
    """
    try:
        st.session_state[key] = value
    except Exception as e:
        print(f"Error al establecer variable de sesión {key}: {str(e)}")


def show_sidebar_info():
    """
    Muestra información en el sidebar de la aplicación, incluyendo
    el logo, información de la versión y datos de la sesión.
    """
    # Logo en sidebar
    try:
        if os.path.exists(LOGO_PATH):
            logo_img = Image.open(LOGO_PATH)
            st.sidebar.image(logo_img, width=200, caption="")
        else:
            st.sidebar.info(f"Logo no encontrado en: {LOGO_PATH}")
    except Exception as e:
        print(f"No se pudo cargar el logo: {e}")

    # Título e información
    st.sidebar.title(f"📝 {APP_NAME}")
    st.sidebar.info(
        f"""
        Versión: {APP_VERSION}
        
        Una herramienta de corrección de textos
        en español con análisis contextual avanzado.
        
        Powered by {ORGANIZATION}
        
        ID de sesión: {get_session_var("session_id", "")[:8]}
        """
    )

    # Si el usuario está autenticado, mostrar su información
    if get_session_var("is_authenticated", False):
        usuario = get_session_var("usuario_actual", "")
        email = get_session_var("email_usuario", "")
        if usuario and email:
            st.sidebar.success(f"Usuario: {usuario} ({email})")

            # Botón para cerrar sesión
            if st.sidebar.button("Cerrar sesión"):
                # Reiniciar variables de sesión relacionadas con el usuario
                set_session_var("usuario_actual", "")
                set_session_var("email_usuario", "")
                set_session_var("uid_usuario", "")
                set_session_var("is_authenticated", False)
                set_session_var("mostrar_login", True)
                st.rerun()

            # Alternancia de prioridad de modelos (solo para administradores)
            if email.endswith('@spanishfactoria.com') or email == 'admin@example.com':
                priorizar_costo = get_session_var("priorizar_costo", True)
                st.sidebar.write("⚙️ Configuración avanzada:")
                if st.sidebar.checkbox("Priorizar modelos económicos", value=priorizar_costo):
                    set_session_var("priorizar_costo", True)
                else:
                    set_session_var("priorizar_costo", False)


# Inicializar session_state al cargar este artefacto
init_session_state()

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 3: Conexiones Seguras a APIs
# ==================================================================================
#
# Este artefacto contiene:
# 1. Implementación del patrón Circuit Breaker para APIs
# 2. Funciones para manejo seguro de claves API
# 3. Conexión segura a OpenAI con detección de modelos
# 4. Funciones de diagnóstico para el estado de conexiones
#
# Es crucial para establecer conexiones seguras y robustas con servicios externos,
# gestionar fallos de manera adecuada, y proporcionar mecanismos para la recuperación.
# Se ha mejorado el manejo de timeouts y reintentos para mayor fiabilidad.
# ==================================================================================


class CircuitBreaker:
    """
    Implementa el patrón Circuit Breaker para APIs externas.
    Previene llamadas repetidas a APIs con fallo.
    """

    def __init__(self, failure_threshold=5, reset_timeout=300):
        self.failure_threshold = failure_threshold  # Número de fallos antes de abrir
        self.reset_timeout = reset_timeout  # Tiempo en segundos antes de reintentar

        # Inicializar contadores para diferentes servicios
        self.services = {
            "openai": {"failures": 0, "last_failure_time": None, "open": False, "error_types": {}},
            "firebase": {"failures": 0, "last_failure_time": None, "open": False, "error_types": {}},
            "elevenlabs": {"failures": 0, "last_failure_time": None, "open": False, "error_types": {}},
            "dalle": {"failures": 0, "last_failure_time": None, "open": False, "error_types": {}}
        }

    def record_failure(self, service_name, error_type="generic"):
        """
        Registra un fallo para el servicio especificado, clasificado por tipo de error.

        Args:
            service_name: Nombre del servicio que falló
            error_type: Tipo de error (timeout, rate_limit, auth, etc.)
        """
        if service_name not in self.services:
            logger.warning(f"Servicio desconocido: {service_name}")
            return

        service = self.services[service_name]
        service["failures"] += 1
        service["last_failure_time"] = time.time()

        # Registrar el tipo de error para análisis
        if error_type not in service["error_types"]:
            service["error_types"][error_type] = 0
        service["error_types"][error_type] += 1

        # Abrir el circuit breaker si se alcanza el umbral de fallos
        if service["failures"] >= self.failure_threshold:
            service["open"] = True
            logger.warning(
                f"Circuit breaker ABIERTO para {service_name} (error: {error_type})")

            # Si es un error de tiempo de espera, podríamos establecer un umbral más bajo
            if error_type == "timeout" and service["error_types"].get("timeout", 0) >= 3:
                logger.warning(
                    f"Múltiples errores de timeout detectados para {service_name}")

    def record_success(self, service_name):
        """
        Registra un éxito y restablece contadores para el servicio

        Args:
            service_name: Nombre del servicio
        """
        if service_name not in self.services:
            return

        service = self.services[service_name]
        service["failures"] = 0
        service["open"] = False
        # Mantenemos el historial de tipos de error para análisis

    def can_execute(self, service_name):
        """
        Determina si se puede ejecutar una llamada al servicio

        Args:
            service_name: Nombre del servicio a verificar

        Returns:
            bool: True si se puede ejecutar, False si está bloqueado
        """
        if service_name not in self.services:
            return True

        service = self.services[service_name]

        # Si el circuit breaker está abierto, verificar si ha pasado el tiempo de reset
        if service["open"]:
            if service["last_failure_time"] is None:
                return True

            elapsed = time.time() - service["last_failure_time"]
            if elapsed > self.reset_timeout:
                # Permitir un reintento
                service["open"] = False
                return True
            else:
                return False

        return True

    def attempt_reset(self, service_name):
        """
        Intenta restablecer el circuito para un servicio específico

        Args:
            service_name: Nombre del servicio

        Returns:
            bool: True si se restableció, False en caso contrario
        """
        if service_name in self.services and self.services[service_name]["open"]:
            self.services[service_name]["failures"] = max(
                0, self.services[service_name]["failures"] - 1)
            if self.services[service_name]["failures"] <= self.failure_threshold // 2:
                self.services[service_name]["open"] = False
                logger.info(
                    f"Circuit breaker RESTABLECIDO para {service_name}")
                return True
        return False

    def get_status(self):
        """
        Devuelve el estado actual de todos los servicios

        Returns:
            dict: Estado actual de todos los servicios
        """
        return {name: {"open": info["open"], "failures": info["failures"], "error_types": info["error_types"]}
                for name, info in self.services.items()}


# Inicializar circuit breaker
circuit_breaker = CircuitBreaker()


def get_api_keys():
    """
    Obtiene las claves de API de los secretos de Streamlit con manejo de errores.
    Permite la operación en modo degradado si faltan claves.

    Returns:
        dict: Diccionario con las claves de API y credenciales
    """
    keys = {
        "openai": None,
        "elevenlabs": {"api_key": None, "voice_id": None},
        "firebase_credentials": None,
        "dalle": None
    }

    try:
        # Intentamos obtener la API key de OpenAI
        keys["openai"] = st.secrets["OPENAI_API_KEY"]
        # Reutilizamos la misma API key para DALLE
        keys["dalle"] = keys["openai"]
        set_session_var("openai_available", True)
    except Exception as e:
        logger.warning(f"Error al obtener API Key de OpenAI: {e}")
        st.sidebar.warning(
            "⚠️ API de OpenAI no configurada. Algunas funciones estarán limitadas.")
        set_session_var("openai_available", False)

    try:
        keys["elevenlabs"]["api_key"] = st.secrets["ELEVENLABS_API_KEY"]
        keys["elevenlabs"]["voice_id"] = st.secrets["ELEVENLABS_VOICE_ID"]
    except Exception as e:
        logger.warning(f"Error al obtener configuración de ElevenLabs: {e}")
        st.sidebar.warning(
            "⚠️ API de ElevenLabs no configurada. La función de audio estará deshabilitada.")

    try:
        keys["firebase_credentials"] = json.loads(
            st.secrets["FIREBASE_CREDENTIALS"])
        set_session_var("firebase_available", True)
    except Exception as e:
        logger.warning(f"Error al obtener credenciales de Firebase: {e}")
        st.sidebar.warning(
            "⚠️ Credenciales de Firebase no configuradas. El guardado de datos estará deshabilitado.")
        set_session_var("firebase_available", False)

    return keys


# Obtener claves de API
api_keys = get_api_keys()


def retry_with_backoff(func, max_retries=3, initial_delay=1, max_delay=60):
    """
    Ejecuta una función con reintentos y backoff exponencial.
    Mejorado para manejar diferentes tipos de errores.

    Args:
        func: Función a ejecutar
        max_retries: Número máximo de reintentos
        initial_delay: Retraso inicial en segundos
        max_delay: Retraso máximo en segundos

    Returns:
        El resultado de la función o levanta la excepción
    """
    for attempt in range(max_retries):
        try:
            return func()
        except (requests.ConnectionError, requests.Timeout) as e:
            # Errores de red específicos - reintentamos
            if attempt == max_retries - 1:
                raise

            # Calcular delay con backoff exponencial limitado
            delay = min(initial_delay * (2 ** attempt), max_delay)

            # Log más detallado
            if isinstance(e, requests.Timeout):
                logger.warning(
                    f"Timeout en intento {attempt+1}/{max_retries}. Reintentando en {delay} segundos.")
            else:
                logger.warning(
                    f"Error de conexión en intento {attempt+1}/{max_retries}. Reintentando en {delay} segundos.")

            time.sleep(delay)
        except requests.HTTPError as e:
            # Para errores HTTP, verificar el código
            status_code = e.response.status_code if hasattr(
                e, 'response') else 0

            # Para rate limits (429), siempre reintentar con backoff
            if status_code == 429 and attempt < max_retries - 1:
                delay = min(initial_delay * (2 ** attempt) * 2,
                            max_delay)  # Backoff más agresivo
                logger.warning(
                    f"Rate limit (429) alcanzado. Reintentando en {delay} segundos.")
                time.sleep(delay)
            elif 500 <= status_code < 600 and attempt < max_retries - 1:
                # Errores de servidor, reintentar
                delay = min(initial_delay * (2 ** attempt), max_delay)
                logger.warning(
                    f"Error del servidor ({status_code}). Reintentando en {delay} segundos.")
                time.sleep(delay)
            else:
                # Otros errores HTTP, no reintentar
                logger.error(f"Error HTTP no recuperable: {status_code}")
                raise
        except Exception as e:
            # Otros errores - no reintentamos
            logger.error(f"Error no recuperable: {str(e)}")
            raise


# --- CONFIGURACIÓN DE OPENAI ---

def list_available_openai_models():
    """
    Lista todos los modelos disponibles de OpenAI.

    Returns:
        list: Lista de nombres de modelos disponibles o lista por defecto
    """
    if api_keys["openai"] is None:
        logger.warning("API key de OpenAI no configurada")
        return []

    if not circuit_breaker.can_execute("openai"):
        logger.warning("Circuit breaker abierto para OpenAI")
        return []

    try:
        # Configurar la API key para OpenAI
        headers = {
            "Authorization": f"Bearer {api_keys['openai']}",
            "Content-Type": "application/json"
        }

        # Función para ejecutar la solicitud con reintentos
        def fetch_models():
            response = requests.get(
                "https://api.openai.com/v1/models",
                headers=headers,
                timeout=10  # Timeout reducido para esta solicitud informativa
            )
            response.raise_for_status()
            return response.json()

        # Ejecutar con reintentos
        models_data = retry_with_backoff(fetch_models, max_retries=2)

        # Filtrar modelos GPT
        gpt_models = [model["id"] for model in models_data.get("data", [])
                      if "gpt" in model["id"].lower()]

        logger.info(f"Modelos OpenAI disponibles: {gpt_models}")

        # Registrar éxito
        circuit_breaker.record_success("openai")

        return gpt_models
    except Exception as e:
        logger.error(f"Error al listar modelos OpenAI: {e}")
        circuit_breaker.record_failure("openai", error_type="list_models")
        # Devolver lista por defecto en caso de error
        return ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"]


def get_best_openai_model(priorizar_costo=True):
    """
    Determina el mejor modelo OpenAI disponible basado en prioridades.
    Puede priorizar costo o capacidad según necesidades.

    Args:
        priorizar_costo: Si es True, prioriza modelos más económicos.

    Returns:
        str: Nombre del mejor modelo disponible o un modelo predeterminado
    """
    if priorizar_costo:
        # Lista de modelos preferidos en orden de prioridad (primero los más económicos)
        preferred_models = [
            "gpt-3.5-turbo",
            "gpt-3.5-turbo-16k",
            "gpt-4-turbo",
            "gpt-4"
        ]
    else:
        # Lista de modelos preferidos en orden de capacidad (primero los más potentes)
        preferred_models = [
            "gpt-4",
            "gpt-4-turbo",
            "gpt-3.5-turbo",
            "gpt-3.5-turbo-16k"
        ]

    # Obtener modelos disponibles
    available_models = list_available_openai_models()

    # Si no hay modelos disponibles, usar un modelo predeterminado
    if not available_models:
        logger.warning(
            "No se encontraron modelos OpenAI disponibles, usando valor predeterminado")
        return "gpt-3.5-turbo"  # Modelo predeterminado

    # Buscar el primer modelo preferido que esté disponible
    for model in preferred_models:
        if any(model in available_model for available_model in available_models):
            return model

    # Si ninguno de los preferidos está disponible, usar el primero de la lista
    return available_models[0]


def configure_openai():
    """
    Configura el cliente de OpenAI con la API key y verifica la conexión.

    Returns:
        tuple: (modelo_seleccionado, éxito_configuración)
    """
    if api_keys["openai"] is None:
        logger.warning("API key de OpenAI no configurada")
        return None, False

    if not circuit_breaker.can_execute("openai"):
        logger.warning("Circuit breaker abierto para OpenAI")
        return None, False

    try:
        # Obtener preferencia de priorización de costos
        priorizar_costo = get_session_var("priorizar_costo", True)

        # Seleccionar mejor modelo según preferencia
        best_model_name = get_best_openai_model(priorizar_costo)
        logger.info(f"Usando modelo OpenAI: {best_model_name}")

        # Verificar que la configuración funcione haciendo una llamada simple
        headers = {
            "Authorization": f"Bearer {api_keys['openai']}",
            "Content-Type": "application/json"
        }

        data = {
            "model": best_model_name,
            "messages": [{"role": "user", "content": "Hello, testing OpenAI connection."}],
            "max_tokens": 10
        }

        # Función para verificar la conexión con reintentos
        def test_connection():
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=15  # Timeout para conexión de prueba
            )
            response.raise_for_status()
            return response.json()

        # Probar conexión con reintentos
        response_data = retry_with_backoff(test_connection, max_retries=2)

        # Si llegamos aquí, la configuración fue exitosa
        circuit_breaker.record_success("openai")
        set_session_var("servicios_disponibles", {
            **get_session_var("servicios_disponibles", {}),
            "openai": True
        })

        # Registrar información del modelo seleccionado
        logger.info(f"Conexión a OpenAI exitosa con modelo {best_model_name}")

        return best_model_name, True
    except requests.Timeout as e:
        logger.error(f"Timeout al configurar OpenAI: {e}")
        circuit_breaker.record_failure("openai", error_type="timeout")
        return None, False
    except requests.HTTPError as e:
        status_code = e.response.status_code if hasattr(e, 'response') else 0
        logger.error(f"Error HTTP {status_code} al configurar OpenAI: {e}")
        circuit_breaker.record_failure(
            "openai", error_type=f"http_{status_code}")
        return None, False
    except Exception as e:
        logger.error(f"Error al configurar OpenAI: {e}")
        circuit_breaker.record_failure("openai", error_type="generic")
        return None, False


# --- DIAGNÓSTICO DE CONEXIONES ---

def show_connection_status():
    """Muestra el estado de conexión de los servicios externos en el sidebar."""
    with st.sidebar.expander("Estado de conexiones", expanded=False):
        status = circuit_breaker.get_status()

        # OpenAI
        if api_keys["openai"] is not None:
            if not status["openai"]["open"]:
                st.sidebar.success("✅ OpenAI: Conectado")
                # Mostrar el modelo actual
                modelo_actual, _ = configure_openai()
                if modelo_actual:
                    st.sidebar.info(f"Modelo actual: {modelo_actual}")
            else:
                error_types = status["openai"]["error_types"]
                error_msg = ", ".join(
                    [f"{tipo}: {count}" for tipo, count in error_types.items()])
                st.sidebar.error(
                    f"❌ OpenAI: Desconectado ({status['openai']['failures']} fallos - {error_msg})")
                # Añadir botón para reintentar
                if st.sidebar.button("Reintentar OpenAI"):
                    if circuit_breaker.attempt_reset("openai"):
                        st.sidebar.info("Reintentando conexión con OpenAI...")
                        configure_openai()
                        st.rerun()
        else:
            st.sidebar.warning("⚠️ OpenAI: No configurado")

        # Firebase
        if api_keys["firebase_credentials"] is not None:
            if not status["firebase"]["open"]:
                st.sidebar.success("✅ Firebase: Conectado")
            else:
                error_types = status["firebase"]["error_types"]
                error_msg = ", ".join(
                    [f"{tipo}: {count}" for tipo, count in error_types.items()])
                st.sidebar.error(
                    f"❌ Firebase: Desconectado ({status['firebase']['failures']} fallos - {error_msg})")
        else:
            st.sidebar.warning("⚠️ Firebase: No configurado")

        # ElevenLabs
        if api_keys["elevenlabs"]["api_key"] is not None:
            if not status["elevenlabs"]["open"]:
                st.sidebar.success("✅ ElevenLabs: Conectado")
            else:
                error_types = status["elevenlabs"]["error_types"]
                error_msg = ", ".join(
                    [f"{tipo}: {count}" for tipo, count in error_types.items()])
                st.sidebar.error(
                    f"❌ ElevenLabs: Desconectado ({status['elevenlabs']['failures']} fallos - {error_msg})")
        else:
            st.sidebar.warning("⚠️ ElevenLabs: No configurado")

        # DALL-E
        if api_keys["dalle"] is not None:
            if not status["dalle"]["open"]:
                st.sidebar.success("✅ DALL-E: Conectado")
            else:
                error_types = status["dalle"]["error_types"]
                error_msg = ", ".join(
                    [f"{tipo}: {count}" for tipo, count in error_types.items()])
                st.sidebar.error(
                    f"❌ DALL-E: Desconectado ({status['dalle']['failures']} fallos - {error_msg})")
        else:
            st.sidebar.warning("⚠️ DALL-E: No configurado")


# Función para diagnóstico proactivo de problemas
def diagnosticar_aplicacion():
    """
    Diagnostica problemas comunes en la aplicación.

    Returns:
        list: Lista de problemas detectados
    """
    problemas = []

    # Verificar conexión a OpenAI
    if api_keys["openai"] is None:
        problemas.append({
            "tipo": "crítico",
            "mensaje": "API Key de OpenAI no configurada",
            "solucion": "Configura la API Key de OpenAI en los secretos de la aplicación"
        })
    else:
        # Verificar si podemos conectar efectivamente
        modelo, conectado = configure_openai()
        if not conectado:
            problemas.append({
                "tipo": "crítico",
                "mensaje": "No se puede conectar con la API de OpenAI",
                "solucion": "Verifica que la API Key de OpenAI sea válida y que el servicio esté disponible"
            })

    # Verificar conexión a Firebase
    if api_keys["firebase_credentials"] is None:
        problemas.append({
            "tipo": "advertencia",
            "mensaje": "Credenciales de Firebase no configuradas",
            "solucion": "Configura las credenciales de Firebase en los secretos de la aplicación para habilitar el guardado de datos"
        })

    # Verificar si hay archivos necesarios
    if not os.path.exists(LOGO_PATH):
        problemas.append({
            "tipo": "advertencia",
            "mensaje": f"Logo no encontrado en {LOGO_PATH}",
            "solucion": "Asegúrate de tener el logo en la carpeta assets"
        })

    # Verificar bibliotecas necesarias
    required_modules = ["docx", "requests", "qrcode"]
    for module_name in required_modules:
        try:
            __import__(module_name.split(".")[0])
        except ImportError:
            problemas.append({
                "tipo": "crítico" if module_name == "requests" else "advertencia",
                "mensaje": f"Biblioteca {module_name} no instalada",
                "solucion": f"Instala {module_name} con pip install {module_name.split('.')[0]}"
            })

    return problemas


# Realizar diagnóstico inicial al cargar este artefacto
diagnostico_inicial = diagnosticar_aplicacion()
for problema in diagnostico_inicial:
    if problema["tipo"] == "crítico":
        logger.error(
            f"Problema crítico detectado: {problema['mensaje']}. Solución: {problema['solucion']}")
    else:
        logger.warning(
            f"Advertencia: {problema['mensaje']}. Solución: {problema['solucion']}")

# Determinar si estamos en modo degradado basado en disponibilidad de servicios
set_session_var("modo_degradado", not get_session_var(
    "openai_available", False))

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 4: Funciones de Firebase
# ==================================================================================
#
# Este artefacto contiene:
# 1. Funciones para inicializar Firebase
# 2. Funciones para guardar y recuperar datos de Firestore
# 3. Manejo de errores y fallbacks para operaciones con Firebase
#
# Estas funciones gestionan la persistencia de datos en la nube, permitiendo
# guardar el historial de correcciones y datos de usuarios.
# ==================================================================================


def initialize_firebase():
    """
    Inicializa la conexión con Firebase usando las credenciales proporcionadas.

    Returns:
        tuple: (firestore_db, success_flag)
    """
    if api_keys["firebase_credentials"] is None:
        logger.warning("Credenciales de Firebase no configuradas")
        return None, False

    if not circuit_breaker.can_execute("firebase"):
        logger.warning("Circuit breaker abierto para Firebase")
        return None, False

    try:
        # Intentar importar Firebase aquí para manejar la importación condicional
        try:
            import firebase_admin
            from firebase_admin import credentials, firestore
        except ImportError as e:
            logger.error(f"Error al importar firebase_admin: {e}")
            return None, False

        # Verificar si Firebase ya está inicializado
        if not firebase_admin._apps:
            # Crear un objeto de credenciales desde el diccionario
            cred = credentials.Certificate(api_keys["firebase_credentials"])

            # Inicializar la aplicación de Firebase
            firebase_admin.initialize_app(cred)

        # Obtener una referencia a la base de datos Firestore
        db = firestore.client()

        # Hacer una prueba simple para verificar la conexión
        test_ref = db.collection('test').document('connection')
        test_ref.set({'timestamp': firestore.SERVER_TIMESTAMP})

        # Si llegamos aquí, la conexión fue exitosa
        circuit_breaker.record_success("firebase")
        set_session_var("servicios_disponibles", {
            **get_session_var("servicios_disponibles", {}),
            "firebase": True
        })
        return db, True
    except Exception as e:
        logger.error(f"Error al inicializar Firebase: {e}")
        circuit_breaker.record_failure("firebase", error_type="initialization")
        return None, False


def guardar_correccion_firestore(nombre, nivel, idioma, texto, resultado_json):
    """
    Guarda los datos de una corrección en Firestore con métricas adicionales.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel de español
        idioma: Idioma de corrección
        texto: Texto original
        resultado_json: Resultado de la corrección en formato JSON (string o dict)

    Returns:
        dict: Resultado de la operación
    """
    try:
        # Intentar importar Firebase en forma segura
        try:
            import firebase_admin
            from firebase_admin import firestore
        except ImportError:
            return {"success": False, "message": "Firebase no está disponible"}

        db, conexion_ok = initialize_firebase()
        if not conexion_ok or db is None:
            return {"success": False, "message": "Conexión a Firestore no disponible"}

        # Obtener UID del usuario
        uid = get_session_var("uid_usuario", "")
        if not uid:
            # Generar un UID temporal basado en la sesión
            uid = f"temp_{get_session_var('session_id', str(uuid.uuid4()))}"

        # Fecha actual para el registro
        fecha = datetime.now().isoformat()

        # Convertir resultado_json a string si es un diccionario
        if isinstance(resultado_json, dict):
            raw_output = json.dumps(resultado_json)
        else:
            raw_output = resultado_json

        # Generar ID único para la corrección
        correccion_id = str(uuid.uuid4())

        # Preparar los datos para guardar
        datos_correccion = {
            "id": correccion_id,
            "nombre": nombre,
            "nivel": nivel,
            "idioma": idioma,
            "fecha": fecha,
            "texto_original": texto,
            "resultado_raw": raw_output,
            "timestamp": firestore.SERVER_TIMESTAMP
        }

        # Extraer estadísticas para guardar en formato estructurado
        try:
            if isinstance(resultado_json, str):
                data_json = extract_json_safely(resultado_json)
            else:
                data_json = resultado_json

            # Extraer datos relevantes
            errores = data_json.get("errores", {})
            analisis = data_json.get("analisis_contextual", {})

            # Contar errores por categoría
            stats = {
                "errores_gramatica": len(errores.get("Gramática", [])),
                "errores_lexico": len(errores.get("Léxico", [])),
                "errores_puntuacion": len(errores.get("Puntuación", [])),
                "errores_estructura": len(errores.get("Estructura textual", [])),
            }

            # Total de errores
            stats["total_errores"] = (
                stats["errores_gramatica"] +
                stats["errores_lexico"] +
                stats["errores_puntuacion"] +
                stats["errores_estructura"]
            )

            # Extraer puntuaciones
            coherencia = analisis.get("coherencia", {})
            cohesion = analisis.get("cohesion", {})
            registro = analisis.get("registro_linguistico", {})
            adecuacion = analisis.get("adecuacion_cultural", {})

            stats["puntuacion_coherencia"] = coherencia.get("puntuacion", 0)
            stats["puntuacion_cohesion"] = cohesion.get("puntuacion", 0)
            stats["puntuacion_registro"] = registro.get("puntuacion", 0)
            stats["puntuacion_adecuacion"] = adecuacion.get("puntuacion", 0)

            # Calcular puntuación global
            stats["puntuacion_global"] = (
                stats["puntuacion_coherencia"] +
                stats["puntuacion_cohesion"] +
                stats["puntuacion_registro"] +
                stats["puntuacion_adecuacion"]
            ) / 4

            # Métricas adicionales
            stats["longitud_texto"] = len(texto.split())
            stats["nivel_complejidad"] = calcular_nivel_complejidad(
                texto) if 'calcular_nivel_complejidad' in globals() else 0
            stats["indice_szigriszt"] = calcular_indice_szigriszt(
                texto) if 'calcular_indice_szigriszt' in globals() else 0

            # Extraer consejo final
            stats["consejo_final"] = data_json.get("consejo_final", "")

            # Añadir estadísticas a los datos
            datos_correccion["stats"] = stats

            # Añadir información del modelo usado
            modelo_usado, _ = configure_openai()
            datos_correccion["modelo_usado"] = modelo_usado if modelo_usado else "desconocido"

        except Exception as stats_error:
            logger.error(f"Error al extraer estadísticas: {str(stats_error)}")
            datos_correccion["stats_error"] = str(stats_error)

        # Guardar en Firestore
        db.collection('usuarios').document(uid).collection(
            'correcciones').document(correccion_id).set(datos_correccion)

        return {"success": True, "message": "Datos guardados correctamente.", "id": correccion_id}

    except Exception as e:
        logger.error(f"Error al guardar en Firestore: {str(e)}")
        circuit_breaker.record_failure(
            "firebase", error_type="save_correction")
        return {"success": False, "message": f"Error al guardar datos: {str(e)}"}


def obtener_historial_estudiante(uid=None, email=None, nombre=None):
    """
    Obtiene el historial de correcciones para un estudiante específico.
    Optimizado con caché para reducir consultas a Firebase.

    Args:
        uid: UID del usuario (opcional)
        email: Email del usuario (opcional)
        nombre: Nombre del estudiante (opcional)

    Returns:
        pd.DataFrame o None: DataFrame con historial o None si no hay datos
    """
    try:
        # Verificar si tenemos historial en caché
        historial_cache = get_session_var("historial_correcciones", None)

        # Si tenemos caché y no se especificaron parámetros diferentes a los usados anteriormente
        cached_uid = get_session_var("cached_historial_uid", None)
        cached_email = get_session_var("cached_historial_email", None)
        cached_nombre = get_session_var("cached_historial_nombre", None)

        # Verificar si podemos usar la caché
        usar_cache = (
            historial_cache is not None and
            cached_uid == uid and
            cached_email == email and
            cached_nombre == nombre
        )

        if usar_cache:
            return historial_cache

        # Importar Firebase en forma segura
        try:
            import firebase_admin
            from firebase_admin import firestore
        except ImportError:
            logger.warning(
                "Firebase no está disponible para obtener historial")
            return None

        db, conexion_ok = initialize_firebase()
        if not conexion_ok or db is None:
            logger.warning("No hay conexión con Firestore")
            return None

        # Determinar el UID a usar
        user_uid = uid

        # Si no hay UID pero hay email, generar el UID
        if not user_uid and email:
            user_uid = generate_user_uid(email)

        # Si no hay UID ni email, usar el UID de la sesión actual
        if not user_uid:
            user_uid = get_session_var("uid_usuario", "")

        # Si aún no hay UID, no podemos continuar
        if not user_uid:
            logger.warning(
                "No se proporcionó UID del usuario para obtener historial")
            return None

        # Obtener todas las correcciones del usuario
        correcciones_ref = db.collection('usuarios').document(
            user_uid).collection('correcciones')

        # Si se especifica un nombre, filtrar por él
        # Usamos where() en lugar de filter() para compatibilidad con versiones anteriores
        if nombre:
            correcciones_ref = correcciones_ref.where('nombre', '==', nombre)

        # Obtener los documentos ordenados por fecha
        correcciones = correcciones_ref.order_by('fecha').get()

        # Si no hay documentos, retornar None
        if not correcciones:
            return None

        # Crear una lista para los datos
        datos = []

        # Procesar cada corrección
        for doc in correcciones:
            correc = doc.to_dict()

            # Extraer los datos que necesitamos
            entrada = {
                'ID': doc.id,
                'Nombre': correc.get('nombre', ''),
                'Nivel': correc.get('nivel', ''),
                'Fecha': correc.get('fecha', ''),
                'Modelo': correc.get('modelo_usado', 'desconocido'),
            }

            # Añadir estadísticas si están disponibles
            stats = correc.get('stats', {})
            if stats:
                entrada['Errores Gramática'] = stats.get(
                    'errores_gramatica', 0)
                entrada['Errores Léxico'] = stats.get('errores_lexico', 0)
                entrada['Errores Puntuación'] = stats.get(
                    'errores_puntuacion', 0)
                entrada['Errores Estructura'] = stats.get(
                    'errores_estructura', 0)
                entrada['Total Errores'] = stats.get('total_errores', 0)
                entrada['Puntuación Coherencia'] = stats.get(
                    'puntuacion_coherencia', 0)
                entrada['Puntuación Cohesión'] = stats.get(
                    'puntuacion_cohesion', 0)
                entrada['Puntuación Registro'] = stats.get(
                    'puntuacion_registro', 0)
                entrada['Puntuación Adecuación Cultural'] = stats.get(
                    'puntuacion_adecuacion', 0)
                entrada['Puntuación Global'] = stats.get(
                    'puntuacion_global', 0)
                entrada['Longitud Texto'] = stats.get('longitud_texto', 0)
                entrada['Consejo Final'] = stats.get('consejo_final', '')

            # Añadir a la lista de datos
            datos.append(entrada)

        # Convertir a DataFrame
        if datos:
            df = pd.DataFrame(datos)

            # Convertir columnas numéricas explícitamente para evitar errores
            columnas_numericas = [
                'Errores Gramática', 'Errores Léxico', 'Errores Puntuación',
                'Errores Estructura', 'Total Errores', 'Puntuación Coherencia',
                'Puntuación Cohesión', 'Puntuación Registro',
                'Puntuación Adecuación Cultural', 'Puntuación Global', 'Longitud Texto'
            ]

            for col in columnas_numericas:
                if col in df.columns:
                    # Convertir a float de manera segura
                    df[col] = pd.to_numeric(
                        df[col], errors='coerce').fillna(0).astype(float)

            # Guardar en caché para futuras consultas
            set_session_var("historial_correcciones", df)
            set_session_var("cached_historial_uid", uid)
            set_session_var("cached_historial_email", email)
            set_session_var("cached_historial_nombre", nombre)

            return df

        return None

    except Exception as e:
        logger.error(f"Error en obtener_historial_estudiante: {str(e)}")
        circuit_breaker.record_failure("firebase", error_type="get_history")
        return None


def guardar_feedback_firebase(rating, feedback_text):
    """
    Guarda el feedback del usuario en Firebase.

    Args:
        rating: Valoración numérica (1-5)
        feedback_text: Texto del feedback

    Returns:
        dict: Resultado de la operación
    """
    try:
        # Importar Firebase en forma segura
        try:
            import firebase_admin
            from firebase_admin import firestore
        except ImportError:
            return {"success": False, "message": "Firebase no está disponible"}

        db, conexion_ok = initialize_firebase()
        if not conexion_ok or db is None:
            return {"success": False, "message": "Conexión a Firestore no disponible"}

        # Generar ID único para el feedback
        feedback_id = str(uuid.uuid4())

        # Preparar los datos para guardar
        feedback_data = {
            "rating": rating,
            "feedback": feedback_text,
            "timestamp": firestore.SERVER_TIMESTAMP,
            "app_version": APP_VERSION,
            "session_id": get_session_var("session_id", ""),
            "usuario": get_session_var("usuario_actual", "Anónimo"),
            "email": get_session_var("email_usuario", "")
        }

        # Guardar en Firestore
        db.collection("feedback").document(feedback_id).set(feedback_data)

        return {"success": True, "message": "Feedback guardado correctamente"}

    except Exception as e:
        logger.error(f"Error al guardar feedback: {str(e)}")
        circuit_breaker.record_failure("firebase", error_type="save_feedback")
        return {"success": False, "message": f"Error al guardar feedback: {str(e)}"}


def guardar_metricas_modelo(modelo, tiempo_respuesta, longitud_texto, resultado_exitoso):
    """
    Guarda métricas de rendimiento del modelo en Firebase para análisis.
    Implementación sin transacciones para evitar problemas de compatibilidad.

    Args:
        modelo: Nombre del modelo utilizado
        tiempo_respuesta: Tiempo de respuesta en segundos
        longitud_texto: Longitud del texto (en tokens o palabras)
        resultado_exitoso: Si la solicitud fue exitosa

    Returns:
        dict: Resultado de la operación
    """
    try:
        # Registrar localmente en session_state para uso en la sesión actual
        if "metricas_modelos" not in st.session_state:
            st.session_state.metricas_modelos = {}

        if modelo not in st.session_state.metricas_modelos:
            st.session_state.metricas_modelos[modelo] = {
                "total_requests": 0,
                "successful_requests": 0,
                "total_time": 0,
                "total_tokens": 0
            }

        metrics = st.session_state.metricas_modelos[modelo]
        metrics["total_requests"] += 1
        if resultado_exitoso:
            metrics["successful_requests"] += 1
        metrics["total_time"] += tiempo_respuesta
        metrics["total_tokens"] += longitud_texto

        # Solo guardar en Firebase si está disponible
        try:
            import firebase_admin
            from firebase_admin import firestore
        except ImportError:
            return {"success": False, "message": "Firebase no está disponible"}

        db, conexion_ok = initialize_firebase()
        if not conexion_ok or db is None:
            return {"success": False, "message": "Conexión a Firestore no disponible"}

        # Generar ID único para la métrica
        metrica_id = str(uuid.uuid4())

        # Preparar los datos para guardar
        metrica_data = {
            "modelo": modelo,
            "tiempo_respuesta": tiempo_respuesta,
            "longitud_texto": longitud_texto,
            "resultado_exitoso": resultado_exitoso,
            "timestamp": firestore.SERVER_TIMESTAMP,
            "app_version": APP_VERSION,
            "session_id": get_session_var("session_id", ""),
            "usuario": get_session_var("usuario_actual", "Anónimo")
        }

        # Guardar en Firestore
        db.collection("metricas_modelo").document(metrica_id).set(metrica_data)

        # SOLUCIÓN C: Actualizar estadísticas sin usar transacciones
        # Obtenemos una referencia al documento de estadísticas del modelo
        stats_ref = db.collection("estadisticas_modelos").document(modelo)

        # Obtener el documento actual (si existe)
        doc = stats_ref.get()

        if doc.exists:
            # Actualizar estadísticas existentes
            stats = doc.to_dict()
            stats["total_requests"] = stats.get("total_requests", 0) + 1
            stats["successful_requests"] = stats.get(
                "successful_requests", 0) + (1 if resultado_exitoso else 0)
            stats["total_time"] = stats.get("total_time", 0) + tiempo_respuesta
            stats["total_tokens"] = stats.get(
                "total_tokens", 0) + longitud_texto

            # Calcular promedios
            stats["avg_response_time"] = stats["total_time"] / \
                stats["total_requests"]
            stats["success_rate"] = stats["successful_requests"] / \
                stats["total_requests"]

            # Actualizar última actualización
            stats["last_updated"] = firestore.SERVER_TIMESTAMP
        else:
            # Crear nuevas estadísticas
            stats = {
                "modelo": modelo,
                "total_requests": 1,
                "successful_requests": 1 if resultado_exitoso else 0,
                "total_time": tiempo_respuesta,
                "total_tokens": longitud_texto,
                "avg_response_time": tiempo_respuesta,
                "success_rate": 1.0 if resultado_exitoso else 0.0,
                "last_updated": firestore.SERVER_TIMESTAMP,
                "created_at": firestore.SERVER_TIMESTAMP
            }

        # Guardar las estadísticas actualizadas directamente sin transacción
        stats_ref.set(stats)

        return {"success": True, "message": "Métricas guardadas correctamente"}

    except Exception as e:
        logger.error(f"Error al guardar métricas del modelo: {str(e)}")
        return {"success": False, "message": f"Error al guardar métricas: {str(e)}"}

# Función auxiliar para calcular nivel de complejidad (placeholder para uso en guardar_correccion_firestore)


def calcular_nivel_complejidad(texto):
    """
    Calcula un nivel aproximado de complejidad para un texto.

    Args:
        texto: Texto a analizar

    Returns:
        float: Valor de complejidad en escala 0-10
    """
    # Si no hay texto, retornar 0
    if not texto:
        return 0

    # Separar en palabras
    palabras = re.findall(r'\b\w+\b', texto.lower())

    if not palabras:
        return 0

    # Métricas básicas
    num_palabras = len(palabras)
    palabras_unicas = len(set(palabras))
    longitud_media = sum(len(palabra) for palabra in palabras) / num_palabras

    # Número de oraciones
    oraciones = re.split(r'[.!?]+', texto)
    oraciones = [o for o in oraciones if o.strip()]
    num_oraciones = len(oraciones) or 1  # Evitar división por cero

    # Longitud media de oración
    palabras_por_oracion = num_palabras / num_oraciones

    # Calcular un índice de complejidad (0-10)
    # Basado en:
    # - Diversidad léxica (palabras únicas / total)
    # - Longitud media de palabra
    # - Palabras por oración

    diversidad_lexica = palabras_unicas / num_palabras

    # Normalizar factores (aproximación simple):
    # - diversidad_lexica: normalmente entre 0.4-0.8 para textos en español
    # - longitud_media: normalmente entre 4-8 caracteres en español
    # - palabras_por_oracion: 10-25 para nivel intermedio

    complejidad = (
        (diversidad_lexica - 0.4) * 10 / 0.4 * 0.4 +  # 40% del peso
        (longitud_media - 4) / 4 * 0.3 +              # 30% del peso
        (palabras_por_oracion - 10) / 15 * 0.3        # 30% del peso
    )

    # Recortar a valores entre 0 y 10
    complejidad = max(0, min(10, complejidad * 10))

    return round(complejidad, 2)

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 4.2: Integración de Asistentes de OpenAI (Implementación Directa)
# ==================================================================================
#
# Este artefacto contiene:
# 1. Implementación directa de Asistentes de OpenAI
# 2. Gestión de asistentes por tipo de tarea
# 3. Caché y reutilización de asistentes
#
# Esta implementación reemplaza directamente el uso de la API tradicional
# por la API de Asistentes de OpenAI.
# ==================================================================================


class OpenAIAssistants:
    """
    Clase mejorada para interactuar con la API de Asistentes de OpenAI.
    Incluye soporte para herramientas como Code Interpreter y Functions.
    """

    # Mapeo de tipos de tarea a IDs de asistentes
    ASSISTANT_IDS = {
        "correccion_texto": None,  # ID para corrección de textos
        "generacion_ejercicios": None,  # ID para generación de ejercicios
        "plan_estudio": None,  # ID para planes de estudio
        "simulacro_examen": None,  # ID para simulacros de examen
        "default": None  # ID por defecto para otras tareas
    }

    def __init__(self, api_key: str):
        """
        Inicializa el gestor de Asistentes de OpenAI.

        Args:
            api_key: API key de OpenAI
        """
        self.api_key = api_key
        self.client = None
        self.current_model = "gpt-4-turbo"  # Modelo predeterminado
        self.assistant_cache = {}  # Caché de asistentes creados

        # Cargar IDs de asistentes desde secrets si están disponibles
        self._load_assistant_ids_from_secrets()

        # Inicializar cliente
        self._initialize_client()

    def _load_assistant_ids_from_secrets(self):
        """
        Carga las IDs de asistentes desde los secrets de Streamlit.
        """
        try:
            # Intentar cargar IDs de los secrets
            if hasattr(st, 'secrets'):
                # Primero intentar cargar desde sección específica OPENAI_ASSISTANTS
                if 'OPENAI_ASSISTANTS' in st.secrets:
                    assistants_config = st.secrets['OPENAI_ASSISTANTS']

                    for task_type, config_key in [
                        ("correccion_texto", "CORRECTION_ASSISTANT_ID"),
                        ("generacion_ejercicios", "EXERCISES_ASSISTANT_ID"),
                        ("plan_estudio", "STUDY_PLAN_ASSISTANT_ID"),
                        ("simulacro_examen", "EXAM_ASSISTANT_ID"),
                        ("default", "DEFAULT_ASSISTANT_ID")
                    ]:
                        if config_key in assistants_config:
                            self.ASSISTANT_IDS[task_type] = assistants_config[config_key]
                            logger.info(
                                f"Cargada ID de asistente para {task_type}: {self.ASSISTANT_IDS[task_type]}")

                # Buscar también en el nivel principal de secrets (para compatibilidad)
                elif 'CORRECTION_ASSISTANT_ID' in st.secrets:
                    self.ASSISTANT_IDS["correccion_texto"] = st.secrets['CORRECTION_ASSISTANT_ID']
                    logger.info(
                        f"Cargada ID de asistente de corrección desde nivel principal: {self.ASSISTANT_IDS['correccion_texto']}")
        except Exception as e:
            logger.warning(
                f"Error al cargar IDs de asistentes desde secrets: {e}")

    def _initialize_client(self):
        """
        Inicializa el cliente de OpenAI.
        """
        try:
            # Importar OpenAI aquí
            try:
                from openai import OpenAI
            except ImportError:
                logger.error(
                    "La biblioteca 'openai' no está instalada. Ejecuta 'pip install openai>=1.0.0'")
                return False

            # Inicializar cliente
            self.client = OpenAI(api_key=self.api_key)

            # Verificar conectividad con una operación simple
            try:
                # Listar asistentes para verificar acceso
                self.client.beta.assistants.list(limit=1)

                # Verificar asistentes configurados
                if any(self.ASSISTANT_IDS.values()):
                    for task_type, assistant_id in self.ASSISTANT_IDS.items():
                        if assistant_id:
                            try:
                                # Verificar que el asistente existe
                                self.client.beta.assistants.retrieve(
                                    assistant_id)
                                logger.info(
                                    f"Asistente para {task_type} verificado: {assistant_id}")
                            except Exception as e:
                                logger.warning(
                                    f"No se pudo verificar el asistente para {task_type} ({assistant_id}): {e}")

                # Registrar éxito
                circuit_breaker.record_success("openai")
                logger.info(
                    f"Conexión exitosa a OpenAI Assistants con modelo {self.current_model}")

                return True
            except Exception as e:
                logger.error(
                    f"Error al verificar conexión a OpenAI Assistants: {str(e)}")
                circuit_breaker.record_failure(
                    "openai", error_type="initialization")
                return False

        except Exception as e:
            logger.error(f"Error al inicializar OpenAI Assistants: {str(e)}")
            circuit_breaker.record_failure(
                "openai", error_type="initialization")
            return False

    def get_assistant_id(self, task_type: str, system_message: str) -> str:
        """
        Determina qué ID de asistente usar basado en el tipo de tarea y mensaje.

        Args:
            task_type: Tipo de tarea ('correccion_texto', 'generacion_ejercicios', etc.)
            system_message: Mensaje del sistema que describe la tarea

        Returns:
            str: ID del asistente a usar
        """
        # 1. Verificar si tenemos un ID configurado para este tipo de tarea
        if task_type in self.ASSISTANT_IDS and self.ASSISTANT_IDS[task_type]:
            return self.ASSISTANT_IDS[task_type]

        # 2. Verificar si tenemos un ID por defecto
        if self.ASSISTANT_IDS["default"]:
            return self.ASSISTANT_IDS["default"]

        # 3. Buscar en caché basado en el hash del mensaje
        instruccion_hash = hashlib.md5(
            f"{system_message}_{self.current_model}".encode()).hexdigest()
        if instruccion_hash in self.assistant_cache:
            return self.assistant_cache[instruccion_hash]

        # 4. Crear un nuevo asistente - MODIFICADO para asegurar que se pueda usar JSON
        try:
            # Verificar si el mensaje del sistema contiene la palabra "json"
            has_json_keyword = "json" in system_message.lower()

            # Configurar formato de respuesta
            response_format = {
                "type": "json_object"} if has_json_keyword else None

            # Si no tiene la palabra "json", añadir una nota
            if not has_json_keyword:
                logger.warning(
                    "El mensaje del sistema no contiene la palabra 'json', usando formato de respuesta abierto")

            # Crear el asistente con la configuración adecuada
            assistant = self.client.beta.assistants.create(
                name=f"TextoCorrector ELE - {task_type}",
                instructions=system_message,
                model=self.current_model,
                response_format=response_format
            )
            assistant_id = assistant.id

            # Guardar en caché
            self.assistant_cache[instruccion_hash] = assistant_id

            logger.info(
                f"Creado nuevo asistente para {task_type} con ID: {assistant_id}")
            return assistant_id
        except Exception as e:
            logger.error(f"Error al crear asistente: {e}")
            raise

    def get_completion(self, system_message: str, user_message: str,
                       max_retries: int = 3, task_type: str = "default") -> tuple:
        """
        Obtiene una respuesta usando OpenAI Assistants con soporte para herramientas.

        Args:
            system_message: Mensaje del sistema (instrucciones)
            user_message: Mensaje del usuario (contenido)
            max_retries: Número máximo de reintentos
            task_type: Tipo de tarea para seleccionar el asistente adecuado

        Returns:
            tuple: (respuesta_raw, resultado_json)
        """
        # Verificar que OpenAI esté configurado
        if self.api_key is None or self.client is None:
            return None, {"error": "API de OpenAI no configurada o inicializada"}

        if not circuit_breaker.can_execute("openai"):
            return None, {"error": "Servicio OpenAI temporalmente no disponible"}

        # MODIFICADO: Verificar si system_message o user_message contienen la palabra "json"
        has_json_keyword = "json" in system_message.lower() or "json" in user_message.lower()
        if not has_json_keyword:
            logger.warning(
                "Ni el mensaje del sistema ni el mensaje del usuario contienen la palabra 'json'")
            # Añadir una referencia a JSON en el mensaje del usuario si es necesario
            if task_type in ["correccion_texto", "generacion_ejercicios", "plan_estudio"]:
                user_message += "\n\nPor favor, proporciona tu respuesta en formato JSON."
                logger.info(
                    "Añadida referencia a JSON en el mensaje del usuario")

        # Registrar tiempo de inicio para métricas
        tiempo_inicio = time.time()
        longitud_estimada = len(user_message.split())

        try:
            # Determinar qué asistente usar
            try:
                assistant_id = self.get_assistant_id(task_type, system_message)
            except Exception as e:
                return None, {"error": f"Error al obtener ID de asistente: {str(e)}"}

            # Crear un nuevo thread para esta conversación
            thread = self.client.beta.threads.create()
            thread_id = thread.id

            # Añadir el mensaje del usuario al thread
            self.client.beta.threads.messages.create(
                thread_id=thread_id,
                role="user",
                content=user_message
            )

            # Ejecutar el asistente con reintentos
            for attempt in range(max_retries):
                try:
                    # Ejecutar el asistente con manejo de ejecuciones activas
                    try:
                        # Intentar crear una nueva ejecución
                        run = self.client.beta.threads.runs.create(
                            thread_id=thread_id,
                            assistant_id=assistant_id
                        )
                    except Exception as run_error:
                        # Verificar si el error es porque ya hay una ejecución activa
                        if "already has an active run" in str(run_error):
                            logger.warning(
                                f"El hilo {thread_id} ya tiene una ejecución activa. Intentando recuperarla...")

                            # Obtener la lista de ejecuciones para este hilo
                            runs_list = self.client.beta.threads.runs.list(
                                thread_id=thread_id)

                            # Buscar la ejecución activa
                            active_run = None
                            for existing_run in runs_list.data:
                                if existing_run.status in ["queued", "in_progress", "requires_action"]:
                                    active_run = existing_run
                                    break

                            if active_run:
                                # Usar la ejecución activa
                                logger.info(
                                    f"Recuperada ejecución activa: {active_run.id} con estado: {active_run.status}")
                                run = active_run
                            else:
                                # Si no podemos encontrar la ejecución activa, intentar cancelar todas las ejecuciones
                                logger.warning(
                                    "No se pudo encontrar la ejecución activa. Intentando cancelar ejecuciones existentes...")

                                for existing_run in runs_list.data:
                                    try:
                                        self.client.beta.threads.runs.cancel(
                                            thread_id=thread_id,
                                            run_id=existing_run.id
                                        )
                                        logger.info(
                                            f"Cancelada ejecución: {existing_run.id}")
                                    except Exception as cancel_error:
                                        logger.warning(
                                            f"No se pudo cancelar ejecución {existing_run.id}: {cancel_error}")

                                # Esperar un momento para que las cancelaciones se procesen
                                time.sleep(2)

                                # Intentar crear una nueva ejecución otra vez
                                run = self.client.beta.threads.runs.create(
                                    thread_id=thread_id,
                                    assistant_id=assistant_id
                                )
                        else:
                            # Si es otro tipo de error, propagarlo
                            raise

                    # Esperar a que termine la ejecución con timeout
                    max_wait_time = 120  # 2 minutos máximo
                    start_time = time.time()

                    polling_interval = 1  # Comenzamos con 1 segundo
                    max_polling_interval = 5  # Máximo intervalo de polling: 5 segundos
                    polling_count = 0

                    while True:
                        # Verificar si hemos excedido el tiempo máximo
                        if time.time() - start_time > max_wait_time:
                            raise TimeoutError(
                                f"Timeout esperando respuesta del asistente después de {max_wait_time}s")

                        # Obtener estado de la ejecución
                        run_status = self.client.beta.threads.runs.retrieve(
                            thread_id=thread_id,
                            run_id=run.id
                        )

                        polling_count += 1
                        status = run_status.status
                        logger.info(
                            f"Estado de ejecución ({polling_count}): {status}")

                        # Si ha terminado, salir del bucle
                        if status == "completed":
                            logger.info(
                                f"Ejecución completada después de {polling_count} consultas")
                            break

                        # Si ha fallado, levantar excepción
                        if status in ["failed", "cancelled", "expired"]:
                            error_detail = ""
                            if hasattr(run_status, 'last_error'):
                                error_detail = f" - {run_status.last_error}"
                            raise Exception(
                                f"Ejecución fallida con estado: {status}{error_detail}")

                        # NUEVO: Manejar estado "requires_action"
                        if status == "requires_action":
                            logger.info(
                                f"El asistente requiere una acción. Procesando...")
                            # Verificar el tipo de acción requerida
                            if hasattr(run_status, 'required_action'):
                                required_action = run_status.required_action

                                # Si requiere una acción de función
                                if required_action.type == "submit_tool_outputs":
                                    # Extraer las llamadas a funciones
                                    tool_calls = required_action.submit_tool_outputs.tool_calls
                                    tool_outputs = []

                                    # Procesar cada llamada a función
                                    for tool_call in tool_calls:
                                        # Extraer información de la llamada
                                        tool_call_id = tool_call.id
                                        function_name = tool_call.function.name
                                        arguments = tool_call.function.arguments

                                        logger.info(
                                            f"Función llamada: {function_name} con argumentos: {arguments}")

                                        # Aquí manejaríamos las funciones específicas que el asistente puede llamar
                                        # Por ahora, devolvemos una respuesta simulada para cada función
                                        if function_name == "corregir_texto":
                                            output = json.dumps({
                                                "success": True,
                                                "message": "Corrección realizada",
                                                "via_function": True
                                            })
                                        elif function_name == "generar_informe":
                                            output = json.dumps({
                                                "success": True,
                                                "message": "Informe generado",
                                                "via_function": True
                                            })
                                        else:
                                            output = json.dumps({
                                                "error": f"Función no implementada: {function_name}"
                                            })

                                        # Añadir la salida a la lista
                                        tool_outputs.append({
                                            "tool_call_id": tool_call_id,
                                            "output": output
                                        })

                                    # Enviar las salidas de las herramientas
                                    if tool_outputs:
                                        self.client.beta.threads.runs.submit_tool_outputs(
                                            thread_id=thread_id,
                                            run_id=run.id,
                                            tool_outputs=tool_outputs
                                        )
                                        logger.info(
                                            f"Enviadas {len(tool_outputs)} salidas de herramientas al asistente")
                                else:
                                    logger.warning(
                                        f"Tipo de acción no soportada: {required_action.type}")
                                    raise Exception(
                                        f"Tipo de acción no soportada: {required_action.type}")
                            else:
                                logger.warning(
                                    "Estado requires_action pero sin required_action especificada")
                                raise Exception(
                                    "Estado requires_action pero sin required_action especificada")

                        # Verificar si está en espera (queued) o en proceso (in_progress)
                        # y ajustar el tiempo de espera en consecuencia
                        if polling_count > 30:  # Si llevamos muchas iteraciones, posible problema
                            logger.warning(
                                f"Demasiadas consultas de estado ({polling_count}). Posible problema con el asistente.")
                            # Opción 1: Continuar con intervalos más largos
                            polling_interval = max_polling_interval

                        # Ajustar intervalo de polling de forma exponencial hasta el máximo
                        polling_interval = min(
                            polling_interval * 1.5, max_polling_interval)
                        logger.info(
                            f"Esperando {polling_interval:.1f}s antes de verificar estado nuevamente")

                        # Esperar antes de verificar de nuevo
                        time.sleep(polling_interval)

                    # Obtener los mensajes del thread
                    messages = self.client.beta.threads.messages.list(
                        thread_id=thread_id
                    )

                    # El primer mensaje será la respuesta más reciente del asistente
                    for message in messages.data:
                        if message.role == "assistant":
                            # Extraer el contenido
                            content_parts = [
                                part.text.value for part in message.content
                                if hasattr(part, 'text')
                            ]
                            content = ''.join(content_parts)

                            # Intentar extraer JSON
                            data_json = extract_json_safely(content)

                            # Calcular tiempo total de respuesta
                            tiempo_total = time.time() - tiempo_inicio

                            # Registrar métricas
                            guardar_metricas_modelo(
                                modelo=self.current_model,
                                tiempo_respuesta=tiempo_total,
                                longitud_texto=longitud_estimada,
                                resultado_exitoso="error" not in data_json
                            )

                            # Registrar éxito
                            circuit_breaker.record_success("openai")

                            logger.info(
                                f"Solicitud a OpenAI Assistants completada en {tiempo_total:.2f}s")
                            return content, data_json

                    # Si llegamos aquí, no hay respuesta del asistente
                    raise Exception("No se encontró respuesta del asistente")

                except TimeoutError as e:
                    logger.warning(
                        f"Timeout en solicitud a OpenAI Assistants (intento {attempt+1}/{max_retries}): {e}")

                    if attempt == max_retries - 1:
                        # Registrar fallo si es el último intento
                        circuit_breaker.record_failure(
                            "openai", error_type="timeout")
                        tiempo_total = time.time() - tiempo_inicio
                        guardar_metricas_modelo(
                            modelo=self.current_model,
                            tiempo_respuesta=tiempo_total,
                            longitud_texto=longitud_estimada,
                            resultado_exitoso=False
                        )
                        return None, {"error": f"Timeout en solicitud a OpenAI Assistants después de {max_retries} intentos"}

                    # Esperar antes de reintentar
                    wait_time = min(60, 4 ** attempt)
                    logger.info(f"Esperando {wait_time}s antes de reintentar.")
                    time.sleep(wait_time)

                except Exception as e:
                    logger.warning(
                        f"Error en solicitud a OpenAI Assistants (intento {attempt+1}/{max_retries}): {e}")

                    if attempt == max_retries - 1:
                        # Registrar fallo si es el último intento
                        circuit_breaker.record_failure(
                            "openai", error_type="general")
                        tiempo_total = time.time() - tiempo_inicio
                        guardar_metricas_modelo(
                            modelo=self.current_model,
                            tiempo_respuesta=tiempo_total,
                            longitud_texto=longitud_estimada,
                            resultado_exitoso=False
                        )
                        return None, {"error": f"Error en solicitud a OpenAI Assistants: {str(e)}"}

                    # Esperar antes de reintentar
                    wait_time = min(60, 2 ** attempt)
                    logger.info(f"Esperando {wait_time}s antes de reintentar.")
                    time.sleep(wait_time)

        except Exception as e:
            # Registrar fallo general
            logger.error(f"Error general en OpenAI Assistants: {str(e)}")
            circuit_breaker.record_failure("openai", error_type="general")
            tiempo_total = time.time() - tiempo_inicio
            guardar_metricas_modelo(
                modelo=self.current_model if hasattr(
                    self, 'current_model') else "desconocido",
                tiempo_respuesta=tiempo_total,
                longitud_texto=longitud_estimada,
                resultado_exitoso=False
            )
            return None, {"error": f"Error en API de OpenAI Assistants: {str(e)}"}


# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 5: Funciones de Procesamiento de Texto
# ==================================================================================
#
# Este artefacto contiene:
# 1. Funciones para procesamiento de JSON
# 2. Funciones para corrección de texto con OpenAI (optimizadas)
# 3. Funciones para generar audio con ElevenLabs
# 4. Funciones para OCR y generación de imágenes
#
# Estas funciones constituyen el núcleo funcional de la aplicación,
# proporcionando las capacidades principales de procesamiento. Se han
# mejorado para manejar errores de timeout y optimizar modelos.
# ==================================================================================

# --- 1. PROCESAMIENTO DE JSON DE OPENAI ---

def extract_json_safely(content):
    """
    Extrae contenido JSON de una respuesta con múltiples estrategias.
    Implementa parsing robusto para evitar errores.

    Args:
        content: Contenido de texto que debería contener JSON

    Returns:
        dict: El contenido parseado como JSON o un diccionario con error
    """
    # Si es None o vacío, retornar error inmediatamente
    if not content:
        return {"error": "Contenido vacío, no se puede extraer JSON"}

    # Intento directo
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        # Limpiar el contenido - eliminar caracteres que puedan causar problemas
        content_clean = re.sub(r'[^\x20-\x7E]', '', content)

        try:
            return json.loads(content_clean)
        except json.JSONDecodeError:
            pass

        # Búsqueda con regex para JSON completo
        # Regex mejorada para JSON anidado
        json_pattern = r'(\{(?:[^{}]|(?1))*\})'
        match = re.search(json_pattern, content_clean, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

        # Segunda estrategia: buscar cualquier JSON entre llaves
        simple_pattern = r'\{.*\}'
        match = re.search(simple_pattern, content_clean, re.DOTALL)
        if match:
            try:
                # Intentar limpiar el JSON encontrado
                potential_json = match.group(0)
                # Eliminar comillas mal formadas, escape chars, etc.
                potential_json = re.sub(r'[\r\n\t]', ' ', potential_json)
                return json.loads(potential_json)
            except json.JSONDecodeError:
                pass

    # Si no se pudo extraer, devolver un objeto error
    logger.warning(f"No se pudo extraer JSON de: {content[:100]}...")
    return {"error": "No se pudo extraer JSON válido", "raw_content": content[:500]}


# --- 2. FUNCIONES DE CORRECCIÓN DE TEXTO CON OPENAI (OPTIMIZADAS) ---

def optimizar_prompt_para_gpt35(system_prompt, modelo):
    """
    Optimiza el prompt del sistema según el modelo utilizado.
    Para GPT-3.5 añade instrucciones más detalladas y estructura.

    Args:
        system_prompt: Prompt original
        modelo: Nombre del modelo a utilizar

    Returns:
        str: Prompt optimizado
    """
    if "gpt-3.5" in modelo:
        # Añadir instrucciones más explícitas para GPT-3.5
        optimized_prompt = system_prompt + """

INSTRUCCIONES ADICIONALES PARA CORRECCIÓN PRECISA:
1. Analiza primero todas las estructuras gramaticales.
2. Identifica cada error y clasifícalo por categoría.
3. Para cada error, proporciona SIEMPRE el fragmento original, la corrección y una explicación concisa.
4. Evalúa la coherencia analizando la progresión lógica de ideas.
5. Evalúa la cohesión considerando conectores, referencias y puntuación.
6. Evalúa el registro según el tipo de texto y contexto cultural.
7. Asegúrate de que tus correcciones respetan el nivel del estudiante.
8. Sé extremadamente preciso en tu análisis contextual.

ESTRUCTURA EXACTA REQUERIDA:
{
  "saludo": "Saludo personalizado al estudiante",
  "tipo_texto": "Tipo de texto analizado",
  "errores": {
    "Gramática": [ {"fragmento_erroneo": "...", "correccion": "...", "explicacion": "..."} ],
    "Léxico": [ {"fragmento_erroneo": "...", "correccion": "...", "explicacion": "..."} ],
    "Puntuación": [ {"fragmento_erroneo": "...", "correccion": "...", "explicacion": "..."} ],
    "Estructura textual": [ {"fragmento_erroneo": "...", "correccion": "...", "explicacion": "..."} ]
  },
  "texto_corregido": "Texto completo con correcciones",
  "analisis_contextual": {
    "coherencia": {"puntuacion": 7, "comentario": "...", "sugerencias": ["...", "..."]},
    "cohesion": {"puntuacion": 8, "comentario": "...", "sugerencias": ["...", "..."]},
    "registro_linguistico": {"puntuacion": 6, "tipo_detectado": "...", "adecuacion": "...", "sugerencias": ["...", "..."]},
    "adecuacion_cultural": {"puntuacion": 7, "comentario": "...", "elementos_destacables": ["...", "..."], "sugerencias": ["...", "..."]}
  },
  "consejo_final": "Consejo personalizado para el estudiante",
  "fin": "Fin de texto corregido."
}

RECUERDA: Debes proporcionar un JSON válido con EXACTAMENTE la estructura solicitada.
"""
        return optimized_prompt

    return system_prompt


# Reemplazar completamente la función existente
# Reemplaza solo la función obtener_json_de_openai con esta versión
# Esta versión incluye todas las variables y funciones auxiliares necesarias dentro de sí misma

def obtener_json_de_openai(system_msg, user_msg, max_retries=3):
    """
    Obtiene una respuesta de OpenAI usando directamente el asistente.
    Versión mejorada con soporte para herramientas como Code Interpreter y Functions.

    Args:
        system_msg: Mensaje del sistema (usado en instrucciones)
        user_msg: Mensaje del usuario con el texto a procesar
        max_retries: Número máximo de reintentos

    Returns:
        tuple: (respuesta raw, contenido JSON parseado)
    """
    # Verificar que OpenAI esté configurado
    if api_keys["openai"] is None:
        return None, {"error": "API de OpenAI no configurada"}

    if not circuit_breaker.can_execute("openai"):
        return None, {"error": "Servicio OpenAI temporalmente no disponible"}

    # Inicializar el cliente de OpenAI
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_keys["openai"])
    except ImportError:
        logger.error(
            "La biblioteca 'openai' no está instalada. Instálala con: pip install openai>=1.0.0")
        return None, {"error": "Biblioteca OpenAI no instalada"}
    except Exception as e:
        logger.error(f"Error al inicializar cliente OpenAI: {e}")
        return None, {"error": f"Error al inicializar OpenAI: {str(e)}"}

    # Obtener ID del asistente
    assistant_id = None
    try:
        if hasattr(st, 'secrets'):
            # Buscar en nivel principal o en sección específica
            if 'OPENAI_ASSISTANTS' in st.secrets and 'CORRECTION_ASSISTANT_ID' in st.secrets['OPENAI_ASSISTANTS']:
                assistant_id = st.secrets['OPENAI_ASSISTANTS']['CORRECTION_ASSISTANT_ID']
            else:
                # Buscar directamente en el nivel principal
                assistant_id = st.secrets.get("CORRECTION_ASSISTANT_ID")

            if assistant_id:
                logger.info(f"Usando asistente ID: {assistant_id}")
            else:
                logger.warning("ID de asistente no configurada en secrets")
                return None, {"error": "ID de asistente no configurado. Configura CORRECTION_ASSISTANT_ID en secrets."}
    except Exception as e:
        logger.error(f"Error al cargar ID de asistente: {e}")
        return None, {"error": f"Error al cargar ID del asistente: {str(e)}"}

    # Si no hay ID de asistente, no podemos continuar
    if not assistant_id:
        logger.error("No se pudo obtener ID del asistente de corrección")
        return None, {"error": "No se pudo obtener ID del asistente de corrección"}

    # Log explícito que podamos buscar
    logger.info(f"🤖 USANDO ASISTENTE DE OPENAI ID: {assistant_id}")

    # Registrar tiempo para métricas
    tiempo_inicio = time.time()
    longitud_estimada = len(user_msg.split())

    try:
        # Crear un thread para esta conversación
        thread = client.beta.threads.create()
        thread_id = thread.id
        logger.info(f"Thread creado: {thread_id}")

        # Añadir el mensaje del usuario
        client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=user_msg
        )
        logger.info("Mensaje usuario añadido al thread")

        # Ejecutar con reintentos
        for attempt in range(max_retries):
            try:
                # Iniciar ejecución del asistente
                run = client.beta.threads.runs.create(
                    thread_id=thread_id,
                    assistant_id=assistant_id
                )
                run_id = run.id
                logger.info(f"Ejecución iniciada: {run_id}")

                # Esperar a que termine
                # 3 minutos (aumentado para dar más tiempo)
                max_wait_time = 180
                start_time = time.time()

                # Control de polling para evitar demasiadas solicitudes
                polling_interval = 1  # Comienza con 1 segundo
                max_polling_interval = 5  # Máximo 5 segundos
                polling_count = 0

                while True:
                    # Verificar timeout
                    if time.time() - start_time > max_wait_time:
                        raise TimeoutError(
                            f"Timeout esperando respuesta del asistente")

                    # Obtener estado
                    run_status = client.beta.threads.runs.retrieve(
                        thread_id=thread_id,
                        run_id=run_id
                    )
                    status = run_status.status
                    polling_count += 1
                    logger.info(
                        f"Estado de ejecución ({polling_count}): {status}")

                    # Si completado, salir del bucle
                    if status == "completed":
                        logger.info(
                            f"Ejecución completada después de {polling_count} consultas")
                        break

                    # Si fallido, levantar excepción
                    if status in ["failed", "cancelled", "expired"]:
                        error_details = ""
                        if hasattr(run_status, 'last_error'):
                            error_details = f" - {run_status.last_error}"
                        raise Exception(
                            f"Ejecución fallida: {status}{error_details}")

                    # NUEVO: Manejar estado "requires_action"
                    if status == "requires_action":
                        logger.info(
                            "El asistente requiere una acción. Procesando...")

                        # Verificar que tenemos la información necesaria
                        if hasattr(run_status, 'required_action'):
                            required_action = run_status.required_action

                            # Verificar que es una acción de herramienta
                            if required_action.type == "submit_tool_outputs":
                                tool_calls = required_action.submit_tool_outputs.tool_calls
                                logger.info(
                                    f"El asistente requiere {len(tool_calls)} llamadas a herramientas")

                                # Procesar cada llamada a herramienta
                                tool_outputs = []
                                for tool_call in tool_calls:
                                    # Obtener información de la llamada
                                    tool_call_id = tool_call.id

                                    # Para el Code Interpreter
                                    if tool_call.type == "code_interpreter":
                                        # Extraer el código a ejecutar
                                        logger.info(
                                            "Llamada a Code Interpreter detectada")
                                        output = json.dumps({
                                            "result": "Código procesado correctamente (simulado)",
                                            "via_code_interpreter": True
                                        })
                                    # Para Functions (corregir_texto, generar_informe)
                                    elif tool_call.type == "function":
                                        function_name = tool_call.function.name
                                        arguments = tool_call.function.arguments
                                        logger.info(
                                            f"Llamada a función {function_name} con argumentos: {arguments}")

                                        # Respuestas simuladas para cada función
                                        if function_name == "corregir_texto":
                                            output = json.dumps({
                                                "success": True,
                                                "message": "Texto corregido correctamente",
                                                "via_function": True
                                            })
                                        elif function_name == "generar_informe":
                                            output = json.dumps({
                                                "success": True,
                                                "message": "Informe generado correctamente",
                                                "via_function": True
                                            })
                                        else:
                                            output = json.dumps({
                                                "error": f"Función no implementada: {function_name}"
                                            })
                                    else:
                                        output = json.dumps({
                                            "error": f"Tipo de herramienta no soportada: {tool_call.type}"
                                        })

                                    # Añadir la salida a la lista
                                    tool_outputs.append({
                                        "tool_call_id": tool_call_id,
                                        "output": output
                                    })

                                # Enviar las salidas al asistente
                                if tool_outputs:
                                    logger.info(
                                        f"Enviando {len(tool_outputs)} salidas de herramientas")
                                    client.beta.threads.runs.submit_tool_outputs(
                                        thread_id=thread_id,
                                        run_id=run_id,
                                        tool_outputs=tool_outputs
                                    )
                                    logger.info(
                                        "Salidas de herramientas enviadas correctamente")
                            else:
                                logger.warning(
                                    f"Tipo de acción no soportada: {required_action.type}")
                                raise Exception(
                                    f"Tipo de acción no soportada: {required_action.type}")
                        else:
                            logger.warning(
                                "Estado requires_action pero sin required_action")
                            raise Exception(
                                "Estado requires_action pero sin required_action")

                    # Ajustar intervalo de polling para no sobrecargar la API
                    if polling_count > 30:
                        polling_interval = max_polling_interval
                    else:
                        polling_interval = min(
                            polling_interval * 1.5, max_polling_interval)

                    logger.info(
                        f"Esperando {polling_interval:.1f}s antes de verificar estado nuevamente")
                    time.sleep(polling_interval)

                # Obtener mensajes
                messages = client.beta.threads.messages.list(
                    thread_id=thread_id
                )
                logger.info(f"Mensajes recibidos: {len(messages.data)}")

                # Buscar respuesta del asistente
                for message in messages.data:
                    if message.role == "assistant":
                        # Extraer contenido
                        content_parts = [
                            part.text.value for part in message.content
                            if hasattr(part, 'text')
                        ]
                        content = ''.join(content_parts)
                        logger.info(
                            "Respuesta del asistente recibida correctamente")

                        # Extraer JSON
                        data_json = extract_json_safely(content)

                        # Métricas
                        tiempo_total = time.time() - tiempo_inicio
                        guardar_metricas_modelo(
                            modelo="gpt-4-turbo",
                            tiempo_respuesta=tiempo_total,
                            longitud_texto=longitud_estimada,
                            resultado_exitoso="error" not in data_json
                        )

                        # Asegurar que tiene marcador via_assistant
                        if isinstance(data_json, dict) and "via_assistant" not in data_json:
                            data_json["via_assistant"] = True
                            logger.info(
                                "Añadido marcador via_assistant al resultado")

                        # Registrar éxito
                        circuit_breaker.record_success("openai")
                        logger.info(
                            f"Proceso completado en {tiempo_total:.2f}s")

                        return content, data_json

                # Si no hay respuesta
                raise Exception("No se encontró respuesta del asistente")

            except TimeoutError as e:
                logger.warning(
                    f"Timeout en intento {attempt+1}/{max_retries}: {e}")

                # Si último intento, reportar error
                if attempt == max_retries - 1:
                    circuit_breaker.record_failure(
                        "openai", error_type="timeout")
                    tiempo_total = time.time() - tiempo_inicio
                    guardar_metricas_modelo(
                        modelo="gpt-4-turbo",
                        tiempo_respuesta=tiempo_total,
                        longitud_texto=longitud_estimada,
                        resultado_exitoso=False
                    )
                    return None, {"error": f"Timeout esperando respuesta del asistente después de {max_retries} intentos"}

                # Esperar antes de reintentar
                wait_time = min(60, 4 ** attempt)
                logger.info(f"Esperando {wait_time}s antes de reintentar")
                time.sleep(wait_time)

            except Exception as e:
                logger.warning(
                    f"Error en intento {attempt+1}/{max_retries}: {e}")

                # Si último intento, reportar error
                if attempt == max_retries - 1:
                    circuit_breaker.record_failure(
                        "openai", error_type="general")
                    tiempo_total = time.time() - tiempo_inicio
                    guardar_metricas_modelo(
                        modelo="gpt-4-turbo",
                        tiempo_respuesta=tiempo_total,
                        longitud_texto=longitud_estimada,
                        resultado_exitoso=False
                    )
                    return None, {"error": f"Error con el asistente: {str(e)}"}

                # Esperar antes de reintentar
                wait_time = min(60, 2 ** attempt)
                logger.info(f"Esperando {wait_time}s antes de reintentar")
                time.sleep(wait_time)

    except Exception as e:
        # Registrar fallo general
        logger.error(f"Error general con asistente: {str(e)}")
        logger.error(traceback.format_exc())
        circuit_breaker.record_failure("openai", error_type="general")
        tiempo_total = time.time() - tiempo_inicio
        guardar_metricas_modelo(
            modelo="gpt-4-turbo",
            tiempo_respuesta=tiempo_total,
            longitud_texto=longitud_estimada,
            resultado_exitoso=False
        )
        return None, {"error": f"Error en Asistente OpenAI: {str(e)}"}


def corregir_texto(texto, nombre, nivel, idioma, tipo_texto, contexto_cultural, info_adicional=""):
    """
    Realiza una corrección completa de un texto con análisis contextual
    utilizando OpenAI. Mejora: asegura que el idioma de feedback se respete.

    Args:
        texto: Texto a corregir
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        idioma: Idioma de corrección (Español, Francés, Inglés)
        tipo_texto: Tipo de texto
        contexto_cultural: Contexto cultural relevante
        info_adicional: Información adicional o contexto

    Returns:
        dict: Resultado de la corrección o mensaje de error
    """
    # Si el modo degradado está activado, usar una versión simplificada
    if get_session_var("modo_degradado", False):
        return _corregir_texto_degradado(texto, nivel)

    # Validar la entrada para evitar problemas
    texto_validado = validar_texto_entrada(texto)
    if not texto_validado:
        return {"error": "El texto proporcionado no es válido o está vacío."}

    # Crear un hash para identificar esta solicitud única
    try:
        input_str = f"{texto_validado}|{nivel}|{idioma}|{tipo_texto}|{contexto_cultural}|{info_adicional}"
        request_hash = hashlib.md5(input_str.encode()).hexdigest()
        cache_key = f"correction_{request_hash}"

        # Verificar si ya tenemos esta corrección en caché de sesión
        if cache_key in st.session_state:
            # Antes de devolver el resultado, verificar que el idioma está correcto
            result = st.session_state[cache_key]
            if resultado_respeta_idioma(result, idioma):
                return result
            else:
                # Si el idioma no es correcto, forzamos una nueva corrección
                logger.warning(
                    f"Resultado en caché no respeta el idioma {idioma}. Forzando nueva corrección.")

        # Si no está en caché o el idioma no es correcto, llamar a la implementación
        result = _corregir_texto_impl(
            texto_validado, nombre, nivel, idioma, tipo_texto, contexto_cultural, info_adicional)

        # Verificar que el idioma se respeta antes de guardar en caché
        if resultado_respeta_idioma(result, idioma):
            # Guardar en caché de sesión para futuras referencias
            st.session_state[cache_key] = result
            return result
        else:
            # Si el idioma no se respeta, intentar corregirlo
            result_corregido = corregir_idioma_resultado(result, idioma)
            # Guardar la versión corregida en caché
            st.session_state[cache_key] = result_corregido
            return result_corregido
    except Exception as e:
        # Si hay error con el caché, llamar directamente a la implementación
        logger.error(f"Error en caché de corrección: {str(e)}")
        return _corregir_texto_impl(texto_validado, nombre, nivel, idioma, tipo_texto, contexto_cultural, info_adicional)


def resultado_respeta_idioma(resultado, idioma_solicitado):
    """
    Verifica si el resultado de la corrección respeta el idioma solicitado.

    Args:
        resultado: Diccionario con el resultado de la corrección
        idioma_solicitado: Idioma solicitado (Español, Francés, Inglés)

    Returns:
        bool: True si respeta el idioma, False en caso contrario
    """
    if not isinstance(resultado, dict):
        return False

    # Si hay error, no verificamos más
    if "error" in resultado:
        return True

    # Verificar que las categorías principales existen
    if not all(k in resultado for k in ["saludo", "errores", "analisis_contextual"]):
        return False

    # Si el idioma solicitado es español, no necesitamos verificar
    if idioma_solicitado.lower() == "español":
        return True

    # Verificar saludo
    saludo = resultado.get("saludo", "")
    if not es_texto_en_idioma(saludo, idioma_solicitado):
        return False

    # Verificar explicaciones de errores
    errores = resultado.get("errores", {})
    if isinstance(errores, dict):
        for categoria, lista_errores in errores.items():
            if isinstance(lista_errores, list):
                for error in lista_errores:
                    if isinstance(error, dict) and "explicacion" in error:
                        if not es_texto_en_idioma(error["explicacion"], idioma_solicitado):
                            return False

    # Verificar análisis contextual
    analisis = resultado.get("analisis_contextual", {})
    if isinstance(analisis, dict):
        for categoria, datos in analisis.items():
            if isinstance(datos, dict):
                if "comentario" in datos and not es_texto_en_idioma(datos["comentario"], idioma_solicitado):
                    return False
                if "sugerencias" in datos and isinstance(datos["sugerencias"], list):
                    for sugerencia in datos["sugerencias"]:
                        if not es_texto_en_idioma(sugerencia, idioma_solicitado):
                            return False
                if "adecuacion" in datos and not es_texto_en_idioma(datos["adecuacion"], idioma_solicitado):
                    return False
                if "elementos_destacables" in datos and isinstance(datos["elementos_destacables"], list):
                    for elemento in datos["elementos_destacables"]:
                        if not es_texto_en_idioma(elemento, idioma_solicitado):
                            return False

    return True


def es_texto_en_idioma(texto, idioma):
    """
    Detecta si un texto está en el idioma especificado.
    Método simple basado en palabras frecuentes de cada idioma.

    Args:
        texto: Texto a analizar
        idioma: Idioma esperado (Español, Francés, Inglés)

    Returns:
        bool: True si el texto parece estar en el idioma esperado
    """
    if not texto or len(texto) < 5:
        return True  # Textos muy cortos son difíciles de detectar

    # Convertir a minúsculas
    texto_lower = texto.lower()

    # Palabras frecuentes por idioma
    palabras_frecuentes = {
        "español": ["el", "la", "los", "las", "en", "y", "que", "por", "para", "con", "un", "una", "es", "son", "su", "del", "al", "este", "esta", "pero", "como", "muy", "bien", "más", "hay"],
        "inglés": ["the", "and", "is", "in", "to", "of", "a", "for", "that", "with", "this", "it", "you", "not", "be", "are", "on", "have", "as", "your", "from", "they", "we", "there", "more", "can", "an", "must"],
        "francés": ["le", "la", "les", "des", "un", "une", "et", "est", "en", "dans", "pour", "que", "qui", "du", "à", "au", "aux", "ce", "cette", "ces", "mais", "ou", "où", "avec", "sur", "nous", "vous", "ils", "elles"]
    }

    # Determinar el idioma esperado
    idioma_key = ""
    if "español" in idioma.lower():
        idioma_key = "español"
    elif "inglés" in idioma.lower() or "ingles" in idioma.lower():
        idioma_key = "inglés"
    elif "francés" in idioma.lower() or "frances" in idioma.lower():
        idioma_key = "francés"
    else:
        return True  # Si no reconocemos el idioma, asumimos que está bien

    # Contar palabras frecuentes del idioma esperado
    contador_idioma_esperado = 0
    for palabra in palabras_frecuentes[idioma_key]:
        if f" {palabra} " in f" {texto_lower} " or texto_lower.startswith(f"{palabra} ") or texto_lower.endswith(f" {palabra}"):
            contador_idioma_esperado += 1

    # Contar palabras frecuentes de otros idiomas
    contadores_otros_idiomas = {}
    for otro_idioma, palabras in palabras_frecuentes.items():
        if otro_idioma != idioma_key:
            contador = 0
            for palabra in palabras:
                if f" {palabra} " in f" {texto_lower} " or texto_lower.startswith(f"{palabra} ") or texto_lower.endswith(f" {palabra}"):
                    contador += 1
            contadores_otros_idiomas[otro_idioma] = contador

    # Determinar si el texto está en el idioma esperado
    if contador_idioma_esperado == 0:
        return True  # No hay palabras frecuentes de ningún idioma, difícil determinar

    for otro_idioma, contador in contadores_otros_idiomas.items():
        if contador > contador_idioma_esperado * 1.5:
            # Si hay muchas más palabras de otro idioma, probablemente no está en el esperado
            return False

    return True


def corregir_idioma_resultado(resultado, idioma_solicitado):
    """
    Intenta corregir el idioma de las explicaciones si no coincide con el solicitado.
    Si no es posible, devuelve el resultado original.

    Args:
        resultado: Diccionario con el resultado de la corrección
        idioma_solicitado: Idioma solicitado (Español, Francés, Inglés)

    Returns:
        dict: Resultado con idioma corregido si es posible
    """
    if not isinstance(resultado, dict) or "error" in resultado:
        return resultado

    # Si el idioma es español, no es necesario corregir
    if idioma_solicitado.lower() == "español":
        return resultado

    try:
        # Obtener OpenAI para traducción
        priorizar_costo = get_session_var("priorizar_costo", True)
        modelo_seleccionado, exito = configure_openai()

        if not exito or not modelo_seleccionado:
            logger.error("No se pudo configurar OpenAI para corregir idioma")
            return resultado

        # Función auxiliar para traducir texto
        def traducir_texto(texto, idioma_destino):
            if not texto or len(texto) < 5:
                return texto

            # Crear prompt para traducción
            system_msg = f"""
            Eres un traductor profesional especializado en terminología de enseñanza de idiomas.
            Traduce el siguiente texto al {idioma_destino} manteniendo el significado original y la terminología técnica.
            Devuelve SOLO el texto traducido, sin explicaciones ni comentarios adicionales.
            """

            user_msg = f"Traduce el siguiente texto al {idioma_destino}:\n\n{texto}"

            # Realizar traducción con OpenAI
            try:
                # Llamada directa a la API para evitar usar obtener_json_de_openai que espera estructura JSON
                headers = {
                    "Authorization": f"Bearer {api_keys['openai']}",
                    "Content-Type": "application/json"
                }

                data = {
                    "model": modelo_seleccionado,
                    "messages": [
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": user_msg}
                    ],
                    "max_tokens": 500,
                    "temperature": 0.3
                }

                response = requests.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=15
                )

                if response.status_code == 200:
                    response_data = response.json()
                    if "choices" in response_data and len(response_data["choices"]) > 0:
                        traduccion = response_data["choices"][0]["message"]["content"].strip(
                        )
                        return traduccion
            except Exception as e:
                logger.error(f"Error al traducir texto: {e}")

            # Si hay cualquier error, devolver el texto original
            return texto

        # Crear una copia del resultado para modificar
        resultado_corregido = copy.deepcopy(resultado)

        # Corregir saludo
        if "saludo" in resultado_corregido:
            resultado_corregido["saludo"] = traducir_texto(
                resultado_corregido["saludo"], idioma_solicitado)

        # Corregir explicaciones de errores
        errores = resultado_corregido.get("errores", {})
        if isinstance(errores, dict):
            for categoria, lista_errores in errores.items():
                if isinstance(lista_errores, list):
                    for error in lista_errores:
                        if isinstance(error, dict) and "explicacion" in error:
                            error["explicacion"] = traducir_texto(
                                error["explicacion"], idioma_solicitado)

        # Corregir análisis contextual
        analisis = resultado_corregido.get("analisis_contextual", {})
        if isinstance(analisis, dict):
            for categoria, datos in analisis.items():
                if isinstance(datos, dict):
                    if "comentario" in datos:
                        datos["comentario"] = traducir_texto(
                            datos["comentario"], idioma_solicitado)

                    if "sugerencias" in datos and isinstance(datos["sugerencias"], list):
                        datos["sugerencias"] = [traducir_texto(
                            sugerencia, idioma_solicitado) for sugerencia in datos["sugerencias"]]

                    if "adecuacion" in datos:
                        datos["adecuacion"] = traducir_texto(
                            datos["adecuacion"], idioma_solicitado)

                    if "elementos_destacables" in datos and isinstance(datos["elementos_destacables"], list):
                        datos["elementos_destacables"] = [traducir_texto(
                            elemento, idioma_solicitado) for elemento in datos["elementos_destacables"]]

        return resultado_corregido
    except Exception as e:
        logger.error(f"Error al corregir idioma: {e}")
        return resultado


def _corregir_texto_impl(texto, nombre, nivel, idioma, tipo_texto, contexto_cultural, info_adicional=""):
    """
    Implementación real de la corrección de texto.
    Actualizada para usar el asistente de OpenAI.

    Args:
        (Mismos parámetros que corregir_texto)

    Returns:
        dict: Resultado de la corrección o mensaje de error
    """
    try:
        # Validar la entrada
        if not texto or not nombre:
            return {"error": "El texto y el nombre son obligatorios."}

        # Mapeo de niveles para instrucciones más específicas
        nivel_map_instrucciones = {
            "Nivel principiante (A1-A2)": {
                "descripcion": "principiante (A1-A2)",
                "enfoque": "Enfócate en estructuras básicas, vocabulario fundamental y errores comunes. Utiliza explicaciones simples y claras. Evita terminología lingüística compleja."
            },
            "Nivel intermedio (B1-B2)": {
                "descripcion": "intermedio (B1-B2)",
                "enfoque": "Puedes señalar errores más sutiles de concordancia, uso de tiempos verbales y preposiciones. Puedes usar alguna terminología lingüística básica en las explicaciones."
            },
            "Nivel avanzado (C1-C2)": {
                "descripcion": "avanzado (C1-C2)",
                "enfoque": "Céntrate en matices, coloquialismos, registro lingüístico y fluidez. Puedes usar terminología lingüística específica y dar explicaciones más detalladas y técnicas."
            }
        }

        # Simplificar nivel si es necesario
        nivel_simple = nivel
        if nivel in ["principiante", "intermedio", "avanzado"]:
            # Ya está en formato simple
            pass
        elif "principiante" in nivel.lower() or "a1" in nivel.lower() or "a2" in nivel.lower():
            nivel_simple = "principiante"
        elif "intermedio" in nivel.lower() or "b1" in nivel.lower() or "b2" in nivel.lower():
            nivel_simple = "intermedio"
        elif "avanzado" in nivel.lower() or "c1" in nivel.lower() or "c2" in nivel.lower():
            nivel_simple = "avanzado"
        else:
            nivel_simple = "intermedio"  # Valor por defecto

        # Usar nivel intermedio como fallback para las instrucciones detalladas
        nivel_info = nivel_map_instrucciones.get(
            nivel, nivel_map_instrucciones["Nivel intermedio (B1-B2)"])

        # Instrucciones para el modelo de IA con análisis contextual avanzado
        system_message = f'''
Eres Diego, un profesor experto en ELE (Español como Lengua Extranjera) especializado en análisis lingüístico contextual.
Tu objetivo es corregir textos adaptando tu feedback al nivel {nivel_info["descripcion"]} del estudiante.
{nivel_info["enfoque"]}

IMPORTANTE: Debes proporcionar tu respuesta ÚNICAMENTE en formato JSON, siguiendo exactamente la estructura que se detalla a continuación.

Cuando corrijas un texto, DEBES devolver la respuesta en un JSON válido, sin texto adicional, con la siguiente estructura EXACTA:

{{
  "via_assistant": true,             // IMPORTANTE: siempre incluir este campo
  "saludo": "string",                // en {idioma}
  "tipo_texto": "string",            // en {idioma}
  "errores": {{
       "Gramática": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
           // más errores de Gramática (o [] si ninguno)
       ],
       "Léxico": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
       ],
       "Puntuación": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
       ],
       "Estructura textual": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
       ]
  }},
  "texto_corregido": "string",       // siempre en español
  "analisis_contextual": {{
       "coherencia": {{
           "puntuacion": number,     // del 1 al 10
           "comentario": "string",   // en {idioma}
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }},
       "cohesion": {{
           "puntuacion": number,     // del 1 al 10
           "comentario": "string",   // en {idioma}
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }},
       "registro_linguistico": {{
           "puntuacion": number,     // del 1 al 10
           "tipo_detectado": "string", // tipo de registro detectado en {idioma}
           "adecuacion": "string",   // evaluación de adecuación en {idioma}
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }},
       "adecuacion_cultural": {{
           "puntuacion": number,     // del 1 al 10
           "comentario": "string",   // en {idioma}
           "elementos_destacables": [  // elementos culturales destacables en {idioma}
               "string",
               "string"
           ],
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }}
  }},
  "consejo_final": "string",         // en español
  "fin": "Fin de texto corregido."
}}

IMPORTANTE:
- Tu respuesta debe estar SOLO en formato JSON como se especifica arriba.
- SIEMPRE incluye "via_assistant": true en la raíz del JSON
- Las explicaciones de los errores deben estar en {idioma}
- Todo el análisis contextual debe estar en {idioma}
- El texto corregido completo SIEMPRE debe estar en español, independientemente del idioma seleccionado
- El consejo final SIEMPRE debe estar en español
- Adapta tus explicaciones y sugerencias al nivel {nivel_info["descripcion"]} del estudiante
- Considera el tipo de texto "{tipo_texto}" y el contexto cultural "{contexto_cultural}" en tu análisis

No devuelvas ningún texto extra fuera de este JSON.
'''

        # Mensaje para el usuario con contexto adicional
        user_message = f'''
Texto del alumno:
"""
{texto}
"""
Nivel: {nivel}
Nombre del alumno: {nombre}
Idioma de corrección: {idioma}
Tipo de texto: {tipo_texto}
Contexto cultural: {contexto_cultural}
{f"Información adicional: {info_adicional}" if info_adicional else ""}

Por favor, provee un análisis completo en formato JSON siguiendo la estructura solicitada.
'''

        # Registrar que usamos el asistente para verificación
        logger.info("🤖 Usando asistente para corrección de texto")

        # SOLUCIÓN 1: Inicializar el gestor directamente en lugar de usar get_assistants_manager
        try:
            # Obtener la API key
            api_key = api_keys.get("openai")
            if api_key is None:
                return {"error": "API key de OpenAI no configurada"}

            # Crear instancia de OpenAIAssistants directamente
            assistants_manager = OpenAIAssistants(api_key)

            # Enviar solicitud al asistente de OpenAI
            raw_output, data_json = assistants_manager.get_completion(
                system_message,
                user_message,
                max_retries=3,
                task_type="correccion_texto"
            )
        except Exception as e:
            logger.error(
                f"Error al inicializar o usar el gestor de asistentes: {e}")
            return {"error": f"Error al procesar con asistentes de OpenAI: {str(e)}"}

        # Verificar si hay error en la respuesta
        if raw_output is None or "error" in data_json:
            error_msg = data_json.get(
                "error", "Error desconocido en el procesamiento")
            logger.error(f"Error en corrección: {error_msg}")
            return {"error": error_msg}

        # Añadir verificación de uso del asistente en el resultado
        if "via_assistant" not in data_json:
            # Añadir marcador manualmente si el asistente no lo incluyó
            data_json["via_assistant"] = True
            logger.info(
                "Añadido marcador via_assistant manualmente al resultado")

        # Guardar corrección en Firestore si está disponible
        if get_session_var("firebase_available", False):
            guardar_resultado = guardar_correccion_firestore(
                nombre, nivel_simple, idioma, texto, raw_output)

            if not guardar_resultado["success"]:
                logger.warning(
                    f"No se pudo guardar la corrección: {guardar_resultado['message']}")

        # Guardar el texto para posible uso futuro
        set_session_var("ultimo_texto", texto)

        # Devolver resultado
        return data_json

    except Exception as e:
        logger.error(f"Error al corregir texto: {str(e)}")
        logger.error(traceback.format_exc())
        return {"error": f"Error al corregir texto: {str(e)}"}


def _corregir_texto_degradado(texto, nivel):
    """
    Versión simplificada de corrección para modo degradado cuando la API no está disponible.

    Args:
        texto: Texto a corregir
        nivel: Nivel del estudiante

    Returns:
        dict: Resultado simplificado de la corrección
    """
    # Análisis básico del texto
    palabras = re.findall(r'\b\w+\b', texto.lower())
    num_palabras = len(palabras)

    # Crear un resultado básico
    return {
        "saludo": "Hola, he analizado tu texto en modo básico.",
        "tipo_texto": "Texto general",
        "errores": {
            "Gramática": [],
            "Léxico": [],
            "Puntuación": [],
            "Estructura textual": []
        },
        "texto_corregido": texto,  # Sin correcciones reales
        "analisis_contextual": {
            "coherencia": {
                "puntuacion": 7,
                "comentario": "No se pudo realizar un análisis detallado en modo básico.",
                "sugerencias": ["Revisa la conexión entre tus ideas.", "Asegúrate de que el texto sigue un orden lógico."]
            },
            "cohesion": {
                "puntuacion": 7,
                "comentario": "No se pudo realizar un análisis detallado en modo básico.",
                "sugerencias": ["Utiliza conectores para unir tus ideas.", "Evita repeticiones innecesarias."]
            },
            "registro_linguistico": {
                "puntuacion": 7,
                "tipo_detectado": "Neutro",
                "adecuacion": "El registro parece adecuado para un texto general.",
                "sugerencias": ["Adapta tu lenguaje al contexto y destinatario."]
            },
            "adecuacion_cultural": {
                "puntuacion": 7,
                "comentario": "No se pudo realizar un análisis cultural detallado.",
                "elementos_destacables": [],
                "sugerencias": ["Considera aspectos culturales relevantes para tu texto."]
            }
        },
        "consejo_final": f"Tu texto tiene aproximadamente {num_palabras} palabras. En modo básico no puedo ofrecer una corrección detallada. Por favor, intenta más tarde cuando el servicio completo esté disponible.",
        "fin": "Fin de texto corregido."
    }


# --- 3. GENERACIÓN DE AUDIO CON ELEVENLABS ---

def generar_audio_consejo(consejo_texto):
    """
    Genera un archivo de audio a partir del texto usando ElevenLabs.

    Args:
        consejo_texto: Texto a convertir en audio

    Returns:
        BytesIO: Buffer con el audio generado, o None si ocurre un error
    """
    if not api_keys["elevenlabs"]["api_key"] or not api_keys["elevenlabs"]["voice_id"]:
        logger.warning("Claves de ElevenLabs no configuradas")
        return None

    if not circuit_breaker.can_execute("elevenlabs"):
        logger.warning("ElevenLabs temporalmente no disponible")
        return None

    if not consejo_texto:
        return None

    # Limpiar el texto
    audio_text = consejo_texto.replace("Consejo final:", "").strip()
    if not audio_text:
        return None

    try:
        elevenlabs_api_key = api_keys["elevenlabs"]["api_key"]
        elevenlabs_voice_id = api_keys["elevenlabs"]["voice_id"]

        tts_url = f"https://api.elevenlabs.io/v1/text-to-speech/{elevenlabs_voice_id}"
        headers = {
            "xi-api-key": elevenlabs_api_key,
            "Content-Type": "application/json"
        }
        data = {
            "text": audio_text,
            "model_id": "eleven_multilingual_v2",
            "voice_settings": {
                "stability": 0.3,
                "similarity_boost": 0.9
            }
        }

        # Función para envío de solicitud con timeout más largo
        def send_request():
            response = requests.post(
                tts_url, headers=headers, json=data, timeout=20)
            response.raise_for_status()  # Levantar excepción si hay error
            return response

        # Usar sistema de reintentos
        response_audio = retry_with_backoff(send_request, max_retries=2)

        if response_audio.ok:
            audio_bytes = BytesIO(response_audio.content)
            circuit_breaker.record_success("elevenlabs")
            set_session_var("servicios_disponibles", {
                **get_session_var("servicios_disponibles", {}),
                "elevenlabs": True
            })
            return audio_bytes
        else:
            logger.error(
                f"Error en ElevenLabs API: {response_audio.status_code}")
            circuit_breaker.record_failure(
                "elevenlabs", error_type="api_error")
            return None

    except requests.exceptions.Timeout:
        logger.error("Timeout en solicitud a ElevenLabs")
        circuit_breaker.record_failure("elevenlabs", error_type="timeout")
        return None
    except Exception as e:
        logger.error(f"Error al generar audio: {str(e)}")
        circuit_breaker.record_failure("elevenlabs", error_type="general")
        return None


# --- 4. OCR PARA TEXTOS MANUSCRITOS CON OPENAI ---

def transcribir_imagen_texto(imagen_bytes, idioma="es"):
    """
    Transcribe texto manuscrito de una imagen utilizando la API de OpenAI.
    Mejorado con manejo robusto de errores.

    Args:
        imagen_bytes: Bytes de la imagen a transcribir
        idioma: Código de idioma (es, en, fr)

    Returns:
        str: Texto transcrito o mensaje de error
    """
    if api_keys["openai"] is None:
        return "Error: API de OpenAI no disponible"

    if not circuit_breaker.can_execute("openai"):
        return "Error: Servicio OpenAI temporalmente no disponible"

    try:
        # Configurar OpenAI
        modelo_seleccionado, exito = configure_openai()
        if not exito:
            return "Error: No se pudo configurar OpenAI correctamente"

        # Codificar la imagen a base64
        import base64
        encoded_image = base64.b64encode(imagen_bytes).decode('utf-8')

        # Preparar la solicitud a la API de visión de OpenAI
        headers = {
            "Authorization": f"Bearer {api_keys['openai']}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "gpt-4-vision-preview",  # Modelo con capacidades de visión
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"Transcribe todo el texto manuscrito visible en esta imagen. Devuelve SOLO el texto transcrito, sin comentarios adicionales. El texto está en {idioma}."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded_image}"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 1000
        }

        # Función para envío de solicitud con timeout adaptativo
        def send_ocr_request():
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=45  # Timeout más largo para procesamiento de imágenes
            )

            if response.status_code != 200:
                raise Exception(
                    f"Error en la API de visión de OpenAI: {response.status_code} - {response.text}")

            response_data = response.json()

            # Extraer el contenido del mensaje de respuesta
            if "choices" in response_data and len(response_data["choices"]) > 0:
                content = response_data["choices"][0]["message"]["content"]
                return content
            else:
                raise Exception("Formato de respuesta inesperado de OpenAI")

        # Usar sistema de reintentos
        response = retry_with_backoff(
            send_ocr_request, max_retries=2, initial_delay=2)

        # Registrar éxito
        circuit_breaker.record_success("openai")

        # Limpiar la respuesta (eliminar comillas si las hay al principio/final)
        cleaned_response = response.strip()
        if cleaned_response.startswith('"') and cleaned_response.endswith('"'):
            cleaned_response = cleaned_response[1:-1]

        return cleaned_response

    except requests.exceptions.Timeout:
        logger.error("Timeout en solicitud de transcripción de imagen")
        circuit_breaker.record_failure("openai", error_type="vision_timeout")
        return "Error: La solicitud de transcripción excedió el tiempo de espera. Por favor, intenta con una imagen más pequeña o en otro momento."
    except Exception as e:
        logger.error(f"Error en transcribir_imagen_texto: {str(e)}")
        logger.error(traceback.format_exc())

        circuit_breaker.record_failure("openai", error_type="vision_error")

        return f"Error en la transcripción: {str(e)}"


# --- 5. GENERACIÓN DE IMÁGENES CON DALL-E ---

def generar_imagen_dalle(tema, nivel):
    """
    Genera una imagen utilizando DALL-E basada en un tema y adaptada al nivel del estudiante.
    Mejorado con manejo robusto de errores y timeouts.

    Args:
        tema: Tema para la imagen
        nivel: Nivel de español (principiante, intermedio, avanzado)

    Returns:
        tuple: (URL de la imagen generada, descripción de la imagen)
    """
    if api_keys["dalle"] is None:
        return None, "API de DALL-E no disponible"

    if not circuit_breaker.can_execute("dalle"):
        return None, "Servicio DALL-E temporalmente no disponible"

    # Adaptar la complejidad del prompt según el nivel
    if "principiante" in nivel.lower():
        complejidad = "simple con objetos y personas claramente identificables"
    elif "intermedio" in nivel.lower():
        complejidad = "con detalles moderados y una escena cotidiana con varios elementos"
    else:
        complejidad = "detallada con múltiples elementos, que pueda generar descripciones complejas"

    # Crear el prompt para DALL-E
    prompt = f"Una escena {complejidad} sobre {tema}. La imagen debe ser clara, bien iluminada, y adecuada para describir en español."

    try:
        # Configurar headers para la API de DALL-E
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_keys['dalle']}"
        }

        # URL de la API
        url = "https://api.openai.com/v1/images/generations"

        # Datos para la solicitud
        data = {
            "model": "dall-e-3",
            "prompt": prompt,
            "n": 1,
            "size": "1024x1024",
            "quality": "standard"
        }

        # Función para envío de solicitud con timeout adaptativo
        def generate_image():
            response = requests.post(
                url,
                headers=headers,
                json=data,
                timeout=60  # Timeout más largo para generación de imágenes
            )
            response.raise_for_status()
            return response.json()

        # Usar sistema de reintentos
        response_data = retry_with_backoff(
            generate_image, max_retries=2, initial_delay=2)

        # Obtener la URL de la imagen
        imagen_url = response_data["data"][0]["url"]

        # Ahora generamos la descripción con OpenAI
        modelo_seleccionado, exito = configure_openai()
        if not exito:
            # Usar una descripción genérica si OpenAI no está disponible
            descripcion = f"Una imagen sobre {tema}. Práctica describiendo lo que ves en español."
            return imagen_url, descripcion

        # Preparar la llamada para generar la descripción
        headers_desc = {
            "Authorization": f"Bearer {api_keys['openai']}",
            "Content-Type": "application/json"
        }

        data_desc = {
            "model": modelo_seleccionado,
            "messages": [
                {
                    "role": "user",
                    "content": f"""
                    Crea una descripción en español de esta imagen generada para un estudiante de nivel {nivel}.
                    
                    La descripción debe:
                    1. Ser apropiada para el nivel {nivel}
                    2. Utilizar vocabulario y estructuras gramaticales de ese nivel
                    3. Incluir entre 3-5 preguntas al final para que el estudiante practique describiendo la imagen
                    
                    Tema de la imagen: {tema}
                    """
                }
            ],
            "temperature": 0.7,
            "max_tokens": 500
        }

        def generate_description():
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers_desc,
                json=data_desc,
                timeout=30
            )

            if response.status_code != 200:
                raise Exception(
                    f"Error en la API de OpenAI: {response.status_code} - {response.text}")

            response_data = response.json()

            # Extraer el contenido del mensaje de respuesta
            if "choices" in response_data and len(response_data["choices"]) > 0:
                content = response_data["choices"][0]["message"]["content"]
                return content
            else:
                raise Exception("Formato de respuesta inesperado de OpenAI")

        descripcion = retry_with_backoff(generate_description, max_retries=2)

        # Registrar éxito
        circuit_breaker.record_success("dalle")
        circuit_breaker.record_success("openai")

        return imagen_url, descripcion

    except requests.exceptions.Timeout:
        logger.error("Timeout en generación de imagen")
        circuit_breaker.record_failure("dalle", error_type="timeout")
        return None, "Error: La generación de imagen ha superado el tiempo de espera. Por favor, inténtalo más tarde."
    except Exception as e:
        handle_exception("generar_imagen_dalle", e)
        circuit_breaker.record_failure("dalle", error_type="general")
        return None, f"Error: {str(e)}"


# Función auxiliar para validar texto de entrada
def validar_texto_entrada(texto, max_length=10000):
    """
    Valida y limpia el texto de entrada para prevenir problemas.

    Args:
        texto: Texto a validar
        max_length: Longitud máxima permitida

    Returns:
        str: Texto validado o mensaje de error
    """
    if not texto:
        return ""

    # Eliminar caracteres potencialmente problemáticos
    texto_limpio = re.sub(r'[^\w\s.,;:?!¿¡()\-\'\"áéíóúÁÉÍÓÚñÑüÜ]', '', texto)

    # Truncar texto muy largo
    if len(texto_limpio) > max_length:
        return texto_limpio[:max_length] + "... (texto truncado por exceder longitud máxima)"

    return texto_limpio

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 6: Análisis y Visualización
# ==================================================================================
#
# Este artefacto contiene:
# 1. Funciones para análisis estadístico de textos
# 2. Funciones para generación de gráficos y visualizaciones
# 3. Funciones para procesamiento de datos en DataFrames
# 4. Utilidades de análisis lingüístico
#
# Estas funciones proporcionan capacidades de análisis avanzado y visualización
# de datos para mejorar la interpretación de correcciones y progreso.
# ==================================================================================

# --- 1. FUNCIONES PARA ANÁLISIS ESTADÍSTICO ---


def analizar_estadisticas_texto(texto, nivel_estudiante="intermedio"):
    """
    Analiza estadísticas básicas de un texto.

    Args:
        texto: Texto a analizar
        nivel_estudiante: Nivel del estudiante

    Returns:
        dict: Estadísticas del texto
    """
    try:
        # Limpiar el texto
        texto_limpio = texto.strip()

        # Contar palabras
        palabras = re.findall(r'\b\w+\b', texto_limpio.lower())
        num_palabras = len(palabras)

        # Contar frases
        frases = re.split(r'[.!?]+', texto_limpio)
        frases = [f for f in frases if f.strip()]
        num_frases = len(frases)

        # Contar párrafos
        parrafos = texto_limpio.split('\n\n')
        parrafos = [p for p in parrafos if p.strip()]
        num_parrafos = len(parrafos)

        # Longitud media de palabra
        long_total = sum(len(palabra) for palabra in palabras)
        long_media_palabra = long_total / num_palabras if num_palabras > 0 else 0

        # Longitud media de frase (palabras por frase)
        palabras_por_frase = num_palabras / num_frases if num_frases > 0 else 0

        # Calcular proporción de palabras únicas (type-token ratio)
        palabras_unicas = len(set(palabras))
        ttr = palabras_unicas / num_palabras if num_palabras > 0 else 0

        # Expectativas por nivel
        expectativas = {
            "principiante": {
                "palabras_min": 50,
                "palabras_max": 150,
                "palabras_por_frase_ideal": 8,
                "ttr_ideal": 0.6
            },
            "intermedio": {
                "palabras_min": 150,
                "palabras_max": 300,
                "palabras_por_frase_ideal": 12,
                "ttr_ideal": 0.7
            },
            "avanzado": {
                "palabras_min": 250,
                "palabras_max": 500,
                "palabras_por_frase_ideal": 15,
                "ttr_ideal": 0.8
            }
        }

        # Usar nivel intermedio como fallback
        nivel_exp = expectativas.get(
            nivel_estudiante, expectativas["intermedio"])

        # Determinar si cumple con expectativas de longitud
        cumple_min = num_palabras >= nivel_exp["palabras_min"]
        cumple_max = num_palabras <= nivel_exp["palabras_max"]
        cumple_longitud = cumple_min and cumple_max

        # Complejidad léxica relativa al nivel
        if ttr < nivel_exp["ttr_ideal"] - 0.15:
            complejidad_lexica = "baja"
        elif ttr > nivel_exp["ttr_ideal"] + 0.15:
            complejidad_lexica = "alta"
        else:
            complejidad_lexica = "adecuada"

        # Complejidad sintáctica basada en palabras por frase
        if palabras_por_frase < nivel_exp["palabras_por_frase_ideal"] - 4:
            complejidad_sintactica = "baja"
        elif palabras_por_frase > nivel_exp["palabras_por_frase_ideal"] + 4:
            complejidad_sintactica = "alta"
        else:
            complejidad_sintactica = "adecuada"

        return {
            "num_palabras": num_palabras,
            "num_frases": num_frases,
            "num_parrafos": num_parrafos,
            "longitud_media_palabra": round(long_media_palabra, 2),
            "palabras_por_frase": round(palabras_por_frase, 2),
            "type_token_ratio": round(ttr, 3),
            "palabras_unicas": palabras_unicas,
            "cumple_expectativas_longitud": cumple_longitud,
            "complejidad_lexica": complejidad_lexica,
            "complejidad_sintactica": complejidad_sintactica,
            "recomendaciones": generar_recomendaciones_estadisticas(
                num_palabras, palabras_por_frase, ttr, nivel_exp, nivel_estudiante
            )
        }

    except Exception as e:
        logger.error(f"Error en análisis estadístico: {str(e)}")
        return {
            "error": f"Error al analizar estadísticas: {str(e)}",
            "num_palabras": 0,
            "num_frases": 0,
            "num_parrafos": 0
        }


def generar_recomendaciones_estadisticas(num_palabras, palabras_por_frase, ttr, nivel_exp, nivel_estudiante):
    """
    Genera recomendaciones basadas en estadísticas del texto.

    Args:
        num_palabras: Número de palabras
        palabras_por_frase: Promedio de palabras por frase
        ttr: Type-token ratio
        nivel_exp: Expectativas para el nivel
        nivel_estudiante: Nivel del estudiante

    Returns:
        list: Recomendaciones generadas
    """
    recomendaciones = []

    # Recomendaciones sobre longitud
    if num_palabras < nivel_exp["palabras_min"]:
        recomendaciones.append(f"El texto es más corto de lo esperado para nivel {nivel_estudiante}. "
                               f"Intenta desarrollar más tus ideas (objetivo: {nivel_exp['palabras_min']}-{nivel_exp['palabras_max']} palabras).")
    elif num_palabras > nivel_exp["palabras_max"]:
        recomendaciones.append(f"El texto es más largo de lo esperado para nivel {nivel_estudiante}. "
                               f"Intenta ser más conciso (objetivo: {nivel_exp['palabras_min']}-{nivel_exp['palabras_max']} palabras).")

    # Recomendaciones sobre complejidad sintáctica
    if palabras_por_frase < nivel_exp["palabras_por_frase_ideal"] - 4:
        recomendaciones.append(
            "Tus frases son muy cortas. Intenta combinar ideas utilizando conectores y subordinadas.")
    elif palabras_por_frase > nivel_exp["palabras_por_frase_ideal"] + 4:
        recomendaciones.append(
            "Tus frases son muy largas. Intenta dividirlas para mejorar la claridad.")

    # Recomendaciones sobre diversidad léxica
    if ttr < nivel_exp["ttr_ideal"] - 0.15:
        recomendaciones.append(
            "Tu vocabulario es poco variado. Intenta utilizar sinónimos y evitar repeticiones.")

    return recomendaciones


# --- 2. FUNCIONES PARA GRÁFICOS Y VISUALIZACIONES ---

def crear_grafico_radar(valores, categorias):
    """
    Crea un gráfico de radar para visualizar puntuaciones en diferentes categorías.

    Args:
        valores: Lista de valores (puntuaciones) para cada categoría
        categorias: Lista de nombres de las categorías

    Returns:
        matplotlib.figure.Figure: Figura con el gráfico generado
    """
    try:
        # Convertir a array numpy
        valores = np.array(valores)

        # Asegurar que todas las puntuaciones estén entre 0 y 10
        valores = np.clip(valores, 0, 10)

        # Número de variables
        N = len(categorias)

        # Crear ángulos para cada categoría (igualmente espaciados)
        angulos = [n / float(N) * 2 * np.pi for n in range(N)]

        # Cerrar el gráfico repitiendo el primer punto
        valores = np.append(valores, valores[0])
        angulos = np.append(angulos, angulos[0])
        categorias = np.append(categorias, categorias[0])

        # Crear figura
        fig = plt.figure(figsize=(8, 8))
        ax = fig.add_subplot(111, polar=True)

        # Dibujar el gráfico
        ax.plot(angulos, valores, 'o-', linewidth=2)
        ax.fill(angulos, valores, alpha=0.25)

        # Añadir etiquetas
        ax.set_thetagrids(np.degrees(angulos), labels=categorias)

        # Fijar límites del eje radial
        ax.set_ylim(0, 10)

        # Añadir título
        plt.title('Análisis por categorías', size=12, fontweight='bold')

        # Mejorar apariencia
        ax.grid(True)

        # Devolver la figura
        return fig

    except Exception as e:
        logger.error(f"Error al crear gráfico radar: {str(e)}")
        # Crear una figura de error simple
        fig = plt.figure(figsize=(4, 3))
        plt.text(0.5, 0.5, f"Error al generar gráfico: {str(e)}",
                 ha='center', va='center', wrap=True)
        plt.axis('off')
        return fig


def crear_grafico_barras_errores(historial_df, tipo_error):
    """
    Crea un gráfico de barras para un tipo específico de error a lo largo del tiempo.

    Args:
        historial_df: DataFrame con historial de correcciones
        tipo_error: Tipo de error a visualizar

    Returns:
        altair.Chart: Gráfico generado con Altair
    """
    try:
        if historial_df is None or historial_df.empty:
            return None

        # Verificar si existe la columna de fecha
        fecha_col = None
        for col in historial_df.columns:
            if 'fecha' in col.lower():
                fecha_col = col
                break

        if fecha_col is None or tipo_error not in historial_df.columns:
            return None

        # Asegurar formato de fecha
        historial_df[fecha_col] = pd.to_datetime(
            historial_df[fecha_col], errors='coerce')

        # Asegurar que el tipo de error es numérico
        historial_df[tipo_error] = pd.to_numeric(
            historial_df[tipo_error], errors='coerce').fillna(0)

        # Ordenar por fecha
        df_ordenado = historial_df.sort_values(by=fecha_col)

        # Crear gráfico de barras
        chart = alt.Chart(df_ordenado).mark_bar().encode(
            x=alt.X(f'{fecha_col}:T', title='Fecha'),
            y=alt.Y(f'{tipo_error}:Q', title=tipo_error),
            tooltip=[f'{fecha_col}:T', f'{tipo_error}:Q']
        ).properties(
            title=f'Evolución de {tipo_error} a lo largo del tiempo',
            width=600,
            height=300
        ).interactive()

        return chart

    except Exception as e:
        logger.error(f"Error al crear gráfico de barras: {str(e)}")
        return None


def crear_grafico_tendencia(historial_df, columna, titulo=None):
    """
    Crea un gráfico de tendencia para una columna específica a lo largo del tiempo.

    Args:
        historial_df: DataFrame con historial de correcciones
        columna: Nombre de la columna a visualizar
        titulo: Título opcional para el gráfico

    Returns:
        altair.Chart: Gráfico generado con Altair
    """
    try:
        if historial_df is None or historial_df.empty:
            return None

        # Verificar si existe la columna de fecha
        fecha_col = None
        for col in historial_df.columns:
            if 'fecha' in col.lower():
                fecha_col = col
                break

        if fecha_col is None or columna not in historial_df.columns:
            return None

        # Asegurar formato de fecha
        historial_df[fecha_col] = pd.to_datetime(
            historial_df[fecha_col], errors='coerce')

        # Asegurar que la columna es numérica
        historial_df[columna] = pd.to_numeric(
            historial_df[columna], errors='coerce').fillna(0)

        # Ordenar por fecha
        df_ordenado = historial_df.sort_values(by=fecha_col)

        # Título por defecto si no se proporciona
        if titulo is None:
            titulo = f'Evolución de {columna}'

        # Crear gráfico de línea
        chart = alt.Chart(df_ordenado).mark_line(point=True).encode(
            x=alt.X(f'{fecha_col}:T', title='Fecha'),
            y=alt.Y(f'{columna}:Q', title=columna),
            tooltip=[f'{fecha_col}:T', f'{columna}:Q']
        ).properties(
            title=titulo,
            width=600,
            height=300
        ).interactive()

        return chart

    except Exception as e:
        logger.error(f"Error al crear gráfico de tendencia: {str(e)}")
        return None


def crear_grafico_evolucion_errores(historial_df):
    """
    Crea un gráfico de líneas para mostrar la evolución de errores por categoría.

    Args:
        historial_df: DataFrame con el historial de correcciones

    Returns:
        altair.Chart: Gráfico generado con Altair
    """
    try:
        if historial_df is None or historial_df.empty:
            return None

        # Asegurar que tenemos una columna de fecha
        if 'Fecha' not in historial_df.columns:
            return None

        # Convertir a datetime y ordenar
        historial_df['Fecha'] = pd.to_datetime(historial_df['Fecha'])
        df_ordenado = historial_df.sort_values('Fecha')

        # Preparar datos para gráfico
        tipos_error = ['Errores Gramática', 'Errores Léxico',
                       'Errores Puntuación', 'Errores Estructura']

        # Verificar que existen las columnas
        tipos_disponibles = [
            col for col in tipos_error if col in historial_df.columns]

        if not tipos_disponibles:
            return None

        # Melt para formato adecuado para Altair
        df_melt = pd.melt(
            df_ordenado,
            id_vars=['Fecha'],
            value_vars=tipos_disponibles,
            var_name='Tipo de Error',
            value_name='Cantidad'
        )

        # Crear gráfico con Altair
        chart = alt.Chart(df_melt).mark_line(point=True).encode(
            x=alt.X('Fecha:T', title='Fecha',
                    axis=alt.Axis(format='%d/%m/%Y')),
            y=alt.Y('Cantidad:Q', title='Número de errores'),
            color=alt.Color('Tipo de Error:N',
                            legend=alt.Legend(orient='top')),
            tooltip=['Fecha:T', 'Tipo de Error:N', 'Cantidad:Q']
        ).properties(
            width=600,
            height=300,
            title='Evolución de errores por categoría'
        ).interactive()

        return chart

    except Exception as e:
        logger.error(
            f"Error al crear gráfico de evolución de errores: {str(e)}")
        return None


def crear_grafico_progreso_puntuaciones(historial_df):
    """
    Crea un gráfico para mostrar el progreso en las puntuaciones de análisis contextual.

    Args:
        historial_df: DataFrame con historial de correcciones

    Returns:
        altair.Chart: Gráfico generado con Altair
    """
    try:
        if historial_df is None or historial_df.empty:
            return None

        # Asegurar que tenemos una columna de fecha
        if 'Fecha' not in historial_df.columns:
            return None

        # Convertir a datetime y ordenar
        historial_df['Fecha'] = pd.to_datetime(historial_df['Fecha'])
        df_ordenado = historial_df.sort_values('Fecha')

        # Preparar datos para gráfico
        puntuaciones = ['Puntuación Coherencia', 'Puntuación Cohesión',
                        'Puntuación Registro', 'Puntuación Adecuación Cultural']

        # Verificar que existen las columnas
        puntuaciones_disponibles = [
            col for col in puntuaciones if col in historial_df.columns]

        if not puntuaciones_disponibles:
            return None

        # Melt para formato adecuado para Altair
        df_melt = pd.melt(
            df_ordenado,
            id_vars=['Fecha'],
            value_vars=puntuaciones_disponibles,
            var_name='Aspecto',
            value_name='Puntuación'
        )

        # Crear gráfico con Altair
        chart = alt.Chart(df_melt).mark_line(point=True).encode(
            x=alt.X('Fecha:T', title='Fecha',
                    axis=alt.Axis(format='%d/%m/%Y')),
            y=alt.Y('Puntuación:Q', scale=alt.Scale(
                domain=[0, 10]), title='Puntuación (0-10)'),
            color=alt.Color('Aspecto:N', legend=alt.Legend(orient='top')),
            tooltip=['Fecha:T', 'Aspecto:N', 'Puntuación:Q']
        ).properties(
            width=600,
            height=300,
            title='Progreso en habilidades lingüísticas'
        ).interactive()

        return chart

    except Exception as e:
        logger.error(
            f"Error al crear gráfico de progreso de puntuaciones: {str(e)}")
        return None


def crear_grafico_distribucion_errores(historial_df):
    """
    Crea un gráfico de barras para mostrar la distribución de errores por categoría.

    Args:
        historial_df: DataFrame con historial de correcciones

    Returns:
        altair.Chart: Gráfico generado con Altair
    """
    try:
        if historial_df is None or historial_df.empty:
            return None

        # Identificar columnas de errores
        columnas_errores = [col for col in historial_df.columns if col.startswith(
            'Errores ') and col != 'Total Errores']

        if not columnas_errores:
            return None

        # Calcular promedio de errores por categoría
        promedios = {}
        for col in columnas_errores:
            # Extraer nombre de categoría (quitar "Errores ")
            categoria = col[8:]
            promedios[categoria] = historial_df[col].mean()

        # Convertir a DataFrame para Altair
        df_promedios = pd.DataFrame({
            'Categoría': list(promedios.keys()),
            'Errores promedio': list(promedios.values())
        })

        # Ordenar por cantidad de errores
        df_promedios = df_promedios.sort_values(
            'Errores promedio', ascending=False)

        # Crear gráfico de barras
        chart = alt.Chart(df_promedios).mark_bar().encode(
            x=alt.X('Errores promedio:Q', title='Errores promedio por texto'),
            y=alt.Y('Categoría:N', sort='-x', title=''),
            color=alt.Color('Categoría:N', legend=None),
            tooltip=['Categoría:N', 'Errores promedio:Q']
        ).properties(
            title='Distribución de errores por categoría',
            width=600,
            height=200
        )

        return chart

    except Exception as e:
        logger.error(
            f"Error al crear gráfico de distribución de errores: {str(e)}")
        return None


# --- 3. FUNCIONES PARA PROCESAMIENTO DE DATAFRAMES ---

def calcular_estadisticas_progreso(historial_df):
    """
    Calcula estadísticas de progreso basadas en un historial de correcciones.

    Args:
        historial_df: DataFrame con historial de correcciones

    Returns:
        dict: Estadísticas de progreso calculadas
    """
    try:
        estadisticas = {
            "num_correcciones": 0,
            "primera_correccion": None,
            "ultima_correccion": None,
            "tendencia_errores": "estable",
            "media_errores_totales": 0,
            "mejora_porcentual": 0,
            "area_mejora_principal": None,
            "puntuacion_promedio": 0
        }

        if historial_df is None or historial_df.empty:
            return estadisticas

        # Número de correcciones
        estadisticas["num_correcciones"] = len(historial_df)

        # Fechas
        fecha_col = None
        for col in historial_df.columns:
            if 'fecha' in col.lower():
                fecha_col = col
                break

        if fecha_col:
            # Convertir a datetime si no lo es ya
            historial_df[fecha_col] = pd.to_datetime(
                historial_df[fecha_col], errors='coerce')

            # Ordenar por fecha
            historial_ordenado = historial_df.sort_values(by=fecha_col)

            # Primera y última corrección
            if not historial_ordenado.empty:
                estadisticas["primera_correccion"] = historial_ordenado.iloc[0][fecha_col]
                estadisticas["ultima_correccion"] = historial_ordenado.iloc[-1][fecha_col]

        # Tendencia de errores y mejora porcentual
        if 'Total Errores' in historial_df.columns and len(historial_df) >= 2:
            # Convertir a numérico
            historial_df['Total Errores'] = pd.to_numeric(
                historial_df['Total Errores'], errors='coerce').fillna(0)

            # Calcular estadísticas
            errores_ordenados = historial_ordenado['Total Errores'].values

            # Media de errores
            estadisticas["media_errores_totales"] = round(
                np.mean(errores_ordenados), 2)

            # Verificar si hay suficientes datos para analizar tendencia
            if len(errores_ordenados) >= 3:
                # Calcular tendencia simple comparando primer tercio vs último tercio
                n = len(errores_ordenados)
                primer_tercio = np.mean(
                    errores_ordenados[:n//3]) if n >= 3 else errores_ordenados[0]
                ultimo_tercio = np.mean(
                    errores_ordenados[-n//3:]) if n >= 3 else errores_ordenados[-1]

                # Determinar tendencia
                if ultimo_tercio < primer_tercio * 0.8:
                    estadisticas["tendencia_errores"] = "mejora"
                elif ultimo_tercio > primer_tercio * 1.2:
                    estadisticas["tendencia_errores"] = "empeora"
                else:
                    estadisticas["tendencia_errores"] = "estable"

                # Calcular mejora porcentual entre primera y última corrección
                primera = errores_ordenados[0]
                ultima = errores_ordenados[-1]

                if primera > 0:
                    mejora = ((primera - ultima) / primera) * 100
                    estadisticas["mejora_porcentual"] = round(mejora, 1)

        # Áreas de mejora
        tipos_errores = ['Errores Gramática', 'Errores Léxico',
                         'Errores Puntuación', 'Errores Estructura']
        errores_por_tipo = {}

        for tipo in tipos_errores:
            if tipo in historial_df.columns:
                # Convertir a numérico
                historial_df[tipo] = pd.to_numeric(
                    historial_df[tipo], errors='coerce').fillna(0)
                errores_por_tipo[tipo] = historial_df[tipo].mean()

        if errores_por_tipo:
            # Encontrar el área con más errores en promedio
            area_mejora = max(errores_por_tipo.items(), key=lambda x: x[1])
            estadisticas["area_mejora_principal"] = area_mejora[0].replace(
                'Errores ', '')

            # Guardar área para usar en recomendaciones
            set_session_var("ultima_area_mejora",
                            estadisticas["area_mejora_principal"])

        # Puntuación promedio (Coherencia, Cohesión, etc.)
        puntuaciones = ['Puntuación Coherencia', 'Puntuación Cohesión',
                        'Puntuación Registro', 'Puntuación Adecuación Cultural']

        puntuaciones_valores = []

        for puntuacion in puntuaciones:
            if puntuacion in historial_df.columns:
                # Convertir a numérico
                historial_df[puntuacion] = pd.to_numeric(
                    historial_df[puntuacion], errors='coerce').fillna(0)

                # Añadir a la lista de valores
                puntuaciones_valores.extend(historial_df[puntuacion].values)

        if puntuaciones_valores:
            estadisticas["puntuacion_promedio"] = round(
                np.mean(puntuaciones_valores), 2)

        # Comprobar si hay una puntuación global
        if 'Puntuación Global' in historial_df.columns:
            historial_df['Puntuación Global'] = pd.to_numeric(
                historial_df['Puntuación Global'], errors='coerce').fillna(0)
            estadisticas["puntuacion_global"] = round(
                historial_df['Puntuación Global'].mean(), 2)

        return estadisticas

    except Exception as e:
        logger.error(f"Error al calcular estadísticas de progreso: {str(e)}")
        return {
            "error": f"Error al calcular estadísticas: {str(e)}",
            "num_correcciones": len(historial_df) if historial_df is not None else 0
        }


def generar_recomendaciones_personalizadas(estadisticas, nivel):
    """
    Genera recomendaciones personalizadas basadas en las estadísticas de progreso.

    Args:
        estadisticas: Estadísticas de progreso
        nivel: Nivel del estudiante

    Returns:
        dict: Recomendaciones personalizadas
    """
    try:
        area_mejora = estadisticas.get("area_mejora_principal")
        tendencia = estadisticas.get("tendencia_errores")
        num_correcciones = estadisticas.get("num_correcciones", 0)

        recomendaciones = {
            "mensaje_principal": "",
            "recursos": [],
            "ejercicios": [],
            "siguiente_paso": ""
        }

        # Mensaje principal según tendencia
        if tendencia == "mejora":
            recomendaciones["mensaje_principal"] = (
                f"¡Felicidades! Has mejorado significativamente en tu escritura en español. "
                f"Continúa practicando regularmente para seguir progresando."
            )
        elif tendencia == "empeora":
            recomendaciones["mensaje_principal"] = (
                f"Has tenido algunos desafíos recientemente. Esto es normal en el proceso de aprendizaje. "
                f"Te recomendamos enfocarte en ejercicios específicos y quizás simplificar temporalmente tus textos."
            )
        else:
            recomendaciones["mensaje_principal"] = (
                f"Tu progreso se mantiene estable. Para avanzar al siguiente nivel, "
                f"intenta incorporar estructuras y vocabulario más variados en tus textos."
            )

        # Recomendaciones específicas según área principal de mejora
        if area_mejora == "Gramática":
            if nivel == "principiante":
                recomendaciones["recursos"].append({
                    "titulo": "Presente de indicativo en español",
                    "url": "https://www.profedeele.es/gramatica/presente-indicativo/"
                })
                recomendaciones["ejercicios"].append(
                    "Practica la conjugación de verbos regulares e irregulares en presente")
                recomendaciones["siguiente_paso"] = "Concentra tu atención en la concordancia sujeto-verbo y género-número"
            elif nivel == "intermedio":
                recomendaciones["recursos"].append({
                    "titulo": "Los tiempos del pasado en español",
                    "url": "https://www.profedeele.es/gramatica/tiempos-pasado-contraste/"
                })
                recomendaciones["ejercicios"].append(
                    "Escribe una narración en pasado utilizando correctamente pretérito indefinido e imperfecto")
                recomendaciones["siguiente_paso"] = "Practica con el subjuntivo en oraciones complejas"
            else:  # avanzado
                recomendaciones["recursos"].append({
                    "titulo": "Usos avanzados del subjuntivo",
                    "url": "https://www.profedeele.es/gramatica/subjuntivo-usos-avanzados/"
                })
                recomendaciones["ejercicios"].append(
                    "Escribe textos utilizando construcciones condicionales complejas")
                recomendaciones["siguiente_paso"] = "Perfecciona el uso de conectores y marcadores discursivos"

        elif area_mejora == "Léxico":
            if nivel == "principiante":
                recomendaciones["recursos"].append({
                    "titulo": "Vocabulario básico de español",
                    "url": "https://cvc.cervantes.es/ensenanza/biblioteca_ele/plan_curricular/niveles/09_nociones_especificas_inventario_a1-a2.htm"
                })
                recomendaciones["ejercicios"].append(
                    "Crea un diccionario personal con palabras nuevas")
                recomendaciones["siguiente_paso"] = "Incorpora 5-10 palabras nuevas en cada texto que escribas"
            elif nivel == "intermedio":
                recomendaciones["recursos"].append({
                    "titulo": "Expresiones idiomáticas comunes",
                    "url": "https://www.spanishpodcast.net/podcasts/expresiones-idiomaticas-en-espanol/"
                })
                recomendaciones["ejercicios"].append(
                    "Busca sinónimos para las palabras que más repites")
                recomendaciones["siguiente_paso"] = "Integra expresiones idiomáticas en tus textos"
            else:  # avanzado
                recomendaciones["recursos"].append({
                    "titulo": "Precisión léxica en español académico",
                    "url": "https://www.profedeele.es/actividad/lexico/precision-lexica/"
                })
                recomendaciones["ejercicios"].append(
                    "Reescribe un texto usando un vocabulario más preciso y técnico")
                recomendaciones["siguiente_paso"] = "Explora campos semánticos especializados según tus intereses"

        elif area_mejora == "Puntuación":
            recomendaciones["recursos"].append({
                "titulo": "Ortografía de la lengua española",
                "url": "https://www.rae.es/obras-academicas/ortografia/ortografia-2010"
            })
            recomendaciones["ejercicios"].append(
                "Practica la lectura en voz alta respetando las pausas marcadas por la puntuación")
            recomendaciones["siguiente_paso"] = "Revisa específicamente la puntuación en cada texto antes de finalizarlo"

        elif area_mejora == "Estructura":
            recomendaciones["recursos"].append({
                "titulo": "Coherencia y cohesión textual",
                "url": "https://cvc.cervantes.es/ensenanza/biblioteca_ele/diccio_ele/diccionario/coherencia.htm"
            })
            recomendaciones["ejercicios"].append(
                "Organiza tus ideas en un esquema antes de escribir")
            recomendaciones["siguiente_paso"] = "Trabaja la estructura de párrafos: idea principal y secundarias"
        else:
            # Recomendación general si no hay área específica
            recomendaciones["recursos"].append({
                "titulo": "Centro Virtual Cervantes - Recursos generales",
                "url": "https://cvc.cervantes.es/ensenanza/default.htm"
            })
            recomendaciones["ejercicios"].append(
                "Práctica variada de diferentes tipos de texto")
            recomendaciones["siguiente_paso"] = "Establece una rutina de escritura regular para mejorar tu fluidez"

        # Si hay pocas correcciones, añadir recomendación para practicar más
        if num_correcciones < 5:
            recomendaciones["mensaje_principal"] += " Para obtener recomendaciones más precisas, continúa practicando y corrigiendo tus textos regularmente."

        return recomendaciones

    except Exception as e:
        logger.error(
            f"Error al generar recomendaciones personalizadas: {str(e)}")
        return {
            "mensaje_principal": "Continúa practicando regularmente para mejorar tu español.",
            "recursos": [],
            "ejercicios": ["Práctica variada de diferentes tipos de texto"],
            "siguiente_paso": "Establece una rutina de escritura regular"
        }


# --- 4. UTILIDADES DE ANÁLISIS LINGÜÍSTICO ---

def calcular_indice_szigriszt(texto):
    """
    Calcula el índice de legibilidad Szigriszt-Pazos para español.

    Args:
        texto: Texto a analizar

    Returns:
        float: Índice de legibilidad Szigriszt
    """
    try:
        # Limpiar el texto
        texto_limpio = re.sub(r'\s+', ' ', texto).strip()

        # Contar sílabas (aproximación para español)
        def contar_silabas(palabra):
            palabra = palabra.lower()
            if not palabra:
                return 0

            # Contar vocales
            count = len(re.findall(r'[aeiouáéíóúü]', palabra))

            # Restar diptongos (aproximación simple)
            diptongos = len(re.findall(r'[aeiouáéíóúü][aeiouáéíóúü]', palabra))
            count -= diptongos

            # Casos especiales
            if palabra.endswith('e'):
                count -= 0.5

            return max(1, round(count))

        # Contar palabras
        palabras = re.findall(r'\b\w+\b', texto_limpio)
        num_palabras = len(palabras)

        # Contar frases
        frases = re.split(r'[.!?]+', texto_limpio)
        frases = [f for f in frases if f.strip()]
        num_frases = len(frases)

        # Contar sílabas totales
        num_silabas = sum(contar_silabas(palabra) for palabra in palabras)

        # Calcular longitud media de frase en palabras
        if num_frases == 0:
            return 0

        palabras_por_frase = num_palabras / num_frases

        # Calcular longitud media de palabra en sílabas
        if num_palabras == 0:
            return 0

        silabas_por_palabra = num_silabas / num_palabras

        # Fórmula Szigriszt-Pazos
        indice = 206.835 - (62.3 * silabas_por_palabra) - (palabras_por_frase)

        return round(indice, 2)

    except Exception as e:
        logger.error(f"Error al calcular índice Szigriszt: {str(e)}")
        return 0


def interpretar_szigriszt(indice):
    """
    Interpreta el índice de legibilidad Szigriszt.

    Args:
        indice: Valor del índice Szigriszt

    Returns:
        str: Interpretación del índice
    """
    if indice > 80:
        return "Muy fácil de leer"
    elif indice > 65:
        return "Bastante fácil de leer"
    elif indice > 50:
        return "Dificultad media"
    elif indice > 35:
        return "Bastante difícil de leer"
    else:
        return "Muy difícil de leer"

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 7: Vista de Perfil de Estudiante
# ==================================================================================
#
# Este artefacto contiene:
# 1. Función principal para mostrar el perfil y progreso del estudiante
# 2. Componentes para visualizar estadísticas y gráficos de progreso
# 3. Sección de recomendaciones personalizadas
# 4. Histórico de correcciones realizadas
#
# Esta nueva funcionalidad permite al estudiante visualizar su progreso,
# identificar áreas de mejora y recibir recomendaciones personalizadas.
# ==================================================================================


def view_perfil_estudiante():
    """
    Muestra el perfil y progreso del estudiante.
    Función principal para la vista de perfil de usuario.
    """
    st.header("👤 Perfil y Progreso del Estudiante")

    # Obtener información del usuario
    usuario = get_session_var("usuario_actual", "")
    email = get_session_var("email_usuario", "")
    nivel = get_session_var("nivel_estudiante", "intermedio")

    if not usuario:
        st.warning("Por favor, inicia sesión para ver tu perfil y progreso.")
        return

    # Mostrar información básica
    col1, col2 = st.columns([1, 2])

    with col1:
        # Imagen de perfil o avatar predeterminado
        st.image(
            "https://www.gravatar.com/avatar/00000000000000000000000000000000?d=mp&f=y", width=150)

    with col2:
        st.markdown(f"## {usuario}")
        st.markdown(
            f"**Correo electrónico:** {email if email else 'No disponible'}")

        # Mostrar nivel con etiqueta visual
        nivel_display = nivel.capitalize()
        if nivel == "principiante":
            st.markdown(
                f"**Nivel:** <span style='background-color:#a8e6cf;padding:3px 8px;border-radius:10px'>{nivel_display} (A1-A2)</span>", unsafe_allow_html=True)
        elif nivel == "intermedio":
            st.markdown(
                f"**Nivel:** <span style='background-color:#ffd3b6;padding:3px 8px;border-radius:10px'>{nivel_display} (B1-B2)</span>", unsafe_allow_html=True)
        else:
            st.markdown(
                f"**Nivel:** <span style='background-color:#ffaaa5;padding:3px 8px;border-radius:10px'>{nivel_display} (C1-C2)</span>", unsafe_allow_html=True)

    # Obtener historial del estudiante
    uid = get_session_var("uid_usuario", "")
    historial_df = obtener_historial_estudiante(uid, email, usuario)

    if historial_df is None or historial_df.empty:
        st.info(
            "Aún no tienes historial de correcciones. Realiza algunas correcciones para ver tu progreso.")

        # Mostrar sugerencias para comenzar
        st.markdown("### 🚀 Comienza tu aprendizaje")
        st.markdown("""
        Para comenzar a ver tu progreso, prueba estas actividades:
        
        1. **Corrige tu primer texto** en la sección "Corrección de texto"
        2. **Practica con consignas** usando el generador de consignas
        3. **Realiza un simulacro de examen** para evaluar tu nivel
        
        Una vez que hayas completado algunas actividades, podrás ver tu progreso aquí.
        """)
        return

    # Mostrar estadísticas generales
    st.subheader("📊 Estadísticas generales")

    # Calcular estadísticas de progreso
    estadisticas = calcular_estadisticas_progreso(historial_df)

    # Crear columnas para mostrar las métricas
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Textos corregidos", estadisticas["num_correcciones"])

    with col2:
        st.metric("Media de errores",
                  f"{estadisticas['media_errores_totales']:.1f}")

    with col3:
        # Calcular puntuación global promedio
        if 'Puntuación Global' in historial_df.columns:
            puntuacion_global = historial_df['Puntuación Global'].mean()
        else:
            # Calcular de otros campos si está disponible
            campos_puntuacion = [
                c for c in historial_df.columns if 'Puntuación' in c and c != 'Puntuación Global']
            if campos_puntuacion:
                puntuacion_global = historial_df[campos_puntuacion].mean(
                ).mean()
            else:
                puntuacion_global = 0

        st.metric("Puntuación media", f"{puntuacion_global:.1f}/10")

    with col4:
        # Calcular tendencia (mejora o empeora)
        if estadisticas["tendencia_errores"] == "mejora":
            st.metric("Tendencia", "Mejorando ↗️")
        elif estadisticas["tendencia_errores"] == "empeora":
            st.metric("Tendencia", "Empeorando ↘️")
        else:
            st.metric("Tendencia", "Estable →")

    # Crear pestañas para diferentes visualizaciones
    tab1, tab2, tab3 = st.tabs(
        ["📈 Evolución", "🎯 Áreas de mejora", "📋 Historial"])

    with tab1:
        st.subheader("Evolución de tu aprendizaje")

        # Crear gráfico de evolución de errores
        grafico_errores = crear_grafico_evolucion_errores(historial_df)
        if grafico_errores:
            st.altair_chart(grafico_errores, use_container_width=True)
        else:
            st.info("No hay suficientes datos para mostrar la evolución de errores.")

        # Crear gráfico de progreso en puntuaciones
        grafico_puntuaciones = crear_grafico_progreso_puntuaciones(
            historial_df)
        if grafico_puntuaciones:
            st.altair_chart(grafico_puntuaciones, use_container_width=True)
        else:
            st.info(
                "No hay suficientes datos para mostrar el progreso en puntuaciones.")

        # Mostrar fechas relevantes si están disponibles
        if estadisticas["primera_correccion"] and estadisticas["ultima_correccion"]:
            primera = estadisticas["primera_correccion"].strftime("%d/%m/%Y")
            ultima = estadisticas["ultima_correccion"].strftime("%d/%m/%Y")

            st.markdown(
                f"**Primera corrección:** {primera} | **Última corrección:** {ultima}")

            # Calcular días entre primera y última corrección
            dias_diferencia = (
                estadisticas["ultima_correccion"] - estadisticas["primera_correccion"]).days
            if dias_diferencia > 0:
                st.markdown(
                    f"Has estado practicando durante **{dias_diferencia} días**. ¡Sigue así!")

    with tab2:
        st.subheader("Áreas que necesitan mejora")

        # Crear gráfico de distribución de errores
        grafico_distribucion = crear_grafico_distribucion_errores(historial_df)
        if grafico_distribucion:
            st.altair_chart(grafico_distribucion, use_container_width=True)
        else:
            st.info(
                "No hay suficientes datos para identificar áreas de mejora específicas.")

        # Obtener área principal de mejora
        area_mejora = estadisticas.get("area_mejora_principal")

        if area_mejora:
            st.markdown(f"### Enfócate en mejorar: {area_mejora}")

            # Generar recomendaciones personalizadas
            recomendaciones = generar_recomendaciones_personalizadas(
                estadisticas, nivel)

            # Mostrar mensaje principal
            st.info(recomendaciones["mensaje_principal"])

            # Mostrar recursos recomendados
            if recomendaciones["recursos"]:
                st.markdown("### Recursos recomendados")
                for recurso in recomendaciones["recursos"]:
                    st.markdown(f"- [{recurso['titulo']}]({recurso['url']})")

            # Mostrar ejercicios sugeridos
            if recomendaciones["ejercicios"]:
                st.markdown("### Ejercicios sugeridos")
                for ejercicio in recomendaciones["ejercicios"]:
                    st.markdown(f"- {ejercicio}")

            # Mostrar siguiente paso
            if recomendaciones["siguiente_paso"]:
                st.markdown(f"### Siguiente paso recomendado")
                st.success(recomendaciones["siguiente_paso"])
        else:
            st.info(
                "Continúa practicando para identificar áreas específicas de mejora.")

    with tab3:
        st.subheader("Historial de textos corregidos")

        # Mostrar tabla con historial de correcciones
        if historial_df is not None and not historial_df.empty:
            # Seleccionar columnas relevantes
            columnas_mostrar = ['Fecha', 'Total Errores',
                                'Puntuación Global', 'Modelo']

            # Añadir más columnas si están disponibles
            for col in ['Puntuación Coherencia', 'Puntuación Registro']:
                if col in historial_df.columns:
                    columnas_mostrar.append(col)

            # Filtrar columnas disponibles
            columnas_disponibles = [
                col for col in columnas_mostrar if col in historial_df.columns]

            if columnas_disponibles:
                # Ordenar por fecha descendente
                historial_mostrar = historial_df[columnas_disponibles].sort_values(
                    'Fecha', ascending=False)

                # Formatear la fecha para mostrar
                if 'Fecha' in historial_mostrar.columns:
                    historial_mostrar['Fecha'] = pd.to_datetime(
                        historial_mostrar['Fecha']).dt.strftime('%d/%m/%Y %H:%M')

                # Mostrar tabla con opción de expandir filas
                st.dataframe(historial_mostrar, use_container_width=True)

                # Botón para descargar historial completo
                col1, col2 = st.columns([4, 1])
                with col2:
                    csv = historial_df.to_csv(index=False).encode('utf-8')

                    st.download_button(
                        label="📥 Descargar CSV",
                        data=csv,
                        file_name=f"historial_{usuario.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv"
                    )
            else:
                st.info("No hay datos suficientes en el historial.")
        else:
            st.info("No tienes correcciones guardadas en tu historial.")

    # Sección de motivación y consejo personalizado
    st.subheader("🔮 Consejo personalizado")

    # Generar consejo basado en las estadísticas
    if estadisticas["num_correcciones"] > 0:
        area_mejora = estadisticas.get("area_mejora_principal", "general")

        if estadisticas["tendencia_errores"] == "mejora":
            st.success(f"""
            ¡Felicidades! Estás progresando muy bien en tu aprendizaje del español. 
            Continúa practicando especialmente en el área de **{area_mejora}**, que es donde 
            todavía puedes mejorar más. Recuerda que la constancia es clave en el aprendizaje de idiomas.
            """)
        elif estadisticas["tendencia_errores"] == "empeora":
            st.warning(f"""
            Has tenido algunos desafíos recientemente, especialmente en el área de **{area_mejora}**. 
            Te recomiendo concentrarte en ejercicios específicos para esta área y quizás reducir 
            la complejidad de tus textos temporalmente. ¡Ánimo! Los retrocesos son parte normal 
            del proceso de aprendizaje.
            """)
        else:
            st.info(f"""
            Tu progreso se mantiene estable. Para seguir avanzando, te recomiendo explorar nuevos 
            tipos de texto y vocabulario, prestando especial atención al área de **{area_mejora}**. 
            Establecer pequeñas metas semanales puede ayudarte a mantener la motivación.
            """)
    else:
        st.info("""
        Bienvenido a tu perfil de estudiante. Aquí podrás seguir tu progreso a medida que 
        realices correcciones de texto. ¡Comienza ahora a practicar tu español!
        """)

    st.markdown("---")
    st.info("ℹ️ Puedes usar la pestaña 'Corrección de texto' en la parte superior para crear una nueva corrección.")

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 8: Exportación de Informes
# ==================================================================================
#
# Este artefacto contiene:
# 1. Funciones para generar informes en formato DOCX (Word)
# 2. Funciones para generar informes en formato HTML
# 3. Funciones para generar informes en formato CSV
# 4. Integración del logo de Spanish FactorIA en los informes
#
# Estas funciones permiten exportar los resultados de las correcciones en
# diferentes formatos para su uso fuera de la aplicación.
# ==================================================================================

# Mejoras en las funciones de generación y descarga de archivos

# 1. Mejorar generar_informe_docx para garantizar que funcione correctamente


def generar_informe_docx(nombre, nivel, fecha, texto_original, texto_corregido, errores_obj, analisis_contextual, consejo_final):
    """
    Genera un informe de corrección en formato Word (DOCX).
    Incluye el logo de Spanish FactorIA.
    Versión mejorada con manejo robusto de errores y mejor formato.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        fecha: Fecha de la corrección
        texto_original: Texto original del estudiante
        texto_corregido: Texto con correcciones
        errores_obj: Objeto con errores detectados
        analisis_contextual: Objeto con análisis contextual
        consejo_final: Consejo final para el estudiante

    Returns:
        BytesIO: Buffer con el documento generado
    """
    try:
        # Asegurar valores válidos para todos los parámetros
        nombre = nombre or "Estudiante"
        nivel = nivel or "No especificado"
        fecha = fecha or datetime.now().strftime("%Y-%m-%d %H:%M")
        texto_original = texto_original or "No disponible"
        texto_corregido = texto_corregido or "No disponible"
        errores_obj = errores_obj or {}
        analisis_contextual = analisis_contextual or {}
        consejo_final = consejo_final or "No disponible"

        # Importar Document solo cuando sea necesario
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
        except ImportError:
            logger.error(
                "No se pudo importar python-docx. Instalando la biblioteca...")
            # Intentar instalar la biblioteca (solo en entornos de desarrollo)
            try:
                import subprocess
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", "python-docx"])
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
            except Exception as e:
                logger.error(f"No se pudo instalar python-docx: {e}")
                raise ImportError(
                    "La biblioteca python-docx es necesaria para generar informes DOCX")

        # Crear el documento desde cero
        doc = Document()

        # Estilo del documento
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # Cabecera con logo
        header = doc.sections[0].header
        header_para = header.paragraphs[0]

        try:
            # Intenta añadir el logo si existe
            if os.path.exists(LOGO_PATH):
                run = header_para.add_run()
                run.add_picture(LOGO_PATH, width=Inches(1.5))

                # Añadir texto a la derecha del logo
                run = header_para.add_run(
                    "\tInforme generado por " + ORGANIZATION)
                run.font.size = Pt(9)
                run.font.italic = True
            else:
                # Si no hay logo, solo añadir texto
                run = header_para.add_run(
                    "Informe generado por " + ORGANIZATION)
                run.font.size = Pt(9)
                run.font.italic = True
                logger.warning(f"Logo no encontrado en: {LOGO_PATH}")
        except Exception as logo_error:
            logger.error(
                f"Error al añadir logo al documento: {str(logo_error)}")
            # Añadir solo texto si falla la inserción de la imagen
            run = header_para.add_run("Informe generado por " + ORGANIZATION)
            run.font.size = Pt(9)
            run.font.italic = True

        # Título
        doc.add_heading('Informe de corrección textual', 0)

        # Información general
        doc.add_heading('Información general', level=1)
        doc.add_paragraph(f'Nombre: {nombre}')
        doc.add_paragraph(f'Nivel: {nivel}')
        doc.add_paragraph(f'Fecha: {fecha}')

        # Texto original
        doc.add_heading('Texto original', level=1)
        doc.add_paragraph(texto_original)

        # Texto corregido
        doc.add_heading('Texto corregido', level=1)
        doc.add_paragraph(texto_corregido)

        # Análisis de errores
        doc.add_heading('Análisis de errores', level=1)

        # Verificación de errores con mejor manejo
        if isinstance(errores_obj, dict):
            # Ordenar categorías para una presentación consistente
            categorias_ordenadas = [
                "Gramática", "Léxico", "Puntuación", "Estructura textual"]
            categorias_encontradas = list(errores_obj.keys())

            # Asegurarse de incluir todas las categorías, incluso las que no estén en el orden predefinido
            for categoria in categorias_encontradas:
                if categoria not in categorias_ordenadas:
                    categorias_ordenadas.append(categoria)

            for categoria in categorias_ordenadas:
                if categoria in errores_obj:
                    errores = errores_obj[categoria]
                    if isinstance(errores, list) and errores:
                        doc.add_heading(categoria, level=2)
                        for i, err in enumerate(errores, 1):
                            if not isinstance(err, dict):
                                continue

                            p = doc.add_paragraph()
                            fragmento = err.get('fragmento_erroneo', '')
                            if fragmento:
                                run = p.add_run('Fragmento erróneo: ')
                                run.bold = True
                                run = p.add_run(fragmento)
                                run.font.color.rgb = RGBColor(255, 0, 0)

                            correccion = err.get('correccion', '')
                            if correccion:
                                p = doc.add_paragraph()
                                run = p.add_run('Corrección: ')
                                run.bold = True
                                run = p.add_run(correccion)
                                run.font.color.rgb = RGBColor(0, 128, 0)

                            explicacion = err.get('explicacion', '')
                            if explicacion:
                                p = doc.add_paragraph()
                                run = p.add_run('Explicación: ')
                                run.bold = True
                                p.add_run(explicacion)

                            doc.add_paragraph()  # Espacio
        else:
            doc.add_paragraph("No se detectaron errores significativos.")

        # Análisis contextual
        doc.add_heading('Análisis contextual', level=1)

        # Tabla de puntuaciones
        doc.add_heading('Puntuaciones', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aspecto'
        hdr_cells[1].text = 'Coherencia'
        hdr_cells[2].text = 'Cohesión'
        hdr_cells[3].text = 'Registro'
        hdr_cells[4].text = 'Adecuación cultural'

        # Verificación mejorada de análisis contextual
        # Datos
        row_cells = table.add_row().cells
        row_cells[0].text = 'Puntuación'

        # Manejar valores de manera más simple y robusta
        coherencia = {}
        cohesion = {}
        registro = {}
        adecuacion = {}

        if isinstance(analisis_contextual, dict):
            coherencia = analisis_contextual.get('coherencia', {}) or {}
            cohesion = analisis_contextual.get('cohesion', {}) or {}
            registro = analisis_contextual.get(
                'registro_linguistico', {}) or {}
            adecuacion = analisis_contextual.get(
                'adecuacion_cultural', {}) or {}

        row_cells[1].text = str(coherencia.get('puntuacion', 'N/A'))
        row_cells[2].text = str(cohesion.get('puntuacion', 'N/A'))
        row_cells[3].text = str(registro.get('puntuacion', 'N/A'))
        row_cells[4].text = str(adecuacion.get('puntuacion', 'N/A'))

        # Añadir comentarios del análisis contextual
        if coherencia:
            doc.add_heading('Coherencia textual', level=3)
            doc.add_paragraph(coherencia.get('comentario', 'No disponible'))

        if cohesion:
            doc.add_heading('Cohesión textual', level=3)
            doc.add_paragraph(cohesion.get('comentario', 'No disponible'))

        if registro:
            doc.add_heading('Registro lingüístico', level=3)
            doc.add_paragraph(
                f"Tipo detectado: {registro.get('tipo_detectado', 'No especificado')}")
            doc.add_paragraph(registro.get('adecuacion', 'No disponible'))

        if adecuacion:
            doc.add_heading('Adecuación cultural', level=3)
            doc.add_paragraph(adecuacion.get('comentario', 'No disponible'))

        # Consejo final
        doc.add_heading('Consejo final', level=1)
        doc.add_paragraph(consejo_final or "No disponible")

        # Manejo del QR mejorado para evitar errores
        try:
            # Verificar si la biblioteca qrcode está disponible
            import qrcode

            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )

            # Generar un ID único para el informe
            informe_id = f"{nombre.replace(' ', '')}_{fecha.replace(' ', '_').replace(':', '-')}"
            qr_data = f"spanishfactoria://informe/{informe_id}"
            qr.add_data(qr_data)
            qr.make(fit=True)

            img = qr.make_image(fill_color="black", back_color="white")

            # Guardar QR como imagen temporal
            qr_buffer = BytesIO()
            img.save(qr_buffer)
            qr_buffer.seek(0)

            # Añadir la imagen del QR al documento
            doc.add_heading('Acceso online', level=1)
            doc.add_paragraph(
                'Escanea este código QR para acceder a este informe online:')

            # Añadir la imagen
            doc.add_picture(qr_buffer, width=Inches(2.0))

            # Cerrar el buffer del QR
            qr_buffer.close()
        except (ImportError, Exception) as qr_error:
            # Si hay error con el QR, simplemente seguimos sin él
            doc.add_heading('Acceso online', level=1)
            doc.add_paragraph('Código QR no disponible en este momento.')
            logger.error(f"Error al generar QR: {str(qr_error)}")

        # Pie de página con información de la organización
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"{ORGANIZATION} - {APP_NAME} v{APP_VERSION}"
        footer_para.style = 'Footer'

        # Guardar el documento en un buffer y devolverlo
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)  # Importante: posicionar al inicio del buffer

        return docx_buffer

    except Exception as e:
        # Si hay un error general, hacemos un log detallado y devolvemos un documento de error
        logger.error(f"Error al generar informe DOCX: {str(e)}")
        logger.error(traceback.format_exc())

        try:
            # Crear un documento simple de error
            doc = Document()
            doc.add_heading('Error al generar informe', 0)
            doc.add_paragraph(
                f"Se produjo un error al generar el informe: {str(e)}")
            doc.add_paragraph(
                "Por favor, contacta con el soporte técnico si el problema persiste.")

            # Guardar en buffer
            error_buffer = BytesIO()
            doc.save(error_buffer)
            error_buffer.seek(0)
            return error_buffer
        except Exception:
            # Si falla incluso el documento de error, devolver None
            return None


# 2. Mejorar mostrar_opciones_exportacion para garantizar que se generen correctamente
def mostrar_opciones_exportacion(resultado_correccion):
    """
    Muestra las opciones de exportación para una corrección.
    Función para UI que maneja la generación y descarga de informes.
    Versión mejorada con mejor manejo de errores y feedback al usuario.

    Args:
        resultado_correccion: Diccionario con los resultados de la corrección
    """
    if not resultado_correccion or "error" in resultado_correccion:
        st.warning("No hay corrección disponible para exportar.")
        return

    # Inicializar opciones
    if "mostrar_exportacion" not in st.session_state:
        st.session_state.mostrar_exportacion = False

    # Botón para mostrar/ocultar opciones de exportación
    exportacion_btn = st.button(
        "📋 Opciones de exportación" if not st.session_state.mostrar_exportacion else "🔼 Ocultar opciones",
        key="toggle_exportacion"
    )

    if exportacion_btn:
        st.session_state.mostrar_exportacion = not st.session_state.mostrar_exportacion

    if not st.session_state.mostrar_exportacion:
        return

    # Crear un contenedor para las opciones de exportación
    with st.container():
        # Obtener información necesaria del resultado
        nombre = get_session_var("usuario_actual", "Estudiante")
        nivel = get_session_var("nivel_estudiante", "intermedio")
        fecha = datetime.now().strftime("%Y-%m-%d %H:%M")

        # Extraer datos con manejo seguro
        texto_original = get_session_var("texto_correccion_corregir", "")
        texto_corregido = resultado_correccion.get("texto_corregido", "")
        errores = resultado_correccion.get("errores", {})
        analisis_contextual = resultado_correccion.get(
            "analisis_contextual", {})
        consejo_final = resultado_correccion.get("consejo_final", "")

        # Mostrar interfaz de exportación
        st.markdown("### 📊 Exportar resultados")
        st.markdown(
            "Selecciona el formato en el que deseas exportar los resultados de esta corrección:")

        # Mostrar mensaje de información
        st.info("Los archivos se generarán y estarán disponibles para descarga. El proceso puede tardar unos segundos dependiendo del tamaño del texto.")

        col1, col2, col3 = st.columns(3)

        with col1:
            try:
                # Contenedor para estado de generación DOCX
                docx_status = st.empty()

                # Botón para generar y descargar Word
                if st.button("Generar Word (DOCX)", key="gen_docx"):
                    with docx_status.container():
                        with st.spinner("Generando documento Word..."):
                            # Generar informe Word (DOCX)
                            docx_buffer = generar_informe_docx(
                                nombre, nivel, fecha, texto_original, texto_corregido,
                                errores, analisis_contextual, consejo_final
                            )

                            if docx_buffer:
                                # Mostrar botón de descarga
                                st.download_button(
                                    label="📎 Descargar DOCX",
                                    data=docx_buffer,
                                    file_name=f"correccion_{nombre.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_docx"
                                )
                                st.success(
                                    "Documento Word generado correctamente")
                            else:
                                st.error(
                                    "No se pudo generar el documento Word")

            except Exception as e:
                st.error(f"Error al generar Word: {str(e)}")

        with col2:
            try:
                # Contenedor para estado de generación HTML
                html_status = st.empty()

                # Botón para generar y descargar HTML
                if st.button("Generar HTML", key="gen_html"):
                    with html_status.container():
                        with st.spinner("Generando documento HTML..."):
                            # Generar informe HTML
                            html_content = generar_informe_html(
                                nombre, nivel, fecha, texto_original, texto_corregido,
                                analisis_contextual, consejo_final
                            )

                            if html_content:
                                # Mostrar botón de descarga
                                st.download_button(
                                    label="🌐 Descargar HTML",
                                    data=html_content,
                                    file_name=f"correccion_{nombre.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                    mime="text/html",
                                    key="download_html"
                                )
                                st.success(
                                    "Documento HTML generado correctamente")
                            else:
                                st.error(
                                    "No se pudo generar el documento HTML")

            except Exception as e:
                st.error(f"Error al generar HTML: {str(e)}")

        with col3:
            try:
                # Contenedor para estado de generación CSV
                csv_status = st.empty()

                # Botón para generar y descargar CSV
                if st.button("Generar CSV", key="gen_csv"):
                    with csv_status.container():
                        with st.spinner("Generando archivo CSV..."):
                            # Generar CSV con datos de análisis
                            csv_buffer = generar_csv_analisis(
                                nombre, nivel, fecha, resultado_correccion
                            )

                            if csv_buffer:
                                # Mostrar botón de descarga
                                st.download_button(
                                    label="📊 Descargar CSV",
                                    data=csv_buffer,
                                    file_name=f"analisis_{nombre.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                                    mime="text/csv",
                                    key="download_csv"
                                )
                                st.success(
                                    "Archivo CSV generado correctamente")
                            else:
                                st.error("No se pudo generar el archivo CSV")

            except Exception as e:
                st.error(f"Error al generar CSV: {str(e)}")

        # Mostrar instrucciones adicionales
        st.markdown("---")
        st.markdown("""
        **Nota:**
        - El documento Word (.docx) contiene el informe completo con formato profesional.
        - El archivo HTML puede abrirse en cualquier navegador web.
        - El archivo CSV contiene datos para análisis estadístico o importación en hojas de cálculo.
        
        Puedes usar estos archivos para guardar tu progreso o compartirlos con tu profesor.
        """)

        # Opción para compartir por correo electrónico (implementación limitada)
        with st.expander("Compartir por correo electrónico", expanded=False):
            st.markdown("""
            Para compartir tu informe por correo electrónico:
            1. Descarga el archivo en el formato deseado
            2. Adjunta el archivo descargado a tu correo electrónico
            """)

            email_destino = st.text_input(
                "Correo electrónico de destino:", key="email_share")

            if st.button("Preparar correo", key="prepare_email"):
                # Esta es una implementación simplificada que solo muestra instrucciones
                # En una implementación real, se podría usar mailto: o un servicio de correo
                mailto_link = f"mailto:{email_destino}?subject=Informe%20de%20corrección%20-%20{nombre}&body=Adjunto%20encontrarás%20mi%20informe%20de%20corrección%20de%20texto%20en%20español."

                st.markdown(f"""
                <div style="background-color: #f0f7ff; padding: 15px; border-radius: 8px; margin-top: 10px;">
                <p>✅ Preparado para compartir con <b>{email_destino}</b></p>
                <p>Descarga el archivo y envíalo como adjunto a la dirección de correo.</p>
                <a href="{mailto_link}" target="_blank">Abrir cliente de correo</a> (funciona solo en equipos con cliente de correo configurado)
                </div>
                """, unsafe_allow_html=True)


def generar_informe_html(nombre, nivel, fecha, texto_original, texto_corregido, analisis_contextual, consejo_final):
    """
    Genera un informe de corrección en formato HTML.
    Versión mejorada con mejor sanitización y manejo de valores nulos.
    Incluye el logo de Spanish FactorIA.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        fecha: Fecha de la corrección
        texto_original: Texto original del estudiante
        texto_corregido: Texto con correcciones
        analisis_contextual: Objeto con análisis contextual
        consejo_final: Consejo final para el estudiante

    Returns:
        str: Contenido HTML del informe
    """
    try:
        # Usar valores seguros por defecto
        nombre = nombre or "Estudiante"
        nivel = nivel or "No especificado"
        fecha = fecha or datetime.now().strftime("%Y-%m-%d %H:%M")
        texto_original = texto_original or "No disponible"
        texto_corregido = texto_corregido or "No disponible"
        consejo_final = consejo_final or "No disponible"
        app_version = APP_VERSION

        # Preparar logo para HTML (codificado en base64)
        logo_base64 = None
        try:
            if os.path.exists(LOGO_PATH):
                with open(LOGO_PATH, "rb") as img_file:
                    logo_bytes = img_file.read()
                    logo_base64 = base64.b64encode(logo_bytes).decode("utf-8")
        except Exception as e:
            logger.error(f"Error al codificar logo en base64: {e}")

        # Función mejorada para sanitizar HTML
        def sanitize_html(text):
            if not text:
                return ""
            # Sanitizar caracteres especiales HTML
            sanitized = text.replace("&", "&amp;").replace(
                "<", "&lt;").replace(">", "&gt;")
            sanitized = sanitized.replace('"', "&quot;").replace("'", "&#39;")
            # Convertir saltos de línea en <br>
            sanitized = sanitized.replace("\n", "<br>")
            return sanitized

        # Sanitizar textos
        texto_original_safe = sanitize_html(texto_original)
        texto_corregido_safe = sanitize_html(texto_corregido)
        consejo_final_safe = sanitize_html(consejo_final)

        # Verificación robusta de análisis contextual
        if not isinstance(analisis_contextual, dict):
            analisis_contextual = {}

        # Extraer datos de análisis contextual con manejo seguro
        coherencia = analisis_contextual.get('coherencia', {}) or {}
        cohesion = analisis_contextual.get('cohesion', {}) or {}
        registro = analisis_contextual.get('registro_linguistico', {}) or {}
        adecuacion = analisis_contextual.get('adecuacion_cultural', {}) or {}

        # Obtener puntuaciones con valores por defecto
        puntuacion_coherencia = coherencia.get('puntuacion', 'N/A')
        puntuacion_cohesion = cohesion.get('puntuacion', 'N/A')
        puntuacion_registro = registro.get('puntuacion', 'N/A')
        puntuacion_adecuacion = adecuacion.get('puntuacion', 'N/A')

        # Crear HTML con estructura mejorada y logo de Spanish FactorIA
        html_content = f'''
        <!DOCTYPE html>
        <html lang="es">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Informe de corrección - {sanitize_html(nombre)}</title>
            <style>
                body {{ 
                    font-family: Arial, sans-serif; 
                    line-height: 1.6; 
                    margin: 0;
                    padding: 20px;
                    color: #333;
                }}
                .container {{ 
                    max-width: 800px; 
                    margin: 0 auto; 
                    padding: 20px; 
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                    background-color: #fff;
                    border-radius: 5px;
                }}
                .header {{
                    display: flex;
                    align-items: center;
                    margin-bottom: 20px;
                    border-bottom: 1px solid #eee;
                    padding-bottom: 10px;
                }}
                .logo {{
                    max-width: 150px;
                    margin-right: 20px;
                }}
                .header-text {{
                    flex-grow: 1;
                }}
                h1 {{ 
                    color: #2c3e50; 
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{ 
                    color: #3498db; 
                    margin-top: 30px; 
                    border-left: 4px solid #3498db;
                    padding-left: 10px;
                }}
                h3 {{ 
                    color: #2980b9; 
                    margin-top: 20px;
                }}
                .original {{ 
                    background-color: #f8f9fa; 
                    padding: 15px; 
                    border-left: 4px solid #6c757d; 
                    white-space: pre-wrap;
                    margin: 15px 0;
                    border-radius: 4px;
                }}
                .corregido {{ 
                    background-color: #e7f4e4; 
                    padding: 15px; 
                    border-left: 4px solid #28a745; 
                    white-space: pre-wrap;
                    margin: 15px 0;
                    border-radius: 4px;
                }}
                .puntuaciones {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 20px 0;
                }}
                .puntuaciones th, .puntuaciones td {{ 
                    border: 1px solid #ddd; 
                    padding: 10px; 
                    text-align: center;
                }}
                .puntuaciones th {{ 
                    background-color: #f2f2f2; 
                    font-weight: bold;
                }}
                .consejo {{ 
                    background-color: #e7f5fe; 
                    padding: 15px; 
                    border-left: 4px solid #17a2b8; 
                    margin: 20px 0;
                    border-radius: 4px;
                }}
                .footer {{ 
                    margin-top: 50px; 
                    padding-top: 20px; 
                    border-top: 1px solid #ddd; 
                    color: #6c757d; 
                    font-size: 0.8em;
                    text-align: center;
                }}
                @media print {{
                    body {{ background-color: #fff; }}
                    .container {{ box-shadow: none; }}
                    a {{ text-decoration: none; color: #000; }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    {'<img class="logo" src="data:image/png;base64,'+logo_base64+'" alt="Spanish FactorIA Logo">' if logo_base64 else ''}
                    <div class="header-text">
                        <h1>Informe de corrección textual</h1>
                        <p>Powered by {ORGANIZATION}</p>
                    </div>
                </div>

                <section>
                    <h2>Información general</h2>
                    <p><strong>Nombre:</strong> {sanitize_html(nombre)}</p>
                    <p><strong>Nivel:</strong> {sanitize_html(nivel)}</p>
                    <p><strong>Fecha:</strong> {sanitize_html(fecha)}</p>
                </section>

                <section>
                    <h2>Texto original</h2>
                    <div class="original">
                        {texto_original_safe}
                    </div>

                    <h2>Texto corregido</h2>
                    <div class="corregido">
                        {texto_corregido_safe}
                    </div>
                </section>

                <section>
                    <h2>Análisis contextual</h2>

                    <h3>Puntuaciones</h3>
                    <table class="puntuaciones">
                        <tr>
                            <th>Coherencia</th>
                            <th>Cohesión</th>
                            <th>Registro</th>
                            <th>Adecuación cultural</th>
                        </tr>
                        <tr>
                            <td>{puntuacion_coherencia}/10</td>
                            <td>{puntuacion_cohesion}/10</td>
                            <td>{puntuacion_registro}/10</td>
                            <td>{puntuacion_adecuacion}/10</td>
                        </tr>
                    </table>
                    
                    <h3>Coherencia textual</h3>
                    <p>{sanitize_html(coherencia.get('comentario', 'No disponible'))}</p>
                    
                    <h3>Cohesión textual</h3>
                    <p>{sanitize_html(cohesion.get('comentario', 'No disponible'))}</p>
                    
                    <h3>Registro lingüístico</h3>
                    <p><strong>Tipo detectado:</strong> {sanitize_html(registro.get('tipo_detectado', 'No especificado'))}</p>
                    <p>{sanitize_html(registro.get('adecuacion', 'No disponible'))}</p>
                    
                    <h3>Adecuación cultural</h3>
                    <p>{sanitize_html(adecuacion.get('comentario', 'No disponible'))}</p>
                </section>

                <section>
                    <h2>Consejo final</h2>
                    <div class="consejo">
                        <p>{consejo_final_safe}</p>
                    </div>
                </section>

                <div class="footer">
                    <p>{ORGANIZATION} - {APP_NAME} v{app_version}</p>
                    <p>Informe generado el {sanitize_html(fecha)}</p>
                </div>
            </div>
        </body>
        </html>
        '''

        return html_content

    except Exception as e:
        logger.error(f"Error al generar informe HTML: {str(e)}")
        logger.error(traceback.format_exc())

        # Crear HTML básico de error
        return f'''
        <!DOCTYPE html>
        <html>
        <head><title>Error en informe</title></head>
        <body>
            <h1>Error al generar informe</h1>
            <p>Se produjo un error al generar el informe. Por favor, inténtelo de nuevo.</p>
            <p>Error: {str(e).replace("<", "&lt;").replace(">", "&gt;")}</p>
            <p>Powered by {ORGANIZATION}</p>
        </body>
        </html>
        '''


def generar_csv_analisis(nombre, nivel, fecha, datos_analisis):
    """
    Genera un archivo CSV con los datos de análisis de una corrección.
    Versión mejorada con mejor manejo de buffers y validación de datos.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        fecha: Fecha de la corrección
        datos_analisis: Diccionario con datos de análisis

    Returns:
        BytesIO: Buffer con el CSV generado
    """
    try:
        # Asegurar que tenemos datos válidos
        nombre = nombre or "Estudiante"
        nivel = nivel or "No especificado"
        fecha = fecha or datetime.now().strftime("%Y-%m-%d %H:%M")

        # Verificar que datos_analisis es un diccionario
        if not isinstance(datos_analisis, dict):
            datos_analisis = {}

        # Extraer datos con manejo seguro
        errores = datos_analisis.get("errores", {}) or {}
        analisis_contextual = datos_analisis.get(
            "analisis_contextual", {}) or {}

        # Contar errores por categoría con validación robusta
        num_gramatica = len(errores.get("Gramática", [])) if isinstance(
            errores.get("Gramática"), list) else 0
        num_lexico = len(errores.get("Léxico", [])) if isinstance(
            errores.get("Léxico"), list) else 0
        num_puntuacion = len(errores.get("Puntuación", [])) if isinstance(
            errores.get("Puntuación"), list) else 0
        num_estructura = len(errores.get("Estructura textual", [])) if isinstance(
            errores.get("Estructura textual"), list) else 0
        total_errores = num_gramatica + num_lexico + num_puntuacion + num_estructura

        # Extraer puntuaciones con validación robusta
        coherencia = analisis_contextual.get("coherencia", {}) or {}
        cohesion = analisis_contextual.get("cohesion", {}) or {}
        registro = analisis_contextual.get("registro_linguistico", {}) or {}
        adecuacion = analisis_contextual.get("adecuacion_cultural", {}) or {}

        # Convertir a valores numéricos o 0 si no están disponibles
        def safe_numeric(val, default=0):
            if val is None:
                return default
            try:
                return float(val)
            except (ValueError, TypeError):
                return default

        coherencia_punt = safe_numeric(coherencia.get("puntuacion"))
        cohesion_punt = safe_numeric(cohesion.get("puntuacion"))
        registro_punt = safe_numeric(registro.get("puntuacion"))
        adecuacion_punt = safe_numeric(adecuacion.get("puntuacion"))
        puntuacion_global = (coherencia_punt + cohesion_punt +
                             registro_punt + adecuacion_punt) / 4

        # Crear un buffer para el CSV
        csv_buffer = StringIO()
        writer = csv.writer(csv_buffer)

        # Escribir los encabezados
        headers = [
            "Fecha", "Nombre", "Nivel",
            "Errores_Gramatica", "Errores_Lexico", "Errores_Puntuacion", "Errores_Estructura", "Total_Errores",
            "Puntuacion_Coherencia", "Puntuacion_Cohesion", "Puntuacion_Registro", "Puntuacion_Adecuacion", "Puntuacion_Global"
        ]
        writer.writerow(headers)

        # Escribir los datos
        row = [
            fecha, nombre, nivel,
            num_gramatica, num_lexico, num_puntuacion, num_estructura, total_errores,
            f"{coherencia_punt:.1f}", f"{cohesion_punt:.1f}", f"{registro_punt:.1f}", f"{adecuacion_punt:.1f}", f"{puntuacion_global:.1f}"
        ]
        writer.writerow(row)

        # Mover el buffer al inicio y devolverlo
        csv_text = csv_buffer.getvalue()
        csv_buffer.close()

        # Convertir a BytesIO para la descarga
        bytes_buffer = BytesIO()
        bytes_buffer.write(csv_text.encode('utf-8'))
        bytes_buffer.seek(0)

        return bytes_buffer

    except Exception as e:
        logger.error(f"Error al generar CSV de análisis: {str(e)}")

        # En caso de error, crear un CSV básico con mensaje de error
        csv_buffer = StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(["Error", "Detalles"])
        writer.writerow(["Error al generar CSV", str(e)])

        csv_text = csv_buffer.getvalue()
        csv_buffer.close()

        bytes_buffer = BytesIO()
        bytes_buffer.write(csv_text.encode('utf-8'))
        bytes_buffer.seek(0)

        return bytes_buffer


def mostrar_opciones_exportacion(resultado_correccion):
    """
    Muestra las opciones de exportación para una corrección.
    Función para UI que maneja la generación y descarga de informes.

    Args:
        resultado_correccion: Diccionario con los resultados de la corrección
    """
    if not resultado_correccion or "error" in resultado_correccion:
        st.warning("No hay corrección disponible para exportar.")
        return

    # Inicializar opciones
    if "mostrar_exportacion" not in st.session_state:
        st.session_state.mostrar_exportacion = False

    if st.button("📋 Opciones de exportación"):
        st.session_state.mostrar_exportacion = not st.session_state.mostrar_exportacion

    if not st.session_state.mostrar_exportacion:
        return

    # Obtener información necesaria del resultado
    nombre = get_session_var("usuario_actual", "Estudiante")
    nivel = get_session_var("nivel_estudiante", "intermedio")
    fecha = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Extraer datos con manejo seguro
    texto_original = get_session_var("texto_correccion_corregir", "")
    texto_corregido = resultado_correccion.get("texto_corregido", "")
    errores = resultado_correccion.get("errores", {})
    analisis_contextual = resultado_correccion.get("analisis_contextual", {})
    consejo_final = resultado_correccion.get("consejo_final", "")

    # Mostrar interfaz de exportación
    st.markdown("### 📊 Exportar resultados")
    st.markdown(
        "Selecciona el formato en el que deseas exportar los resultados de esta corrección:")

    col1, col2, col3 = st.columns(3)

    with col1:
        try:
            # Generar informe Word (DOCX)
            docx_buffer = generar_informe_docx(
                nombre, nivel, fecha, texto_original, texto_corregido, errores, analisis_contextual, consejo_final)

            if docx_buffer:
                st.download_button(
                    label="📎 Descargar como Word (DOCX)",
                    data=docx_buffer,
                    file_name=f"correccion_{nombre.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No se pudo generar el documento Word")

        except Exception as e:
            st.error(f"Error al generar Word: {str(e)}")

    with col2:
        try:
            # Generar informe HTML
            html_content = generar_informe_html(
                nombre, nivel, fecha, texto_original, texto_corregido, analisis_contextual, consejo_final)

            if html_content:
                st.download_button(
                    label="🌐 Descargar como HTML",
                    data=html_content,
                    file_name=f"correccion_{nombre.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                    mime="text/html"
                )
            else:
                st.warning("No se pudo generar el HTML")

        except Exception as e:
            st.error(f"Error al generar HTML: {str(e)}")

    with col3:
        try:
            # Generar CSV con datos de análisis
            csv_buffer = generar_csv_analisis(
                nombre, nivel, fecha, resultado_correccion)

            if csv_buffer:
                st.download_button(
                    label="📊 Descargar como CSV",
                    data=csv_buffer,
                    file_name=f"analisis_{nombre.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv"
                )
            else:
                st.warning("No se pudo generar el CSV")

        except Exception as e:
            st.error(f"Error al generar CSV: {str(e)}")

    # Opción para compartir por correo electrónico
    st.markdown("### 📧 Compartir resultados")
    email_destino = st.text_input("Correo electrónico de destino:")

    if st.button("Enviar por correo") and email_destino:
        # Aquí normalmente implementaríamos el envío por correo,
        # pero como es solo una demostración, mostramos mensaje informativo
        st.success(f"Se enviaría el informe al correo: {email_destino}")
        st.info(
            "Nota: Esta función requiere configuración de servidor SMTP que no está incluida en esta versión.")

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 9: Interfaz de Usuario de Login
# ==================================================================================
#
# Este artefacto contiene:
# 1. Función para mostrar la pantalla de login/bienvenida
# 2. Funciones para manejar la autenticación de usuarios
# 3. Funciones UI para la cabecera y navegación principal
#
# Estas funciones gestionan la experiencia de usuario inicial y la navegación
# entre las diferentes secciones de la aplicación.
# ==================================================================================


def ui_login_screen():
    """
    Muestra la pantalla de bienvenida y login.
    Gestiona la autenticación simple de usuarios.

    Returns:
        bool: True si el usuario ha iniciado sesión correctamente
    """
    # Configurar layout para pantalla de bienvenida
    st.markdown(
        """
        <style>
        .welcome-container {
            text-align: center;
            padding: 2rem;
            max-width: 800px;
            margin: 0 auto;
        }
        .welcome-header {
            font-size: 2.5rem;
            color: #3498db;
            margin-bottom: 1rem;
        }
        .welcome-subheader {
            font-size: 1.5rem;
            color: #555;
            margin-bottom: 2rem;
        }
        .feature-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-top: 2rem;
            margin-bottom: 2rem;
        }
        .feature-card {
            border: 1px solid #eaeaea;
            border-radius: 10px;
            padding: 1rem;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
            transition: transform 0.3s;
        }
        .feature-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }
        .feature-icon {
            font-size: 2rem;
            margin-bottom: 0.5rem;
            color: #3498db;
        }
        .login-container {
            max-width: 400px;
            margin: 0 auto;
            padding: 2rem;
            border: 1px solid #eaeaea;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        }
        @media (max-width: 768px) {
            .feature-grid {
                grid-template-columns: 1fr;
            }
        }
        
        /* CSS para eliminar los recuadros blancos vacíos */
        div.element-container:has(div.stTextInput:empty) {
            display: none !important;
            height: 0 !important;
            min-height: 0 !important;
            visibility: hidden !important;
        }
        
        /* Ocultar contenedores específicos de la pantalla de login */
        #root > div:nth-child(1) > div.withScreencast > div > div > div > section > div > div:nth-child(1) > div > div:nth-child(10),
        #root > div:nth-child(1) > div.withScreencast > div > div > div > section > div > div:nth-child(1) > div > div:nth-child(16) {
            display: none !important;
        }
        
        /* Solución general para elementos vacíos */
        .login-container > div:empty,
        .login-container ~ div:empty,
        .stTextInput:has(input:placeholder-shown) + div:empty {
            display: none !important;
        }
        
        /* Ocultar los contenedores vacíos que se generan para los campos de entrada */
        .stTextInput div[data-baseweb="input-container"]:empty,
        .stTextInput div[data-baseweb="input"] div:empty {
            display: none !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Logo o imagen de bienvenida
    logo_col, _ = st.columns([1, 1])  # Para centrar el logo
    with logo_col:
        try:
            if os.path.exists(LOGO_PATH):
                logo_img = Image.open(LOGO_PATH)
                st.image(logo_img, width=300)
            else:
                st.title(f"{APP_NAME}")
                logger.warning(f"Logo no encontrado en: {LOGO_PATH}")
        except Exception as e:
            st.title(f"{APP_NAME}")
            logger.error(f"Error al cargar logo: {e}")

    # Información de bienvenida
    st.markdown(
        f"""
        <div class="welcome-container">
            <h1 class="welcome-header">¡Bienvenido a {APP_NAME}!</h1>
            <p class="welcome-subheader">
                Una herramienta avanzada para corregir textos en español y mejorar tus habilidades lingüísticas
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Funcionalidades principales
    st.markdown(
        """
        <div class="welcome-container">
            <h2>Nuestras funcionalidades</h2>
            <div class="feature-grid">
                <div class="feature-card">
                    <div class="feature-icon">✏️</div>
                    <h3>Corrección de textos</h3>
                    <p>Corrección detallada con análisis gramatical, léxico y contextual</p>
                </div>
                <div class="feature-card">
                    <div class="feature-icon">📊</div>
                    <h3>Seguimiento de progreso</h3>
                    <p>Visualiza tu evolución y áreas de mejora con estadísticas detalladas</p>
                </div>
                <div class="feature-card">
                    <div class="feature-icon">🔍</div>
                    <h3>Análisis contextual</h3>
                    <p>Evaluación de coherencia, cohesión y adecuación cultural</p>
                </div>
                <div class="feature-card">
                    <div class="feature-icon">🎓</div>
                    <h3>Feedback personalizado</h3>
                    <p>Consejos adaptados a tu nivel para mejorar tu español</p>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Sección de acceso
    st.markdown(
        """
        <div class="welcome-container">
            <h2>Accede a tu cuenta</h2>
        </div>
        """,
        unsafe_allow_html=True
    )

    # CSS adicional para eliminar específicamente los recuadros marcados
    st.markdown("""
    <style>
    /* Solución específica para eliminar los recuadros blancos marcados */
    form div.row-widget.stTextInput:has(input:placeholder-shown):empty {
        visibility: hidden !important;
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }

    /* Selectores específicos para los contenedores vacíos */
    .element-container:has(.login-container + div:empty),
    .login-container + div:empty {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # Formulario de login o registro simple
    login_col, registro_col = st.columns(2)

    # Manejo de login existente
    with login_col:
        st.markdown("<div class='login-container'>", unsafe_allow_html=True)
        st.subheader("Iniciar sesión")

        login_email = st.text_input("Correo electrónico", key="login_email")
        login_nombre = st.text_input("Nombre", key="login_nombre")

        # Selector de nivel
        nivel_options = [
            "Nivel principiante (A1-A2)",
            "Nivel intermedio (B1-B2)",
            "Nivel avanzado (C1-C2)"
        ]
        login_nivel = st.selectbox(
            "Nivel de español",
            options=nivel_options,
            index=1,  # Default: intermedio
            key="login_nivel"
        )

        # Mapeo simplificado de nivel
        nivel_simple = "intermedio"  # Default
        if login_nivel:
            if "principiante" in login_nivel.lower():
                nivel_simple = "principiante"
            elif "intermedio" in login_nivel.lower():
                nivel_simple = "intermedio"
            elif "avanzado" in login_nivel.lower():
                nivel_simple = "avanzado"

        if st.button("Iniciar sesión", type="primary"):
            if login_email and login_nombre:
                # Generar UID basado en el email
                uid = generate_user_uid(login_email)

                # Almacenar datos del usuario en session_state
                set_session_var("usuario_actual", login_nombre)
                set_session_var("email_usuario", login_email)
                set_session_var("uid_usuario", uid)
                set_session_var("nivel_estudiante", nivel_simple)
                set_session_var("is_authenticated", True)
                set_session_var("fecha_inicio_sesion",
                                datetime.now().isoformat())
                set_session_var("mostrar_login", False)

                # Registrar evento de login
                logger.info(
                    f"Usuario {login_nombre} ({login_email}) ha iniciado sesión")

                # Indicar éxito de login
                return True
            else:
                st.error("Por favor, introduce tu correo electrónico y nombre")

        st.markdown("</div>", unsafe_allow_html=True)

    # Información de registro
    with registro_col:
        st.markdown("<div class='login-container'>", unsafe_allow_html=True)
        st.subheader("¿Eres nuevo?")
        st.markdown(
            """
            Simplemente introduce tu correo electrónico y nombre para comenzar.
            
            Con tu cuenta podrás:
            - Guardar tu historial de correcciones
            - Realizar un seguimiento de tu progreso
            - Recibir recomendaciones personalizadas
            - Exportar y compartir tus resultados
            """
        )
        st.markdown("</div>", unsafe_allow_html=True)

    # Opciones para acceso sin registro (modo invitado)
    st.markdown("<div class='welcome-container'>", unsafe_allow_html=True)
    st.markdown("---")

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.write("¿Prefieres probar sin registrarte?")
        if st.button("Continuar como invitado"):
            # Generar ID único para invitado
            guest_id = str(uuid.uuid4())
            guest_name = f"Invitado_{guest_id[:6]}"

            # Configurar sesión como invitado
            set_session_var("usuario_actual", guest_name)
            set_session_var("email_usuario", "")
            set_session_var("uid_usuario", f"guest_{guest_id}")
            set_session_var("nivel_estudiante", "intermedio")
            set_session_var("is_authenticated", True)
            set_session_var("mostrar_login", False)

            # Indicar éxito
            return True

    st.markdown("</div>", unsafe_allow_html=True)

    # Pie de página
    st.markdown(
        f"""
        <div style="text-align: center; margin-top: 50px; color: #666; font-size: 0.8rem;">
            <p>{ORGANIZATION} - Versión {APP_VERSION}</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Si llegamos aquí, no se ha completado el login
    return False


def ui_header():
    """
    Muestra la cabecera principal de la aplicación.
    Incluye información del usuario y controles de navegación.
    Versión actualizada sin indicador de estado de servicios.
    """
    # Información del usuario actual
    usuario = get_session_var("usuario_actual", "")
    nivel = get_session_var("nivel_estudiante", "intermedio")

    # Ajustamos la proporción de las columnas sin la tercera columna
    col1, col2 = st.columns([3, 7])

    with col1:
        st.markdown(f"### 👋 Hola, {usuario}")

        # Mostrar nivel con etiqueta visual
        nivel_display = nivel.capitalize()
        if nivel == "principiante":
            st.markdown(
                f"<span style='background-color:#a8e6cf;padding:3px 8px;border-radius:10px'>{nivel_display} (A1-A2)</span>", unsafe_allow_html=True)
        elif nivel == "intermedio":
            st.markdown(
                f"<span style='background-color:#ffd3b6;padding:3px 8px;border-radius:10px'>{nivel_display} (B1-B2)</span>", unsafe_allow_html=True)
        else:
            st.markdown(
                f"<span style='background-color:#ffaaa5;padding:3px 8px;border-radius:10px'>{nivel_display} (C1-C2)</span>", unsafe_allow_html=True)

    with col2:
        # Título centrado
        st.markdown(
            f"<h1 style='text-align: center;'>{APP_NAME}</h1>", unsafe_allow_html=True)

    # Separador para distinguir header del contenido
    st.markdown("---")


def view_correccion_texto():
    """
    Muestra la interfaz para corrección de textos.
    Esta es la función principal de interfaz para la corrección.
    """
    st.header("✏️ Corrección de Texto")

    # Comprobar si hay texto para corregir de otras herramientas
    texto_para_corregir = get_session_var("texto_correccion_corregir", "")
    info_adicional = get_session_var("info_adicional_corregir", "")
    mostrar_correccion_imagen = get_session_var(
        "mostrar_correccion_imagen", False)
    mostrar_correccion_transcripcion = get_session_var(
        "mostrar_correccion_transcripcion", False)

    # Verificar si tenemos textos pendientes para corrección desde otras herramientas
    if mostrar_correccion_imagen or mostrar_correccion_transcripcion:
        st.info(f"Texto recibido para corrección: {info_adicional}")

        # Mostrar configuración de la corrección
        col1, col2, col3 = st.columns(3)

        with col1:
            idioma_opciones = ["Español", "Inglés", "Francés"]
            idioma = st.selectbox(
                "Idioma de corrección",
                options=idioma_opciones,
                index=0,
                help="Idioma en el que quieres recibir las explicaciones",
                key="idioma_texto_recibido"
            )

        with col2:
            tipo_texto_opciones = [
                "General/No especificado",
                "Académico",
                "Profesional",
                "Formal",
                "Informal",
                "Creativo/Literario",
                "Periodístico",
                "Email/Carta"
            ]
            # Nuevo bloque con multiselect
            tipo_texto_seleccionados = st.multiselect(
                "Tipo de texto (selección múltiple)",
                options=tipo_texto_opciones,
                default=["General/No especificado"],
                help="Ayuda al sistema a entender el contexto de tu texto. Puedes seleccionar varios tipos.",
                key="tipo_texto_recibido_multi"
            )
            # Convertir la selección múltiple a string para mantener compatibilidad
            tipo_texto = ", ".join(
                tipo_texto_seleccionados) if tipo_texto_seleccionados else "General/No especificado"

        with col3:
            contexto_cultural_opciones = [
                "General/Internacional",
                "España",
                "México",
                "Argentina",
                "Colombia",
                "Otro país hispano"
            ]
            # Nuevo bloque con multiselect
            contexto_cultural_seleccionados = st.multiselect(
                "Contexto cultural (selección múltiple)",
                options=contexto_cultural_opciones,
                default=["General/Internacional"],
                help="Contexto cultural relevante para tu texto. Puedes seleccionar varios contextos.",
                key="contexto_cultural_recibido_multi"
            )
            # Convertir la selección múltiple a string para mantener compatibilidad
            contexto_cultural = ", ".join(
                contexto_cultural_seleccionados) if contexto_cultural_seleccionados else "General/Internacional"

        # Asegurar que el texto se muestra en el área de texto para corrección
        texto_correccion = st.text_area(
            "Texto a corregir:",
            value=texto_para_corregir,  # Usar el texto recibido
            height=200,
            key="texto_corrector_recibido"
        )

        # Actualizar la variable si el usuario modifica el texto
        if texto_correccion != texto_para_corregir:
            set_session_var("texto_correccion_corregir", texto_correccion)

        # Botón para corregir
        if st.button("Corregir texto", key="btn_corregir_recibido", type="primary"):
            if not texto_correccion:
                st.error("Por favor, introduce un texto para corregir")
            else:
                with st.spinner("Analizando texto..."):
                    try:
                        # Obtener el nombre y nivel del usuario
                        nombre = get_session_var(
                            "usuario_actual", "Estudiante")
                        nivel = get_session_var(
                            "nivel_estudiante", "intermedio")

                        # Llamar a la función de corrección
                        resultado = corregir_texto(
                            texto_correccion,
                            nombre,
                            nivel,
                            idioma,
                            tipo_texto,
                            contexto_cultural,
                            info_adicional
                        )

                        # Almacenar el resultado en session_state
                        set_session_var("correction_result", resultado)
                        set_session_var("last_correction_time",
                                        datetime.now().isoformat())
                        set_session_var("last_correction_data", {
                            "texto": texto_correccion,
                            "nombre": nombre,
                            "nivel": nivel,
                            "idioma": idioma,
                            "tipo_texto": tipo_texto,
                            "contexto_cultural": contexto_cultural
                        })

                        # Limpiar flags para evitar mostrar este texto nuevamente
                        set_session_var("mostrar_correccion_imagen", False)
                        set_session_var(
                            "mostrar_correccion_transcripcion", False)

                        # Forzar refresco para mostrar resultados
                        st.rerun()
                    except Exception as e:
                        handle_exception("view_correccion_texto", e)
                        st.error(
                            f"Se produjo un error durante la corrección: {str(e)}")
    else:
        # Mostrar instrucciones básicas
        st.markdown("""
        Introduce tu texto en español para recibir una corrección detallada con análisis contextual.
        """)

        # Configuración de la corrección
        col1, col2, col3 = st.columns(3)

        with col1:
            idioma_opciones = ["Español", "Inglés", "Francés"]
            idioma = st.selectbox(
                "Idioma de corrección",
                options=idioma_opciones,
                index=0,
                help="Idioma en el que quieres recibir las explicaciones"
            )

        with col2:
            tipo_texto_opciones = [
                "General/No especificado",
                "Académico",
                "Profesional",
                "Formal",
                "Informal",
                "Creativo/Literario",
                "Periodístico",
                "Email/Carta"
            ]
            # Reemplazar selectbox por multiselect
            tipo_texto_seleccionados = st.multiselect(
                "Tipo de texto (selección múltiple)",
                options=tipo_texto_opciones,
                default=["General/No especificado"],
                help="Ayuda al sistema a entender el contexto de tu texto. Puedes seleccionar varios tipos."
            )
            # Convertir la selección múltiple a string para mantener compatibilidad
            tipo_texto = ", ".join(
                tipo_texto_seleccionados) if tipo_texto_seleccionados else "General/No especificado"

        with col3:
            contexto_cultural_opciones = [
                "General/Internacional",
                "España",
                "México",
                "Argentina",
                "Colombia",
                "Otro país hispano"
            ]
            # Reemplazar selectbox por multiselect
            contexto_cultural_seleccionados = st.multiselect(
                "Contexto cultural (selección múltiple)",
                options=contexto_cultural_opciones,
                default=["General/Internacional"],
                help="Contexto cultural relevante para tu texto. Puedes seleccionar varios contextos."
            )
            # Convertir la selección múltiple a string para mantener compatibilidad
            contexto_cultural = ", ".join(
                contexto_cultural_seleccionados) if contexto_cultural_seleccionados else "General/Internacional"

        # Campo para el texto a corregir
        texto_correccion = st.text_area(
            "Texto a corregir",
            value=get_session_var("texto_correccion_corregir", ""),
            height=200,
            help="Escribe o pega aquí el texto que quieres que corrijamos"
        )

        # Guardar el texto en session_state para persistencia
        set_session_var("texto_correccion_corregir", texto_correccion)

        # Campo para información adicional
        info_adicional = st.text_area(
            "Información adicional (opcional)",
            value=get_session_var("info_adicional_corregir", ""),
            height=100,
            help="Añade cualquier información relevante o contexto específico"
        )

        # Guardar info adicional en session_state
        set_session_var("info_adicional_corregir", info_adicional)

        # Botón para corregir
        col1, col2 = st.columns([5, 1])

        with col2:
            corregir_clicked = st.button(
                "Corregir texto", type="primary", use_container_width=True)

        # Mostrar spinner y procesar la corrección
        if corregir_clicked:
            if not texto_correccion:
                st.error("Por favor, introduce un texto para corregir")
            else:
                with st.spinner("Analizando texto..."):
                    try:
                        # Obtener el nombre y nivel del usuario
                        nombre = get_session_var(
                            "usuario_actual", "Estudiante")
                        nivel = get_session_var(
                            "nivel_estudiante", "intermedio")

                        # Llamar a la función de corrección
                        resultado = corregir_texto(
                            texto_correccion,
                            nombre,
                            nivel,
                            idioma,
                            tipo_texto,
                            contexto_cultural,
                            info_adicional
                        )

                        # Almacenar el resultado en session_state
                        set_session_var("correction_result", resultado)
                        set_session_var("last_correction_time",
                                        datetime.now().isoformat())
                        set_session_var("last_correction_data", {
                            "texto": texto_correccion,
                            "nombre": nombre,
                            "nivel": nivel,
                            "idioma": idioma,
                            "tipo_texto": tipo_texto,
                            "contexto_cultural": contexto_cultural
                        })
                    except Exception as e:
                        handle_exception("view_correccion_texto", e)
                        st.error(
                            f"Se produjo un error durante la corrección: {str(e)}")

    # Mostrar resultados si existen
    resultado = get_session_var("correction_result", None)

    if resultado:
        # Verificar si hay error
        if "error" in resultado:
            st.error(f"Error en la corrección: {resultado['error']}")
            return

        st.markdown("---")
        st.markdown("## 📝 Resultados de la corrección")

        # Verificar y mostrar si se usó el asistente
        if "via_assistant" in resultado and resultado["via_assistant"] == True:
            st.success(
                "✅ Esta corrección fue generada por un **Asistente de OpenAI personalizado**")
        else:
            st.warning(
                "⚠️ Esta corrección fue procesada usando la API estándar de OpenAI")

        # Verificar que el resultado tenga estructura esperada
        if not all(k in resultado for k in ["saludo", "texto_corregido", "consejo_final"]):
            st.error(
                "El formato de los resultados es incorrecto. Inténtalo de nuevo.")
            return

        # Mostrar saludo personalizado
        st.markdown(f"### {resultado['saludo']}")

        # Mostrar texto corregido y original en pestañas
        tab1, tab2 = st.tabs(["Texto corregido", "Texto original"])

        with tab1:
            st.markdown(resultado["texto_corregido"])

        with tab2:
            texto_original = get_session_var(
                "last_correction_data", {}).get("texto", texto_correccion)
            st.markdown(texto_original)

        # Mostrar análisis de errores
        st.markdown("### 🔍 Análisis de errores")

        # Expandir por defecto si hay errores
        hay_errores = False

        # Verificar y mostrar errores por categoría
        errores = resultado.get("errores", {})
        if isinstance(errores, dict):
            for categoria, lista_errores in errores.items():
                if lista_errores:
                    hay_errores = True
                    break

        with st.expander("Ver errores específicos", expanded=hay_errores):
            errores_encontrados = False

            # Usar pestañas para categorías de errores
            if errores and isinstance(errores, dict):
                # Determinar categorías con errores
                tabs = []
                for categoria, lista_errores in errores.items():
                    if lista_errores:
                        tabs.append(categoria)

                if tabs:
                    # Crear pestañas dinámicamente
                    error_tabs = st.tabs(tabs)

                    # Mostrar errores en cada pestaña
                    for i, categoria in enumerate(tabs):
                        with error_tabs[i]:
                            lista_errores = errores[categoria]
                            for idx, error in enumerate(lista_errores):
                                col1, col2 = st.columns(2)

                                with col1:
                                    st.markdown(f"**Fragmento erróneo:**")
                                    st.markdown(
                                        f"<span style='color:red'>{error.get('fragmento_erroneo', '')}</span>", unsafe_allow_html=True)

                                with col2:
                                    st.markdown(f"**Corrección:**")
                                    st.markdown(
                                        f"<span style='color:green'>{error.get('correccion', '')}</span>", unsafe_allow_html=True)

                                st.markdown(
                                    f"**Explicación:** {error.get('explicacion', '')}")

                                if idx < len(lista_errores) - 1:
                                    st.markdown("---")

                            errores_encontrados = True

            if not errores_encontrados:
                st.success(
                    "¡Felicidades! No se encontraron errores significativos en tu texto.")

        # Mostrar análisis contextual
        st.markdown("### 📊 Análisis contextual")

        analisis = resultado.get("analisis_contextual", {})
        if isinstance(analisis, dict):
            # Extraer puntuaciones para visualización
            coherencia = analisis.get("coherencia", {})
            cohesion = analisis.get("cohesion", {})
            registro = analisis.get("registro_linguistico", {})
            adecuacion = analisis.get("adecuacion_cultural", {})

            # Verificar y extraer puntuaciones
            puntuaciones = []
            categorias = []

            if isinstance(coherencia, dict) and "puntuacion" in coherencia:
                puntuaciones.append(coherencia["puntuacion"])
                categorias.append("Coherencia")

            if isinstance(cohesion, dict) and "puntuacion" in cohesion:
                puntuaciones.append(cohesion["puntuacion"])
                categorias.append("Cohesión")

            if isinstance(registro, dict) and "puntuacion" in registro:
                puntuaciones.append(registro["puntuacion"])
                categorias.append("Registro")

            if isinstance(adecuacion, dict) and "puntuacion" in adecuacion:
                puntuaciones.append(adecuacion["puntuacion"])
                categorias.append("Adecuación")

            # Crear gráfico radar si hay puntuaciones
            if puntuaciones and categorias:
                try:
                    # Visualización en columnas: gráfico y detalles
                    col1, col2 = st.columns([2, 3])

                    with col1:
                        # Generar y mostrar gráfico radar
                        fig = crear_grafico_radar(puntuaciones, categorias)
                        st.pyplot(fig)

                    with col2:
                        # Mostrar valoraciones numéricas
                        valores_cols = st.columns(len(categorias))

                        for i, (cat, punt) in enumerate(zip(categorias, puntuaciones)):
                            with valores_cols[i]:
                                st.metric(cat, f"{punt}/10")

                        # Calcular puntuación media
                        puntuacion_media = sum(
                            puntuaciones) / len(puntuaciones)
                        st.metric("Valoración global",
                                  f"{puntuacion_media:.1f}/10")

                except Exception as grafico_error:
                    logger.error(
                        f"Error al crear gráfico radar: {str(grafico_error)}")
                    st.warning("No se pudo generar el gráfico de análisis.")

            # Mostrar comentarios detallados
            with st.expander("Comentarios detallados", expanded=True):
                # Coherencia
                if coherencia:
                    st.markdown("#### Coherencia")
                    st.markdown(coherencia.get("comentario", "No disponible"))

                    # Sugerencias
                    sugerencias = coherencia.get("sugerencias", [])
                    if sugerencias:
                        st.markdown("**Sugerencias:**")
                        for sug in sugerencias:
                            st.markdown(f"- {sug}")

                # Cohesión
                if cohesion:
                    st.markdown("#### Cohesión")
                    st.markdown(cohesion.get("comentario", "No disponible"))

                    # Sugerencias
                    sugerencias = cohesion.get("sugerencias", [])
                    if sugerencias:
                        st.markdown("**Sugerencias:**")
                        for sug in sugerencias:
                            st.markdown(f"- {sug}")

                # Registro lingüístico
                if registro:
                    st.markdown("#### Registro lingüístico")
                    st.markdown(
                        f"**Tipo detectado:** {registro.get('tipo_detectado', 'No especificado')}")
                    st.markdown(registro.get("adecuacion", "No disponible"))

                    # Sugerencias
                    sugerencias = registro.get("sugerencias", [])
                    if sugerencias:
                        st.markdown("**Sugerencias:**")
                        for sug in sugerencias:
                            st.markdown(f"- {sug}")

                # Adecuación cultural
                if adecuacion:
                    st.markdown("#### Adecuación cultural")
                    st.markdown(adecuacion.get("comentario", "No disponible"))

                    # Elementos destacables
                    elementos = adecuacion.get("elementos_destacables", [])
                    if elementos:
                        st.markdown("**Elementos destacables:**")
                        for elem in elementos:
                            st.markdown(f"- {elem}")

                    # Sugerencias
                    sugerencias = adecuacion.get("sugerencias", [])
                    if sugerencias:
                        st.markdown("**Sugerencias:**")
                        for sug in sugerencias:
                            st.markdown(f"- {sug}")

        # Consejo final
        st.markdown("### 💡 Consejo final")
        st.success(resultado.get("consejo_final", ""))

        # Generar audio del consejo si ElevenLabs está disponible
        try:
            audio_bytes = None

            # Verificar disponibilidad de API
            if api_keys["elevenlabs"]["api_key"] and api_keys["elevenlabs"]["voice_id"]:
                if circuit_breaker.can_execute("elevenlabs"):
                    # Generar audio
                    audio_bytes = generar_audio_consejo(
                        resultado.get("consejo_final", ""))

            # Mostrar reproductor de audio si se generó correctamente
            if audio_bytes:
                st.audio(audio_bytes, format="audio/mp3")

                # Opción para descargar
                st.download_button(
                    label="⬇️ Descargar audio",
                    data=audio_bytes,
                    file_name=f"consejo_{datetime.now().strftime('%Y%m%d_%H%M')}.mp3",
                    mime="audio/mp3"
                )
        except Exception as audio_error:
            logger.error(f"Error al generar audio: {str(audio_error)}")
            # No mostrar mensaje al usuario para evitar confusión

        # Opciones de exportación
        mostrar_opciones_exportacion(resultado)

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 10: Función Principal y Punto de Entrada
# ==================================================================================
#
# Este artefacto contiene:
# 1. Función main() como punto de entrada de la aplicación
# 2. Integración de todas las vistas y componentes
# 3. Manejo de navegación entre diferentes secciones
# 4. Diagnóstico e inicialización de la aplicación
#
# Esta función es el punto central donde todas las partes de la aplicación
# se integran y se gestiona el flujo de navegación del usuario.
# ==================================================================================


def view_acerca_de():
    """Muestra información acerca de la aplicación con formato mejorado."""
    st.header("ℹ️ Acerca de")

    col1, col2 = st.columns([1, 2])

    with col1:
        # Logo en grande
        try:
            if os.path.exists(LOGO_PATH):
                logo_img = Image.open(LOGO_PATH)
                st.image(logo_img, width=200)
            else:
                # Alternativa online
                st.image(
                    "https://raw.githubusercontent.com/spanishfactoria/assets/main/Spanish_FactorIA_Logo.png", width=200)
        except Exception as e:
            logger.error(f"Error al cargar logo: {e}")
            st.write("Spanish FactorIA")

    with col2:
        st.markdown(f"""
        # {APP_NAME}
        
        **Versión:** {APP_VERSION}
        
        Una herramienta avanzada para la corrección de textos en español, diseñada específicamente para estudiantes de español como lengua extranjera (ELE).
        """)

    # Sección para Estado de Servicios con estilo mejorado
    st.markdown("""
    <style>
    .servicios-header {
        background-color: #f8f9fa;
        padding: 10px 15px;
        border-radius: 8px;
        border-left: 4px solid #3498db;
        margin: 15px 0;
        font-weight: bold;
    }
    </style>
    <div class="servicios-header">🖥️ Estado de los Servicios del Sistema</div>
    """, unsafe_allow_html=True)

    with st.expander("Ver estado detallado de los servicios", expanded=False):
        # Mostrar estado de conexiones
        status = circuit_breaker.get_status()

        col1, col2 = st.columns(2)

        with col1:
            # OpenAI
            if api_keys["openai"] is not None:
                if not status["openai"]["open"]:
                    st.success("✅ OpenAI: Conectado")
                    # Mostrar el modelo actual
                    modelo_actual, _ = configure_openai()
                    if modelo_actual:
                        st.info(f"Modelo actual: {modelo_actual}")
                else:
                    error_types = status["openai"]["error_types"]
                    error_msg = ", ".join(
                        [f"{tipo}: {count}" for tipo, count in error_types.items()])
                    st.error(
                        f"❌ OpenAI: Desconectado ({status['openai']['failures']} fallos - {error_msg})")
                    # Añadir botón para reintentar
                    if st.button("Reintentar OpenAI"):
                        if circuit_breaker.attempt_reset("openai"):
                            st.info("Reintentando conexión con OpenAI...")
                            configure_openai()
                            st.rerun()
            else:
                st.warning("⚠️ OpenAI: No configurado")

            # Firebase
            if api_keys["firebase_credentials"] is not None:
                if not status["firebase"]["open"]:
                    st.success("✅ Firebase: Conectado")
                else:
                    error_types = status["firebase"]["error_types"]
                    error_msg = ", ".join(
                        [f"{tipo}: {count}" for tipo, count in error_types.items()])
                    st.error(
                        f"❌ Firebase: Desconectado ({status['firebase']['failures']} fallos - {error_msg})")
            else:
                st.warning("⚠️ Firebase: No configurado")

        with col2:
            # ElevenLabs
            if api_keys["elevenlabs"]["api_key"] is not None:
                if not status["elevenlabs"]["open"]:
                    st.success("✅ ElevenLabs: Conectado")
                else:
                    error_types = status["elevenlabs"]["error_types"]
                    error_msg = ", ".join(
                        [f"{tipo}: {count}" for tipo, count in error_types.items()])
                    st.error(
                        f"❌ ElevenLabs: Desconectado ({status['elevenlabs']['failures']} fallos - {error_msg})")
            else:
                st.warning("⚠️ ElevenLabs: No configurado")

            # DALL-E
            if api_keys["dalle"] is not None:
                if not status["dalle"]["open"]:
                    st.success("✅ DALL-E: Conectado")
                else:
                    error_types = status["dalle"]["error_types"]
                    error_msg = ", ".join(
                        [f"{tipo}: {count}" for tipo, count in error_types.items()])
                    st.error(
                        f"❌ DALL-E: Desconectado ({status['dalle']['failures']} fallos - {error_msg})")
            else:
                st.warning("⚠️ DALL-E: No configurado")

    # Características principales
    st.markdown("""
    ### Características principales
    
    - **Corrección detallada** con análisis gramatical, léxico y contextual
    - **Análisis contextual** de coherencia, cohesión, registro lingüístico y adecuación cultural
    - **Seguimiento de progreso** con estadísticas y visualización de mejora
    - **Recomendaciones personalizadas** adaptadas al nivel del estudiante
    - **Exportación de informes** en diferentes formatos (Word, HTML, CSV)
    
    ### Tecnologías utilizadas
    
    - **Streamlit**: Para la interfaz de usuario
    - **OpenAI**: Para análisis lingüístico y procesamiento de textos
    - **Firebase**: Para almacenamiento de datos del usuario
    - **ElevenLabs**: Para síntesis de voz de alta calidad
    
    ### Desarrollado por
    
    Spanish FactorIA - Herramientas avanzadas para la enseñanza y aprendizaje del español
    """)

    # Mostrar opciones de colaboración y feedback en un diseño más atractivo
    st.markdown("---")
    st.subheader("📣 Tu opinión es importante")

    # Diseño de dos columnas para el feedback
    feedback_col1, feedback_col2 = st.columns([3, 2])

    with feedback_col1:
        with st.form(key="feedback_form"):
            st.write(
                "Comparte tu opinión para ayudarnos a mejorar esta herramienta:")

            rating = st.slider("¿Cómo valorarías esta aplicación?",
                               min_value=1, max_value=5, value=4)
            feedback_text = st.text_area(
                "¿Tienes algún comentario o sugerencia?", height=100)

            submit_button = st.form_submit_button(
                label="Enviar feedback", type="primary")

            if submit_button:
                try:
                    # Guardar feedback en Firebase si está disponible
                    if get_session_var("firebase_available", False):
                        result = guardar_feedback_firebase(
                            rating, feedback_text)

                        if result["success"]:
                            st.success(
                                "¡Gracias por tu feedback! Lo hemos recibido correctamente.")
                        else:
                            st.warning(
                                "No pudimos guardar tu feedback. Gracias de todos modos.")
                    else:
                        st.info(
                            "Gracias por tu feedback. El guardado de datos no está disponible en este momento.")
                except Exception as e:
                    st.error(f"Error al procesar feedback: {str(e)}")

    with feedback_col2:
        st.markdown("""
        ### Contacto
        
        ¿Necesitas ayuda o tienes una consulta específica? Contáctanos a través de:
        
        📧 **Email**: soporte@spanishfactoria.com
        
        🌐 **Web**: [www.spanishfactoria.com](https://www.spanishfactoria.com)
        
        🔗 **LinkedIn**: [Spanish FactorIA](https://linkedin.com/company/spanishfactoria)
        """)

    # Sección final con detalles técnicos
    with st.expander("Detalles técnicos", expanded=False):
        st.markdown("""
        #### Información técnica
        
        - **Plataforma**: Streamlit Cloud
        - **Tipo de aplicación**: Web
        - **Backend**: Python 3.8+
        - **Almacenamiento**: Firebase Firestore
        - **Modelos IA**: OpenAI GPT 3.5/4
        
        #### Recursos y limitaciones
        
        Esta aplicación utiliza APIs externas con ciertas cuotas y limitaciones. El rendimiento puede variar según la disponibilidad de estos servicios.
        
        #### Privacidad
        
        Los datos almacenados se utilizan únicamente para mejorar la experiencia del usuario y el funcionamiento de la aplicación. No se comparten con terceros.
        """)


def view_herramientas_ele():
    """Muestra herramientas adicionales para ELE."""
    st.header("🛠️ Herramientas ELE")

    # Lista de herramientas disponibles
    herramientas = [
        "Generador de imágenes para descripción",
        "Transcripción de textos manuscritos",
        "Generador de consignas"
    ]

    # Seleccionar herramienta
    herramienta_seleccionada = st.selectbox(
        "Selecciona una herramienta:",
        options=herramientas,
        index=get_session_var("active_tools_tab", 0)
    )

    # Guardar selección en session_state
    set_session_var("active_tools_tab", herramientas.index(
        herramienta_seleccionada))

    # Mostrar la herramienta seleccionada
    st.markdown("---")

    if herramienta_seleccionada == "Generador de imágenes para descripción":
        mostrar_herramienta_imagenes()
    elif herramienta_seleccionada == "Transcripción de textos manuscritos":
        mostrar_herramienta_transcripcion()
    elif herramienta_seleccionada == "Generador de consignas":
        mostrar_herramienta_consignas()


def mostrar_herramienta_imagenes():
    """Muestra la herramienta de generación de imágenes para descripción con ejemplos por nivel."""
    st.subheader("🖼️ Generador de imágenes para descripción")

    st.markdown("""
    Esta herramienta genera imágenes adaptadas a tu nivel de español, 
    junto con una descripción y preguntas para practicar.
    """)

    # Obtener nivel actual del usuario
    nivel_estudiante = get_session_var("nivel_estudiante", "intermedio")

    # Opciones para generación
    col1, col2 = st.columns(2)

    with col1:
        # Tema para la imagen con ejemplos predefinidos por nivel
        tema_ejemplos = {
            "principiante": [
                "una familia en el parque",
                "mi rutina diaria",
                "un mercado local",
                "mi casa y habitación",
                "las cuatro estaciones"
            ],
            "intermedio": [
                "un festival cultural",
                "un viaje en tren",
                "un encuentro entre amigos",
                "un día en la ciudad",
                "un paisaje natural"
            ],
            "avanzado": [
                "una manifestación cultural",
                "un debate sobre medio ambiente",
                "la brecha tecnológica",
                "las diferencias socioeconómicas",
                "la preservación de tradiciones"
            ]
        }

        # Obtener ejemplos para el nivel actual
        ejemplos = tema_ejemplos.get(
            nivel_estudiante, tema_ejemplos["intermedio"])
        ejemplos_texto = ", ".join([f'"{ej}"' for ej in ejemplos])

        # Campo de texto con ejemplos sugeridos
        tema = st.text_input(
            "Tema para la imagen:",
            value=get_session_var("tema_imagen_state", ""),
            help=f"Ejemplos para tu nivel: {ejemplos_texto}"
        )

        # Guardar tema en session_state
        set_session_var("tema_imagen_state", tema)

        # Mostrar ejemplos como botones para facilitar la selección
        st.markdown("#### Ejemplos de temas para tu nivel:")

        # Crear filas de botones para los ejemplos
        ejemplo_cols = st.columns(len(ejemplos))
        for i, ejemplo in enumerate(ejemplos):
            with ejemplo_cols[i]:
                if st.button(ejemplo, key=f"btn_ejemplo_{i}"):
                    # Actualizar el tema y el valor del input
                    set_session_var("tema_imagen_state", ejemplo)
                    # Usar rerun en lugar de experimental_rerun
                    st.rerun()

    with col2:
        # Nivel del estudiante
        nivel_options = [
            "Nivel principiante (A1-A2)",
            "Nivel intermedio (B1-B2)",
            "Nivel avanzado (C1-C2)"
        ]

        nivel_index = 0
        if nivel_estudiante == "principiante":
            nivel_index = 0
        elif nivel_estudiante == "intermedio":
            nivel_index = 1
        elif nivel_estudiante == "avanzado":
            nivel_index = 2

        nivel = st.selectbox(
            "Nivel de español:",
            options=nivel_options,
            index=nivel_index
        )

    # Botón para generar imagen
    generar_clicked = st.button("Generar imagen", type="primary")

    # Verificar si debemos generar una imagen
    if generar_clicked:
        if not tema:
            st.error("Por favor, introduce un tema para la imagen")
        else:
            with st.spinner("Generando imagen..."):
                try:
                    # Generar imagen
                    imagen_url, descripcion = generar_imagen_dalle(tema, nivel)

                    if imagen_url:
                        # Guardar resultados en session_state
                        set_session_var("imagen_url_state", imagen_url)
                        set_session_var("descripcion_state", descripcion)
                        set_session_var("imagen_generada_state", True)
                    else:
                        st.error(
                            "No se pudo generar la imagen. Por favor, intenta de nuevo.")
                except Exception as e:
                    handle_exception("mostrar_herramienta_imagenes", e)
                    st.error(f"Error al generar imagen: {str(e)}")

    # Mostrar imagen generada si existe
    imagen_url = get_session_var("imagen_url_state", None)
    descripcion = get_session_var("descripcion_state", None)
    imagen_generada = get_session_var("imagen_generada_state", False)

    if imagen_generada and imagen_url and descripcion:
        st.markdown("---")
        st.subheader("Imagen generada:")

        # Mostrar imagen
        st.image(imagen_url, caption=f"Imagen sobre: {tema}")

        # Mostrar descripción y preguntas
        st.markdown("### Descripción y preguntas")
        st.markdown(descripcion)

        # Campo para que el estudiante escriba su descripción
        descripcion_estudiante = st.text_area(
            "Escribe tu propia descripción de la imagen:",
            value=get_session_var("descripcion_estudiante_state", ""),
            height=150,
            key="descripcion_estudiante_input"
        )

        # Guardar descripción en session_state inmediatamente
        if descripcion_estudiante != get_session_var("descripcion_estudiante_state", ""):
            set_session_var("descripcion_estudiante_state",
                            descripcion_estudiante)

        # Opción para corregir la descripción del estudiante
        if st.button("Corregir mi descripción"):
            if not descripcion_estudiante:
                st.warning("Por favor, escribe tu descripción primero")
            else:
                # Extraer nivel simple
                nivel_simple = "intermedio"
                if "principiante" in nivel.lower():
                    nivel_simple = "principiante"
                elif "avanzado" in nivel.lower():
                    nivel_simple = "avanzado"

                # Redirigir a la función de corrección principal
                set_session_var("texto_correccion_corregir",
                                descripcion_estudiante)
                set_session_var("info_adicional_corregir",
                                f"Descripción de una imagen sobre: {tema}")

                # Marcar para mostrar corrección
                set_session_var("mostrar_correccion_imagen", True)

                # Cambiar a la pestaña de corrección
                set_session_var("active_tab", 0)

                # Notificar al usuario sobre la redirección
                st.success(
                    "✅ Tu descripción ha sido enviada a la sección 'Corrección de texto'")
                st.info(
                    "Por favor, navega a la pestaña 'Corrector de texto' para ver los resultados")

                # Forzar el refresco para aplicar los cambios
                st.rerun()


def mostrar_herramienta_transcripcion():
    """Muestra la herramienta de transcripción de textos manuscritos con notificación clara."""
    st.subheader("📝 Transcripción de textos manuscritos")

    st.markdown("""
    Esta herramienta te permite subir una imagen de un texto manuscrito
    para obtener su transcripción digital, que luego podrás corregir.
    """)

    # Subir imagen
    uploaded_file = st.file_uploader(
        "Sube una imagen con texto manuscrito",
        type=["png", "jpg", "jpeg"],
        help="Formatos soportados: PNG, JPG, JPEG"
    )

    # Selector de idioma
    idioma_opciones = ["Español", "Inglés", "Francés"]
    idioma_map = {"Español": "es", "Inglés": "en", "Francés": "fr"}

    idioma = st.selectbox(
        "Idioma del texto manuscrito",
        options=idioma_opciones,
        index=0
    )

    # Obtener código de idioma
    idioma_codigo = idioma_map.get(idioma, "es")

    # Procesar imagen si se ha subido
    if uploaded_file is not None:
        # Mostrar imagen subida
        st.image(uploaded_file, caption="Imagen subida",
                 use_container_width=True)

        # Botón para transcribir
        if st.button("Transcribir texto", type="primary"):
            with st.spinner("Transcribiendo texto..."):
                try:
                    # Leer bytes de la imagen
                    imagen_bytes = uploaded_file.getvalue()

                    # Transcribir texto
                    texto_transcrito = transcribir_imagen_texto(
                        imagen_bytes, idioma_codigo)

                    if texto_transcrito and "Error" not in texto_transcrito:
                        # Guardar resultado en session_state
                        set_session_var(
                            "ultimo_texto_transcrito", texto_transcrito)
                        st.success("¡Transcripción completada!")
                    else:
                        st.error(
                            f"Error en la transcripción: {texto_transcrito}")
                except Exception as e:
                    handle_exception("mostrar_herramienta_transcripcion", e)
                    st.error(f"Error al transcribir: {str(e)}")

    # Mostrar resultado de transcripción si existe
    texto_transcrito = get_session_var("ultimo_texto_transcrito", "")

    if texto_transcrito:
        st.markdown("---")
        st.subheader("Texto transcrito:")

        # Mostrar transcripción
        st.markdown("### Resultado:")
        st.text_area(
            "Texto transcrito (puedes editarlo si es necesario):",
            value=texto_transcrito,
            height=200,
            key="texto_transcrito_editado"
        )

        # Opciones para corregir el texto transcrito
        if st.button("Corregir este texto"):
            # Obtener texto editado
            texto_final = st.session_state.texto_transcrito_editado

            # Preparar para corrección
            set_session_var("texto_correccion_corregir", texto_final)
            set_session_var("info_adicional_corregir",
                            "Texto transcrito de imagen manuscrita")

            # Marcar para mostrar corrección
            set_session_var("mostrar_correccion_transcripcion", True)

            # Cambiar a pestaña de corrección
            set_session_var("active_tab", 0)

            # Notificar al usuario sobre la redirección
            st.success(
                "✅ El texto transcrito ha sido enviado a la sección 'Corrección de texto'")
            st.info(
                "Por favor, navega a la pestaña 'Corrector de texto' para ver los resultados")


def mostrar_herramienta_consignas():
    """Muestra la herramienta de generación de consignas para escritura."""
    st.subheader("📋 Generador de consignas")

    st.markdown("""
    Esta herramienta genera consignas de escritura adaptadas a tu nivel,
    para que practiques la expresión escrita en español.
    """)

    # Opciones para generación de consigna
    col1, col2 = st.columns(2)

    with col1:
        # Nivel del estudiante
        nivel_options = [
            "Nivel principiante (A1-A2)",
            "Nivel intermedio (B1-B2)",
            "Nivel avanzado (C1-C2)"
        ]

        nivel_index = 0
        nivel_estudiante = get_session_var("nivel_estudiante", "intermedio")

        if nivel_estudiante == "principiante":
            nivel_index = 0
        elif nivel_estudiante == "intermedio":
            nivel_index = 1
        elif nivel_estudiante == "avanzado":
            nivel_index = 2

        nivel = st.selectbox(
            "Nivel de español:",
            options=nivel_options,
            index=nivel_index
        )

    with col2:
        # Tipo de texto
        tipo_options = [
            "General",
            "Narrativo",
            "Descriptivo",
            "Argumentativo",
            "Email/Carta",
            "Diálogo",
            "Resumen",
            "Opinión"
        ]

        tipo_texto = st.selectbox(
            "Tipo de texto:",
            options=tipo_options,
            index=0
        )

    # Botón para generar consigna
    generar_clicked = st.button("Generar consigna", type="primary")

    # Verificar si debemos generar una consigna
    if generar_clicked:
        with st.spinner("Generando consigna..."):
            try:
                # Extraer nivel simple
                nivel_simple = "intermedio"
                if "principiante" in nivel.lower():
                    nivel_simple = "principiante"
                elif "avanzado" in nivel.lower():
                    nivel_simple = "avanzado"

                # Generar consigna (usando OpenAI)
                system_msg = """
                Eres un profesor de español especializado en crear consignas de escritura.
                Genera una consigna breve, clara y motivadora para practicar escritura en español.
                La consigna debe ser apropiada para el nivel indicado y el tipo de texto solicitado.
                Incluye:
                1. Título claro
                2. Contexto o situación
                3. Tarea específica
                4. Extensión recomendada (palabras)
                5. Elementos lingüísticos a incluir (apropiados para el nivel)
                
                Devuelve SOLO la consigna, sin comentarios adicionales ni explicaciones.
                """

                user_msg = f"""
                Genera una consigna de escritura para:
                - Nivel: {nivel}
                - Tipo de texto: {tipo_texto}
                
                Asegúrate que sea adecuada para el nivel del estudiante.
                """

                # Llamar a OpenAI para generar consigna
                _, resultado = obtener_json_de_openai(system_msg, user_msg)

                # Extraer consigna - manejar diferentes formatos de respuesta posibles
                consigna = ""
                if isinstance(resultado, dict):
                    if "error" in resultado:
                        raise Exception(resultado["error"])
                    consigna = resultado.get("content", "")
                    if not consigna:
                        consigna = resultado.get("consigna", "")
                    if not consigna:
                        # Buscar el primer valor de texto largo en el diccionario
                        for key, value in resultado.items():
                            if isinstance(value, str) and len(value) > 50:
                                consigna = value
                                break
                elif isinstance(resultado, str):
                    consigna = resultado

                # Verificar que tenemos una consigna
                if consigna:
                    # Guardar en session_state
                    set_session_var("consigna_actual", consigna)
                else:
                    st.error(
                        "No se pudo generar una consigna. Inténtalo de nuevo.")

            except Exception as e:
                handle_exception("mostrar_herramienta_consignas", e)
                st.error(f"Error al generar consigna: {str(e)}")

    # Mostrar consigna generada si existe
    consigna = get_session_var("consigna_actual", "")

    if consigna:
        st.markdown("---")
        st.subheader("Consigna generada:")

        # Formato mejorado para la consigna
        st.markdown(f"""
        <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #5c7cfa;">
            {consigna}
        </div>
        """, unsafe_allow_html=True)

        # Opciones para usar la consigna
        col1, col2 = st.columns([1, 1])

        with col1:
            if st.button("Escribir respuesta a esta consigna"):
                # Preparar para corrección
                set_session_var("texto_correccion_corregir", "")
                set_session_var("info_adicional_corregir",
                                f"Respuesta a la consigna: {consigna[:100]}...")
                set_session_var("usar_consigna_como_texto", True)

                # Cambiar a pestaña de corrección
                set_session_var("active_tab", 0)
                st.rerun()

        with col2:
            if st.button("Generar otra consigna"):
                # Limpiar consigna actual
                set_session_var("consigna_actual", "")
                st.rerun()


def mostrar_configuracion_simulacro():
    """Muestra la configuración para iniciar un simulacro."""
    # Opciones para el simulacro
    col1, col2 = st.columns(2)

    with col1:
        # Nivel del examen
        nivel_options = [
            "A1 - Acceso",
            "A2 - Plataforma",
            "B1 - Umbral",
            "B2 - Avanzado",
            "C1 - Dominio operativo eficaz",
            "C2 - Maestría"
        ]

        # Determinar índice por defecto basado en nivel del usuario
        nivel_defecto = 2  # B1 por defecto
        nivel_usuario = get_session_var("nivel_estudiante", "intermedio")

        if nivel_usuario == "principiante":
            nivel_defecto = 1  # A2
        elif nivel_usuario == "intermedio":
            nivel_defecto = 2  # B1
        elif nivel_usuario == "avanzado":
            nivel_defecto = 4  # C1

        nivel_examen = st.selectbox(
            "Nivel del examen:",
            options=nivel_options,
            index=nivel_defecto
        )

    with col2:
        # Duración del simulacro
        duracion_options = [
            "15 minutos",
            "30 minutos",
            "45 minutos",
            "60 minutos"
        ]

        duracion = st.selectbox(
            "Duración del simulacro:",
            options=duracion_options,
            index=1  # 30 minutos por defecto
        )

        # Convertir duración a minutos
        duracion_minutos = int(duracion.split()[0])

    # Botón para iniciar simulacro
    if st.button("Iniciar simulacro", type="primary"):
        with st.spinner("Preparando simulacro..."):
            try:
                # Generar tarea de examen con OpenAI
                # Obtener solo el código (A1, B2, etc.)
                nivel_corto = nivel_examen.split()[0]

                system_msg = """
                Eres un profesor especializado en exámenes DELE (Diplomas de Español como Lengua Extranjera).
                Genera una tarea de expresión escrita realista que podría aparecer en un examen DELE del nivel indicado.
                
                La tarea debe incluir:
                1. Instrucciones claras
                2. Contexto o situación
                3. Extensión requerida
                4. Elementos a incluir (si aplica)
                5. Criterios que se evaluarán
                
                Asegúrate de que el nivel de dificultad, vocabulario y estructuras gramaticales sean 
                apropiados para el nivel solicitado. Sigue el formato y estilo oficial de los exámenes DELE.
                """

                user_msg = f"""
                Genera una tarea de expresión escrita para el examen DELE {nivel_corto}.
                Debe ser una tarea realista y completa que evalúe el nivel de expresión escrita del candidato.
                """

                # Llamar a OpenAI
                _, resultado = obtener_json_de_openai(system_msg, user_msg)

                # Extraer tarea
                tarea = ""
                if isinstance(resultado, dict):
                    if "error" in resultado:
                        raise Exception(resultado["error"])

                    # Buscar un campo que pueda contener la tarea
                    campos_posibles = ["tarea", "content",
                                       "texto", "instrucciones"]
                    for campo in campos_posibles:
                        if campo in resultado and isinstance(resultado[campo], str):
                            tarea = resultado[campo]
                            break

                    # Si no encontramos en campos específicos, buscar texto largo
                    if not tarea:
                        for key, value in resultado.items():
                            if isinstance(value, str) and len(value) > 100:
                                tarea = value
                                break
                elif isinstance(resultado, str):
                    tarea = resultado

                # Verificar que tenemos una tarea
                if tarea:
                    # Guardar datos del simulacro
                    set_session_var("inicio_simulacro",
                                    datetime.now().isoformat())
                    set_session_var("duracion_simulacro", duracion_minutos)
                    set_session_var("tarea_simulacro", tarea)
                    set_session_var("simulacro_respuesta_texto", "")

                    # Recargar para mostrar el simulacro
                    st.rerun()
                else:
                    st.error("No se pudo generar la tarea. Inténtalo de nuevo.")

            except Exception as e:
                handle_exception("mostrar_configuracion_simulacro", e)
                st.error(f"Error al generar la tarea: {str(e)}")


def ui_countdown_timer(total_seconds, start_time=None):
    """
    Muestra un temporizador de cuenta regresiva.

    Args:
        total_seconds: Tiempo total en segundos
        start_time: Tiempo de inicio (None = ahora)

    Returns:
        dict: Estado del temporizador
    """
    # Manejar el caso donde total_seconds es None
    if total_seconds is None:
        total_seconds = 0  # Usar 0 como valor por defecto

    if start_time is None:
        start_time = time.time()

    # Calcular tiempo transcurrido
    tiempo_transcurrido = time.time() - start_time
    tiempo_restante_segundos = max(0, total_seconds - tiempo_transcurrido)

    # Formatear tiempo restante
    minutos = int(tiempo_restante_segundos // 60)
    segundos = int(tiempo_restante_segundos % 60)
    tiempo_formateado = f"{minutos:02d}:{segundos:02d}"

    # Calcular porcentaje (evitar división por cero)
    if total_seconds > 0:
        porcentaje = 1 - (tiempo_restante_segundos / total_seconds)
    else:
        porcentaje = 1  # Si no hay tiempo total, consideramos que está completo
    porcentaje = max(0, min(1, porcentaje))  # Asegurar entre 0 y 1

    # Determinar color según tiempo restante
    if tiempo_restante_segundos > total_seconds * 0.5:  # Más del 50% restante
        color = "normal"  # Verde/Normal
    elif tiempo_restante_segundos > total_seconds * 0.25:  # Entre 25% y 50%
        color = "warning"  # Amarillo/Advertencia
    else:  # Menos del 25%
        color = "error"  # Rojo/Error

    return {
        "tiempo_restante": tiempo_restante_segundos,
        "tiempo_formateado": tiempo_formateado,
        "porcentaje": porcentaje,
        "color": color,
        "terminado": tiempo_restante_segundos <= 0
    }


def mostrar_simulacro_en_progreso():
    """
    Muestra un simulacro de examen en progreso con un cronómetro que funciona correctamente.
    Implementación mejorada basada en la función ui_countdown_timer.
    """
    # Obtener datos del simulacro
    inicio_str = get_session_var("inicio_simulacro", None)
    duracion_minutos = get_session_var("duracion_simulacro", 30)
    tarea = get_session_var("tarea_simulacro", "")

    # Verificar si es la primera vez que carga o hay que inicializar
    if inicio_str is None:
        # Podría ocurrir en casos especiales, vamos a reinicializar
        set_session_var("inicio_simulacro", time.time())
        inicio_timestamp = time.time()
    else:
        # Intentar convertir a timestamp
        try:
            # Si el inicio está en formato ISO, convertirlo a timestamp
            if isinstance(inicio_str, str) and '-' in inicio_str:
                inicio_datetime = datetime.fromisoformat(inicio_str)
                inicio_timestamp = inicio_datetime.timestamp()
            else:
                # Si es posible que ya sea un número (timestamp)
                inicio_timestamp = float(inicio_str)
        except (ValueError, TypeError):
            # Si hay error en la conversión, usar el tiempo actual
            inicio_timestamp = time.time()
            set_session_var("inicio_simulacro", inicio_timestamp)

    # Calcular tiempo total del simulacro en segundos
    duracion_total_segundos = duracion_minutos * 60

    # Obtener el estado del temporizador
    timer_state = ui_countdown_timer(duracion_total_segundos, inicio_timestamp)

    # Mostrar cabecera del simulacro
    st.subheader("🕒 Simulacro de examen DELE en progreso")

    # Mostrar temporizador con estilo según estado
    col1, col2 = st.columns([3, 1])

    with col1:
        # Barra de progreso
        st.progress(timer_state["porcentaje"])

    with col2:
        # Mostrar tiempo restante con color adecuado
        if timer_state["terminado"]:
            st.error(f"⏰ ¡TIEMPO TERMINADO!")
        elif timer_state["color"] == "warning":
            st.warning(f"⏱️ {timer_state['tiempo_formateado']}")
        elif timer_state["color"] == "error":
            st.error(f"⏱️ {timer_state['tiempo_formateado']}")
        else:
            st.info(f"⏱️ {timer_state['tiempo_formateado']}")

    # Mostrar tarea
    st.markdown("### Tarea:")
    st.markdown(f"""
    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #5c7cfa;">
        {tarea}
    </div>
    """, unsafe_allow_html=True)

    # Campo para la respuesta
    respuesta = st.text_area(
        "Tu respuesta:",
        value=get_session_var("simulacro_respuesta_texto", ""),
        height=300,
        key="simulacro_respuesta"
    )

    # Guardar respuesta en session_state
    set_session_var("simulacro_respuesta_texto", respuesta)

    # Botones de acción
    col1, col2 = st.columns([1, 1])

    with col1:
        if st.button("Finalizar y evaluar", type="primary"):
            if not respuesta.strip():
                st.warning(
                    "Por favor, escribe tu respuesta antes de finalizar.")
            else:
                # Preparar para corrección
                set_session_var("texto_correccion_corregir", respuesta)
                set_session_var("info_adicional_corregir",
                                f"Simulacro de examen DELE. Tarea: {tarea[:100]}...")

                # Reiniciar el simulacro
                set_session_var("inicio_simulacro", None)
                set_session_var("duracion_simulacro", None)
                set_session_var("tarea_simulacro", None)

                # Cambiar a pestaña de corrección
                set_session_var("active_tab", 0)
                st.rerun()

    with col2:
        if st.button("Cancelar simulacro"):
            # Reiniciar el simulacro
            set_session_var("inicio_simulacro", None)
            set_session_var("duracion_simulacro", None)
            set_session_var("tarea_simulacro", None)

            st.rerun()

    # Agregar un mecanismo para auto-actualizar - elemento clave para que funcione
    # Este placeholder obligará a Streamlit a re-renderizar la página y actualizar el temporizador
    refresh_placeholder = st.empty()

    # Si usamos placeholder y tenemos contenido simple que cambia en cada renderizado
    # Streamlit actualizará la UI naturalmente
    with refresh_placeholder:
        # Este texto invisible cambia en cada renderizado y está destinado
        # a "engañar" a Streamlit para que actualice el estado del componente
        st.markdown(
            f'<div style="display:none">{time.time()}</div>', unsafe_allow_html=True)


def view_simulador_examen():
    """
    Muestra el simulador de examen DELE con cronómetro optimizado.
    """
    st.header("📝 Simulador de examen")

    st.markdown("""
    Practica para el examen DELE con simulacros de expresión escrita
    adaptados a tu nivel de español.
    """)

    # Verificar si hay un simulacro en progreso
    inicio_simulacro = get_session_var("inicio_simulacro", None)
    duracion_simulacro = get_session_var("duracion_simulacro", None)
    tarea_simulacro = get_session_var("tarea_simulacro", None)

    if inicio_simulacro and duracion_simulacro and tarea_simulacro:
        # Mostrar simulacro en progreso con cronómetro mejorado
        mostrar_simulacro_en_progreso()
    else:
        # Mostrar configuración de simulacro
        mostrar_configuracion_simulacro()


def mostrar_configuracion_simulacro():
    """
    Muestra la configuración para iniciar un simulacro.
    Optimizado para usar timestamps en vez de strings para controlar el tiempo.
    """
    # Opciones para el simulacro
    col1, col2 = st.columns(2)

    with col1:
        # Nivel del examen
        nivel_options = [
            "A1 - Acceso",
            "A2 - Plataforma",
            "B1 - Umbral",
            "B2 - Avanzado",
            "C1 - Dominio operativo eficaz",
            "C2 - Maestría"
        ]

        # Determinar índice por defecto basado en nivel del usuario
        nivel_defecto = 2  # B1 por defecto
        nivel_usuario = get_session_var("nivel_estudiante", "intermedio")

        if nivel_usuario == "principiante":
            nivel_defecto = 1  # A2
        elif nivel_usuario == "intermedio":
            nivel_defecto = 2  # B1
        elif nivel_usuario == "avanzado":
            nivel_defecto = 4  # C1

        nivel_examen = st.selectbox(
            "Nivel del examen:",
            options=nivel_options,
            index=nivel_defecto
        )

    with col2:
        # Duración del simulacro
        duracion_options = [
            "15 minutos",
            "30 minutos",
            "45 minutos",
            "60 minutos"
        ]

        duracion = st.selectbox(
            "Duración del simulacro:",
            options=duracion_options,
            index=1  # 30 minutos por defecto
        )

        # Convertir duración a minutos
        duracion_minutos = int(duracion.split()[0])

    # Botón para iniciar simulacro
    if st.button("Iniciar simulacro", type="primary"):
        with st.spinner("Preparando simulacro..."):
            try:
                # Generar tarea de examen con OpenAI
                # Obtener solo el código (A1, B2, etc.)
                nivel_corto = nivel_examen.split()[0]

                system_msg = """
                Eres un profesor especializado en exámenes DELE (Diplomas de Español como Lengua Extranjera).
                Genera una tarea de expresión escrita realista que podría aparecer en un examen DELE del nivel indicado.
                
                La tarea debe incluir:
                1. Instrucciones claras
                2. Contexto o situación
                3. Extensión requerida
                4. Elementos a incluir (si aplica)
                5. Criterios que se evaluarán
                
                Asegúrate de que el nivel de dificultad, vocabulario y estructuras gramaticales sean 
                apropiados para el nivel solicitado. Sigue el formato y estilo oficial de los exámenes DELE.
                """

                user_msg = f"""
                Genera una tarea de expresión escrita para el examen DELE {nivel_corto}.
                Debe ser una tarea realista y completa que evalúe el nivel de expresión escrita del candidato.
                """

                # Llamar a OpenAI
                _, resultado = obtener_json_de_openai(system_msg, user_msg)

                # Extraer tarea
                tarea = ""
                if isinstance(resultado, dict):
                    if "error" in resultado:
                        raise Exception(resultado["error"])

                    # Buscar un campo que pueda contener la tarea
                    campos_posibles = ["tarea", "content",
                                       "texto", "instrucciones"]
                    for campo in campos_posibles:
                        if campo in resultado and isinstance(resultado[campo], str):
                            tarea = resultado[campo]
                            break

                    # Si no encontramos en campos específicos, buscar texto largo
                    if not tarea:
                        for key, value in resultado.items():
                            if isinstance(value, str) and len(value) > 100:
                                tarea = value
                                break
                elif isinstance(resultado, str):
                    tarea = resultado

                # Verificar que tenemos una tarea
                if tarea:
                    # Guardar datos del simulacro - usar timestamp en lugar de string para el inicio
                    # Timestamp actual
                    set_session_var("inicio_simulacro", time.time())
                    set_session_var("duracion_simulacro", duracion_minutos)
                    set_session_var("tarea_simulacro", tarea)
                    set_session_var("simulacro_respuesta_texto", "")

                    # Recargar para mostrar el simulacro
                    st.rerun()
                else:
                    st.error("No se pudo generar la tarea. Inténtalo de nuevo.")

            except Exception as e:
                handle_exception("mostrar_configuracion_simulacro", e)
                st.error(f"Error al generar la tarea: {str(e)}")


# Mejoras para la función view_plan_estudio()

def view_plan_estudio():
    """Muestra planes de estudio recomendados con mejor manejo de errores y visualización."""
    st.header("📚 Plan de estudio")

    # Obtener nivel del estudiante
    nivel = get_session_var("nivel_estudiante", "intermedio")

    # Mostrar introducción
    st.markdown(f"""
    Aquí encontrarás un plan de estudio personalizado para mejorar tu español.
    Este plan está adaptado a tu nivel actual: **{nivel.capitalize()}**.
    """)

    # Estado para mostrar el progreso de generación
    generacion_plan_estado = st.empty()
    plan_generado = get_session_var("plan_estudio_generado", "")

    # Crear el botón para generar plan
    if not plan_generado and st.button("✨ Generar plan de estudio personalizado", type="primary"):
        # Iniciar un spinner en el contenedor vacío
        with generacion_plan_estado.container():
            st.info(
                "Generando plan de estudio personalizado... Esto puede tomar unos segundos.")

            progress_bar = st.progress(0)
            for i in range(101):
                # Simular progreso para dar feedback al usuario
                progress_bar.progress(i)
                time.sleep(0.01)  # Pequeña pausa para mostrar progreso visual

            try:
                # Generar plan con OpenAI con un timeout controlado
                system_msg = """
                Eres un experto en enseñanza de español como lengua extranjera (ELE).
                Genera un plan de estudio personalizado para mejorar el español del estudiante,
                adaptado a su nivel. El plan debe incluir:
                
                1. Objetivos específicos para el nivel
                2. Actividades recomendadas (escritura, lectura, conversación, etc.)
                3. Recursos específicos (libros, páginas web, etc.)
                4. Cronograma sugerido (diario, semanal, mensual)
                5. Consejos para la práctica independiente
                
                Organiza el plan de manera clara y motivadora. Asegúrate de que sea realista
                y adaptado al nivel indicado.
                
                Formatea tu respuesta en Markdown para mejor legibilidad.
                """

                user_msg = f"""
                Genera un plan de estudio para un estudiante de español de nivel {nivel}.
                El plan debe ser completo pero conciso, con recomendaciones prácticas y específicas.
                """

                # Utilizamos una versión con timeout para evitar bloqueos
                def generate_plan_with_timeout():
                    # Llamar a OpenAI
                    try:
                        raw_output, resultado = obtener_json_de_openai(
                            system_msg, user_msg)
                        return raw_output, resultado
                    except Exception as e:
                        logger.error(
                            f"Error en generate_plan_with_timeout: {str(e)}")
                        return None, {"error": str(e)}

                # Ejecutar con timeout
                import concurrent.futures
                from concurrent.futures import ThreadPoolExecutor

                with ThreadPoolExecutor() as executor:
                    future = executor.submit(generate_plan_with_timeout)
                    try:
                        raw_output, resultado = future.result(
                            timeout=30)  # 30 segundos máximo
                    except concurrent.futures.TimeoutError:
                        # Si tarda demasiado, usar plan de fallback
                        raw_output = None
                        resultado = {"error": "Timeout en generación del plan"}

                # Extraer plan
                plan = ""
                if raw_output:
                    plan = raw_output
                elif isinstance(resultado, dict):
                    if "error" in resultado:
                        raise Exception(resultado["error"])

                    # Buscar un campo que pueda contener el plan
                    campos_posibles = [
                        "plan", "content", "texto", "plan_estudio"]
                    for campo in campos_posibles:
                        if campo in resultado and isinstance(resultado[campo], str):
                            plan = resultado[campo]
                            break

                    # Si no encontramos en campos específicos, buscar texto largo
                    if not plan:
                        for key, value in resultado.items():
                            if isinstance(value, str) and len(value) > 200:
                                plan = value
                                break
                elif isinstance(resultado, str):
                    plan = resultado

                # Verificar que tenemos un plan
                if plan:
                    # Guardar plan
                    set_session_var("plan_estudio_generado", plan)
                    generacion_plan_estado.success(
                        "✅ Plan de estudio generado correctamente")
                else:
                    # Plan de fallback básico si no hay contenido
                    generacion_plan_estado.warning(
                        "⚠️ No se pudo generar un plan personalizado. Mostrando plan general.")
                    set_session_var("plan_estudio_generado",
                                    crear_plan_fallback(nivel))

            except Exception as e:
                logger.error(f"Error al generar plan de estudio: {str(e)}")
                # Plan de fallback en caso de error
                generacion_plan_estado.error(
                    f"❌ Error al generar plan: {str(e)}. Mostrando plan general.")
                set_session_var("plan_estudio_generado",
                                crear_plan_fallback(nivel))

            # Recargar la página para mostrar el plan generado
            st.rerun()

    # Obtener el plan actualizado después de la posible generación
    plan = get_session_var("plan_estudio_generado", "")

    # Mostrar el plan
    if plan:
        # Usar un contenedor expandido con estilo para mejor visibilidad
        st.markdown("""
        <style>
        .plan-container {
            background-color: #f8f9fa;
            border-left: 5px solid #3FF2E1;
            padding: 20px;
            border-radius: 10px;
            margin: 20px 0;
            line-height: 1.6;
        }
        </style>
        """, unsafe_allow_html=True)

        st.markdown('<div class="plan-container">', unsafe_allow_html=True)
        st.markdown(plan)
        st.markdown('</div>', unsafe_allow_html=True)

        # Botones para acciones adicionales
        col1, col2 = st.columns([1, 1])

        with col1:
            # Botón para descargar plan
            contenido_descarga = plan
            if isinstance(plan, dict):
                contenido_descarga = json.dumps(plan)
            elif not isinstance(plan, str):
                contenido_descarga = str(plan)

            st.download_button(
                label="📥 Descargar plan de estudio",
                data=contenido_descarga,
                file_name=f"plan_estudio_{nivel}_{datetime.now().strftime('%Y%m%d')}.md",
                mime="text/markdown"
            )

        with col2:
            # Botón para regenerar el plan
            if st.button("🔄 Generar nuevo plan"):
                # Limpiar plan existente
                set_session_var("plan_estudio_generado", None)
                st.rerun()
    else:
        # Mensaje para generar el primer plan
        st.info("👆 Haz clic en el botón para generar un plan de estudio personalizado.")

    # Mostrar enlace a otras secciones
    st.markdown("---")
    st.markdown("""
    ### Complementa tu plan de estudio
    
    Utiliza las herramientas de la aplicación para mejorar tu español:
    
    - Corrige tus textos para recibir feedback detallado
    - Practica con simulacros para prepararte para exámenes
    - Utiliza las herramientas ELE para actividades adicionales
    """)


def crear_plan_fallback(nivel):
    """
    Crea un plan de estudio básico en caso de que la generación con OpenAI falle.

    Args:
        nivel: Nivel del estudiante

    Returns:
        str: Plan de estudio básico
    """
    # Planes por nivel
    planes = {
        "principiante": """
# Plan de Estudio para Nivel Principiante (A1-A2)

## Objetivos
- Desarrollar vocabulario básico para situaciones cotidianas
- Aprender estructuras gramaticales fundamentales
- Practicar la comprensión de textos y audios sencillos
- Comunicarse en situaciones de la vida diaria

## Actividades Recomendadas
### Diarias (20-30 minutos)
- Practicar con flashcards de vocabulario nuevo
- Escuchar audios cortos en español
- Escribir oraciones sencillas con el vocabulario aprendido

### Semanales
- Leer un texto corto adaptado a tu nivel
- Ver un video o episodio de serie para principiantes
- Completar ejercicios de gramática básica
- Practicar conversación con frases memorizadas

## Recursos
- Aplicaciones: Duolingo, Memrise, Babbel
- Sitios web: ProfeDeELE.es, VerdesMontañas.com
- Canales de YouTube: Español con Juan, Dreaming Spanish
- Libros: "Aula Internacional 1", "Gente Hoy 1"

## Cronograma Sugerido
- **Lunes**: Vocabulario nuevo + ejercicios
- **Martes**: Práctica de comprensión auditiva
- **Miércoles**: Gramática básica
- **Jueves**: Lectura y comprensión
- **Viernes**: Práctica de expresión escrita
- **Fin de semana**: Repaso general y contenido multimedia

## Consejos
- Utiliza el Textocorrector ELE para revisar tus textos
- Aprende frases completas, no solo palabras sueltas
- Intenta pensar directamente en español
- No te preocupes por los errores, son parte del aprendizaje
        """,

        "intermedio": """
# Plan de Estudio para Nivel Intermedio (B1-B2)

## Objetivos
- Ampliar vocabulario para expresar opiniones y sentimientos
- Dominar los tiempos verbales más comunes
- Mejorar fluidez en conversaciones cotidianas
- Comprender textos y audios de complejidad media

## Actividades Recomendadas
### Diarias (30-45 minutos)
- Leer noticias o artículos adaptados
- Escuchar podcasts en español
- Practicar escritura con diario personal
- Repasar conjugaciones verbales

### Semanales
- Ver un episodio de serie o película en español
- Participar en intercambios de idiomas
- Escribir textos de 250-300 palabras sobre temas diversos
- Practicar situaciones comunicativas específicas

## Recursos
- Podcasts: Radio Ambulante, Español Automático
- Sitios web: Profedeele.es, NoticiasParaTuNivel
- Canales de YouTube: WhyNotSpanish, Spanish with Vicente
- Libros: "Gente Hoy 2", "Aula Internacional 3-4"

## Cronograma Sugerido
- **Lunes**: Vocabulario temático + expresiones
- **Martes**: Gramática avanzada
- **Miércoles**: Comprensión auditiva y expresión oral
- **Jueves**: Lectura de textos y análisis
- **Viernes**: Redacción y práctica escrita
- **Fin de semana**: Inmersión en contenido auténtico

## Consejos
- Utiliza el Textocorrector ELE para identificar patrones de error
- Integra expresiones idiomáticas en tu vocabulario
- Practica el subjuntivo en contextos reales
- Busca temas que te interesen para mantener la motivación
        """,

        "avanzado": """
# Plan de Estudio para Nivel Avanzado (C1-C2)

## Objetivos
- Perfeccionar matices de expresión y registro
- Dominar estructuras complejas y usos idiomáticos
- Comunicarse con fluidez sobre temas especializados
- Comprender textos literarios y académicos

## Actividades Recomendadas
### Diarias (45-60 minutos)
- Leer literatura, prensa o textos académicos
- Escuchar podcasts o ver contenido sin subtítulos
- Practicar escritura argumentativa o creativa
- Ampliar vocabulario especializado

### Semanales
- Debatir sobre temas de actualidad
- Analizar textos literarios o periodísticos
- Redactar ensayos de 500+ palabras
- Practicar diferentes registros lingüísticos

## Recursos
- Literatura: Autores contemporáneos hispanohablantes
- Prensa: El País, BBC Mundo, periódicos locales
- Podcasts: El Podcast de Historia, TED en Español
- Academias online: Baselang, AIL Madrid (cursos específicos)

## Cronograma Sugerido
- **Lunes**: Análisis de textos complejos
- **Martes**: Refuerzo de puntos gramaticales específicos
- **Miércoles**: Práctica avanzada de conversación
- **Jueves**: Redacción especializada
- **Viernes**: Comprensión auditiva de material auténtico
- **Fin de semana**: Inmersión cultural con cine, literatura o música

## Consejos
- Utiliza el Textocorrector ELE para perfeccionar tus textos
- Especialízate en tu área de interés profesional o académico
- Practica la mediación lingüística y cultural
- Involúcrate en comunidades de hablantes nativos
        """
    }

    # Devolver el plan correspondiente al nivel o el de intermedio por defecto
    return planes.get(nivel, planes["intermedio"])

# Sistema de generación automática de ejercicios personalizados
# Añadir este código como un nuevo artefacto o integrado en artefactos existentes


def generar_ejercicio_personalizado(tipo_ejercicio, nivel, area_mejora):
    """
    Genera un ejercicio personalizado según el tipo, nivel y área de mejora.

    Args:
        tipo_ejercicio: Tipo de ejercicio a generar
        nivel: Nivel del estudiante
        area_mejora: Área específica en la que enfocarse

    Returns:
        dict: Ejercicio generado con instrucciones, contenido y solución
    """
    try:
        # Configurar prompt para OpenAI
        system_msg = f"""
        Eres un experto en enseñanza de español como lengua extranjera (ELE).
        Genera un ejercicio personalizado de tipo {tipo_ejercicio} para un estudiante de nivel {nivel},
        enfocado específicamente en mejorar el área de {area_mejora}.
        
        El ejercicio debe incluir:
        1. Instrucciones claras
        2. El contenido del ejercicio
        3. Solución completa para autocorrección
        
        Formatea la respuesta como un JSON con la siguiente estructura exacta:
        {{
            "titulo": "Título del ejercicio",
            "tipo": "{tipo_ejercicio}",
            "nivel": "{nivel}",
            "area_mejora": "{area_mejora}",
            "instrucciones": "Instrucciones detalladas para el estudiante",
            "contenido": "Contenido del ejercicio (texto, oraciones, etc.)",
            "solucion": "Solución completa del ejercicio"
        }}
        
        Asegúrate de que:
        - El nivel de dificultad sea apropiado para {nivel}
        - El ejercicio se enfoque específicamente en {area_mejora}
        - Las instrucciones sean claras y detalladas
        - La solución incluya explicaciones cuando sea necesario
        """

        user_msg = f"""
        Crea un ejercicio personalizado de {tipo_ejercicio} para nivel {nivel} enfocado en mejorar {area_mejora}.
        El ejercicio debe ser práctico, específico y con solución detallada.
        """

        # Realizar la llamada a OpenAI
        _, resultado = obtener_json_de_openai(system_msg, user_msg)

        # Verificar si hay un error en la respuesta
        if "error" in resultado:
            logger.error(f"Error al generar ejercicio: {resultado['error']}")
            return {
                "error": f"No se pudo generar el ejercicio: {resultado['error']}",
                "titulo": f"Ejercicio de {tipo_ejercicio} (nivel {nivel})",
                "tipo": tipo_ejercicio,
                "nivel": nivel,
                "area_mejora": area_mejora,
                "instrucciones": "No disponible",
                "contenido": "No disponible",
                "solucion": "No disponible"
            }

        # Si el resultado es válido, devolverlo
        return resultado

    except Exception as e:
        logger.error(f"Error en generar_ejercicio_personalizado: {str(e)}")
        return {
            "error": str(e),
            "titulo": f"Ejercicio de {tipo_ejercicio} (nivel {nivel})",
            "tipo": tipo_ejercicio,
            "nivel": nivel,
            "area_mejora": area_mejora,
            "instrucciones": "No disponible",
            "contenido": "No disponible",
            "solucion": "No disponible"
        }


def view_ejercicios_personalizados():
    """
    Muestra la interfaz para generar y resolver ejercicios personalizados.
    Esta función puede integrarse como una nueva pestaña o dentro de una sección existente.
    """
    st.header("🎯 Ejercicios Personalizados")

    # Obtener nivel y área de mejora del usuario
    nivel = get_session_var("nivel_estudiante", "intermedio")
    area_mejora = get_session_var("ultima_area_mejora", None)

    # Introducción
    st.markdown(f"""
    Genera ejercicios personalizados adaptados a tu nivel y necesidades específicas.
    Estos ejercicios te ayudarán a practicar y mejorar tus habilidades en español.
    """)

    # Configuración del ejercicio
    st.subheader("Configuración del ejercicio")
    col1, col2, col3 = st.columns(3)

    with col1:
        tipo_ejercicio_opciones = [
            "Completar huecos",
            "Ordenar palabras/frases",
            "Elección múltiple",
            "Corrección de errores",
            "Transformación de frases",
            "Vocabulario en contexto"
        ]
        tipo_ejercicio = st.selectbox(
            "Tipo de ejercicio",
            options=tipo_ejercicio_opciones
        )

    with col2:
        nivel_options = [
            "principiante (A1-A2)",
            "intermedio (B1-B2)",
            "avanzado (C1-C2)"
        ]

        nivel_index = 0
        if nivel == "principiante":
            nivel_index = 0
        elif nivel == "intermedio":
            nivel_index = 1
        elif nivel == "avanzado":
            nivel_index = 2

        nivel_seleccionado = st.selectbox(
            "Nivel",
            options=nivel_options,
            index=nivel_index
        )

    with col3:
        areas_mejora_opciones = [
            "Gramática",
            "Léxico",
            "Puntuación",
            "Estructura textual",
            "Coherencia",
            "Cohesión",
            "Registro lingüístico"
        ]

        # Establecer el área de mejora por defecto según el historial del usuario
        area_index = 0
        if area_mejora:
            try:
                area_index = areas_mejora_opciones.index(area_mejora)
            except ValueError:
                area_index = 0

        area_mejora_seleccionada = st.selectbox(
            "Área de mejora",
            options=areas_mejora_opciones,
            index=area_index,
            help="Área específica en la que quieres enfocarte"
        )

    # Botón para generar ejercicio
    generar_clicked = st.button("Generar ejercicio", type="primary")

    # Estado para almacenar el ejercicio generado
    if "ejercicio_actual" not in st.session_state:
        st.session_state.ejercicio_actual = None

    if "mostrar_solucion" not in st.session_state:
        st.session_state.mostrar_solucion = False

    if "respuesta_usuario" not in st.session_state:
        st.session_state.respuesta_usuario = ""

    # Generar ejercicio si se hace clic en el botón
    if generar_clicked:
        with st.spinner("Generando ejercicio personalizado..."):
            st.session_state.ejercicio_actual = generar_ejercicio_personalizado(
                tipo_ejercicio,
                nivel_seleccionado,
                area_mejora_seleccionada
            )
            # Reiniciar estado de solución y respuesta
            st.session_state.mostrar_solucion = False
            st.session_state.respuesta_usuario = ""

            # Si el ejercicio fue generado con éxito, mostrar texto confirmando generación
            if st.session_state.ejercicio_actual and "error" not in st.session_state.ejercicio_actual:
                st.success("¡Ejercicio generado correctamente!")

    # Mostrar ejercicio generado
    if st.session_state.ejercicio_actual:
        ejercicio = st.session_state.ejercicio_actual

        # Verificar si hay error
        if "error" in ejercicio and ejercicio.get("instrucciones") == "No disponible":
            st.error(
                f"No se pudo generar el ejercicio: {ejercicio.get('error', 'Error desconocido')}")
        else:
            # Mostrar ejercicio con formato mejorado
            st.markdown("---")
            st.subheader(
                f"📝 {ejercicio.get('titulo', 'Ejercicio personalizado')}")

            # Aplicar CSS para el contenedor de ejercicio
            st.markdown("""
            <style>
            .ejercicio-container {
                background-color: #f0f7ff;
                border-left: 5px solid #3785e5;
                padding: 20px;
                border-radius: 10px;
                margin: 20px 0;
            }
            .instrucciones {
                font-weight: bold;
                margin-bottom: 15px;
                color: #3785e5;
            }
            .contenido {
                background-color: white;
                padding: 15px;
                border-radius: 8px;
                border: 1px solid #e0e0e0;
                margin-top: 10px;
            }
            .solucion-container {
                background-color: #eefbf4;
                border-left: 5px solid #28a745;
                padding: 20px;
                border-radius: 10px;
                margin: 20px 0;
            }
            </style>
            """, unsafe_allow_html=True)

            # Renderizar instrucciones y contenido con HTML seguro
            container_html = f"""
            <div class="ejercicio-container">
                <div class="instrucciones">{ejercicio.get('instrucciones', '')}</div>
                <div class="contenido">{ejercicio.get('contenido', '')}</div>
            </div>
            """
            st.markdown(container_html, unsafe_allow_html=True)

            # Campo para la respuesta del usuario
            st.session_state.respuesta_usuario = st.text_area(
                "Tu respuesta:",
                value=st.session_state.respuesta_usuario,
                height=150
            )

            # Botones para verificar y mostrar solución
            col1, col2 = st.columns([1, 1])

            with col1:
                if st.button("Ver solución"):
                    st.session_state.mostrar_solucion = True

            with col2:
                # Guardar ejercicio en historial (si se implementa esa funcionalidad)
                if st.button("Guardar ejercicio"):
                    st.success("Ejercicio guardado en tu historial")
                    # Aquí se podría implementar la lógica de guardado

            # Mostrar solución si se solicita - SOLO cuando el usuario lo pide
            if st.session_state.mostrar_solucion:
                solution_html = f"""
                <div class="solucion-container">
                    <h4>✅ Solución</h4>
                    {ejercicio.get('solucion', '')}
                </div>
                """
                st.markdown(solution_html, unsafe_allow_html=True)
    else:
        # Mensaje inicial
        st.info("👆 Configura y genera un ejercicio personalizado para practicar.")

    # Sección para generación automática según áreas de mejora (sugerida)
    st.markdown("---")
    st.subheader("🤖 Generación automática basada en tus áreas de mejora")

    # Obtener área de mejora detectada del perfil
    area_detectada = get_session_var("ultima_area_mejora", None)

    if area_detectada:
        st.info(
            f"Según tu perfil, tu principal área de mejora es: **{area_detectada}**")
        if st.button("Generar ejercicio personalizado para mi área de mejora"):
            with st.spinner("Generando ejercicio específico..."):
                tipo_ejercicio_recomendado = "Vocabulario en contexto" if area_detectada == "Léxico" else "Corrección de errores"
                st.session_state.ejercicio_actual = generar_ejercicio_personalizado(
                    tipo_ejercicio_recomendado,
                    nivel_seleccionado,
                    area_detectada
                )
                st.session_state.mostrar_solucion = False
                st.session_state.respuesta_usuario = ""
                st.rerun()
    else:
        st.info("Completa algunas correcciones para que podamos identificar tus áreas de mejora y generar ejercicios específicos.")


def generar_ejercicio_personalizado(tipo_ejercicio, nivel, area_mejora):
    """
    Genera un ejercicio personalizado según el tipo, nivel y área de mejora.
    Versión mejorada con mejor separación de contenido y solución.

    Args:
        tipo_ejercicio: Tipo de ejercicio a generar
        nivel: Nivel del estudiante
        area_mejora: Área específica en la que enfocarse

    Returns:
        dict: Ejercicio generado con instrucciones, contenido y solución
    """
    try:
        # Configurar prompt para OpenAI con instrucciones explícitas para mejor formato
        system_msg = f"""
        Eres un experto en enseñanza de español como lengua extranjera (ELE).
        Genera un ejercicio personalizado de tipo {tipo_ejercicio} para un estudiante de nivel {nivel},
        enfocado específicamente en mejorar el área de {area_mejora}.
        
        El ejercicio debe incluir:
        1. Un título descriptivo y atractivo
        2. Instrucciones claras y directas para el estudiante
        3. El contenido del ejercicio (texto, oraciones, preguntas, etc.)
        4. Solución completa para autocorrección - IMPORTANTE: Esta debe estar separada claramente
        
        Formatea la respuesta como un JSON con la siguiente estructura exacta:
        {{
            "titulo": "Título descriptivo del ejercicio",
            "tipo": "{tipo_ejercicio}",
            "nivel": "{nivel}",
            "area_mejora": "{area_mejora}",
            "instrucciones": "Instrucciones detalladas para el estudiante",
            "contenido": "Contenido del ejercicio (texto, oraciones, etc.)",
            "solucion": "Solución completa del ejercicio"
        }}
        
        Asegúrate de que:
        - El nivel de dificultad sea apropiado para {nivel}
        - El ejercicio se enfoque específicamente en {area_mejora}
        - Las instrucciones NO mencionen la solución
        - El contenido NO contenga las respuestas o soluciones
        - La solución esté COMPLETAMENTE separada del contenido
        - Todo el texto HTML esté correctamente escapado para evitar problemas de visualización
        """

        user_msg = f"""
        Crea un ejercicio personalizado de {tipo_ejercicio} para nivel {nivel} enfocado en mejorar {area_mejora}.
        El ejercicio debe ser práctico, específico y con solución completamente separada.
        No incluyas la solución junto al ejercicio.
        """

        # Realizar la llamada a OpenAI
        _, resultado = obtener_json_de_openai(system_msg, user_msg)

        # Verificar si hay un error en la respuesta
        if "error" in resultado:
            logger.error(f"Error al generar ejercicio: {resultado['error']}")
            return {
                "error": f"No se pudo generar el ejercicio: {resultado['error']}",
                "titulo": f"Ejercicio de {tipo_ejercicio} (nivel {nivel})",
                "tipo": tipo_ejercicio,
                "nivel": nivel,
                "area_mejora": area_mejora,
                "instrucciones": "No disponible",
                "contenido": "No disponible",
                "solucion": "No disponible"
            }

        # Procesar el resultado: limpiar HTML o markdown potencialmente problemático
        # y asegurar que la solución esté completamente separada
        if isinstance(resultado, dict):
            # Sanitizar contenido HTML
            for campo in ['instrucciones', 'contenido', 'solucion', 'titulo']:
                if campo in resultado and isinstance(resultado[campo], str):
                    # Mantener el formato HTML pero eliminar posibles scripts maliciosos
                    texto = resultado[campo]
                    # Eliminar etiquetas script
                    texto = re.sub(r'<script.*?>.*?</script>',
                                   '', texto, flags=re.DOTALL)
                    # Asegurar que etiquetas como <div> estén correctamente cerradas
                    resultado[campo] = texto
                else:
                    # Si el campo no existe o no es string, asignar valor por defecto
                    resultado[campo] = resultado.get(campo, "No disponible")

            # Verificar que la solución no esté en el contenido
            if "solucion" in resultado and "contenido" in resultado:
                contenido = resultado["contenido"].lower()
                # Buscar palabras clave que podrían indicar que la solución está en el contenido
                palabras_clave = ["solución", "solution",
                                  "respuesta", "answer", "clave"]

                if any(palabra in contenido for palabra in palabras_clave):
                    # Si se detecta una posible solución en el contenido, intentar separar
                    partes = re.split(r'(solución|solution|respuesta|answer|clave)',
                                      resultado["contenido"],
                                      flags=re.IGNORECASE)
                    if len(partes) > 1:
                        # Mantener solo la primera parte como contenido
                        resultado["contenido"] = partes[0].strip()
                        # Si no hay solución definida, usar el resto como solución
                        if not resultado["solucion"] or resultado["solucion"] == "No disponible":
                            resultado["solucion"] = " ".join(
                                partes[1:]).strip()

        return resultado

    except Exception as e:
        logger.error(f"Error en generar_ejercicio_personalizado: {str(e)}")
        return {
            "error": str(e),
            "titulo": f"Ejercicio de {tipo_ejercicio} (nivel {nivel})",
            "tipo": tipo_ejercicio,
            "nivel": nivel,
            "area_mejora": area_mejora,
            "instrucciones": "No disponible",
            "contenido": "No disponible",
            "solucion": "No disponible"
        }


def generar_ejercicio_desde_area_mejora(nivel):
    """
    Genera un ejercicio personalizado basado en el área de mejora detectada
    en el perfil del estudiante.

    Args:
        nivel: Nivel del estudiante

    Returns:
        dict: Ejercicio generado
    """
    # Obtener área de mejora del perfil del estudiante
    area_mejora = get_session_var("ultima_area_mejora", None)

    if not area_mejora:
        # Si no hay área detectada, usar una general
        area_mejora = "Gramática"

    # Determinar el tipo de ejercicio más adecuado según el área
    tipo_ejercicio_map = {
        "Gramática": "Corrección de errores",
        "Léxico": "Vocabulario en contexto",
        "Puntuación": "Completar huecos",
        "Estructura textual": "Ordenar palabras/frases",
        "Coherencia": "Transformación de frases",
        "Cohesión": "Ordenar palabras/frases",
        "Registro lingüístico": "Elección múltiple"
    }

    tipo_ejercicio = tipo_ejercicio_map.get(area_mejora, "Completar huecos")

    # Generar el ejercicio
    return generar_ejercicio_personalizado(tipo_ejercicio, nivel, area_mejora)


def main():
    """
    Función principal que ejecuta la aplicación.
    Configura la interfaz de usuario y gestiona la navegación entre secciones.
    Versión actualizada con todas las nuevas funcionalidades.
    """
    # Cargar estilos personalizados
    load_custom_css()

    # Verificar login/pantalla de bienvenida
    if get_session_var("mostrar_login", True):
        if ui_login_screen():
            set_session_var("mostrar_login", False)
            st.rerun()
        return

    # Mostrar el sidebar con información
    show_sidebar_info()

    # Mostrar diagnóstico de conexiones en el sidebar
    if get_session_var("mostrar_login", False) == False:
        show_connection_status()

    # Mostrar header
    ui_header()

    # Crear pestañas para las diferentes secciones
    tabs = st.tabs([
        "✏️ Corrección de texto",
        "👤 Perfil y Progreso",
        "🛠️ Herramientas ELE",
        "📝 Simulador de examen",
        "📚 Plan de estudio",
        "🎯 Ejercicios Personalizados",  # Nueva pestaña añadida
        "ℹ️ Acerca de"
    ])

    # Cargar cada vista en su pestaña correspondiente
    with tabs[0]:
        view_correccion_texto()

    with tabs[1]:
        view_perfil_estudiante()

    with tabs[2]:
        view_herramientas_ele()

    with tabs[3]:
        view_simulador_examen()

    with tabs[4]:
        view_plan_estudio()

    # Nueva pestaña para Ejercicios Personalizados
    with tabs[5]:
        view_ejercicios_personalizados()

    with tabs[6]:  # El índice cambia debido a la nueva pestaña
        view_acerca_de()

    # Utilizar un callback para detectar cambios de pestaña (opcional)
    # Esta es una alternativa basada en interacciones de usuario
    for i, tab_name in enumerate(["correccion", "perfil", "herramientas",
                                  "simulador", "plan", "ejercicios", "acerca"]):  # Actualizada con "ejercicios"
        if st.session_state.get(f"tab_{tab_name}_clicked", False):
            # Restablecer todas las banderas de clic
            for tab in ["correccion", "perfil", "herramientas", "simulador", "plan", "ejercicios", "acerca"]:
                st.session_state[f"tab_{tab}_clicked"] = False

            # Actualizar pestaña activa
            set_session_var("active_tab", i)
            break

    # Añadir footer personalizado
    add_footer_local()


# Ejecutar la aplicación si este script es el punto de entrada
if __name__ == "__main__":
    # Ejecutar la aplicación con mejor manejo de errores
    try:
        main()
    except Exception as e:
        # Mostrar error grave de forma elegante
        st.error(f"Se ha producido un error en la aplicación: {str(e)}")
        st.warning(
            "La aplicación intentará reiniciarse. Si el problema persiste, contacta con soporte.")

        # Registrar el error
        logger.error(f"Error crítico en la aplicación: {str(e)}")
        logger.error(traceback.format_exc())

        # Intentar limpiar el estado para evitar errores persistentes
        try:
            for key in list(st.session_state.keys()):
                if key not in ["usuario_actual", "email_usuario", "uid_usuario", "nivel_estudiante"]:
                    del st.session_state[key]
        except:
            pass

        # Opción para reiniciar completamente
        if st.button("Reiniciar aplicación", type="primary"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()


# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 11: Funciones Auxiliares de Cálculos Lingüísticos
# ==================================================================================
#
# Este artefacto contiene:
# 1. Funciones adicionales para análisis lingüístico
# 2. Cálculo de índices de legibilidad
# 3. Métricas específicas para el español
# 4. Utilidades de análisis textual
#
# Estas funciones complementan las capacidades de análisis de textos
# de la aplicación, proporcionando métricas objetivas adicionales.
# ==================================================================================

def calcular_indice_szigriszt(texto):
    """
    Calcula el índice de legibilidad Szigriszt-Pazos adaptado para español.
    Este índice es una adaptación del índice Flesch para el español.

    Args:
        texto: Texto a analizar

    Returns:
        float: Índice de legibilidad Szigriszt (0-100)
    """
    if not texto:
        return 0

    try:
        # Limpiar el texto
        texto_limpio = re.sub(r'\s+', ' ', texto).strip()

        # Contar sílabas (aproximación para español)
        def contar_silabas_palabra(palabra):
            # Normalizar la palabra
            palabra = palabra.lower().strip()
            if not palabra:
                return 0

            # Remover signos de puntuación
            palabra = re.sub(r'[^\w\sáéíóúüñ]', '', palabra)

            # Contar vocales
            vocales = "aeiouáéíóúü"
            num_vocales = sum(1 for letra in palabra if letra in vocales)

            # Contar diptongos (aproximación)
            diptongos_patrones = [
                "ai", "au", "ei", "eu", "io", "iu", "oi", "ou", "ui", "ia", "ua", "ie", "ue", "uo"
            ]
            diptongos = sum(palabra.count(dip) for dip in diptongos_patrones)

            # Ajustar conteo de sílabas
            silabas = num_vocales - diptongos

            # Ajustes especiales
            if len(palabra) > 2 and palabra[-1] in "aeiou" and palabra[-2] not in vocales:
                silabas -= 0.5  # Ajuste para palabras que terminan en vocal

            return max(1, round(silabas))

        # Dividir en palabras
        palabras = re.findall(r'\b\w+\b', texto_limpio.lower())
        num_palabras = len(palabras)

        # Contar sílabas totales
        num_silabas = sum(contar_silabas_palabra(palabra)
                          for palabra in palabras)

        # Dividir en frases
        frases = re.split(r'[.!?;]', texto_limpio)
        frases = [f for f in frases if f.strip()]
        num_frases = len(frases) or 1  # Evitar división por cero

        # Calcular longitud media de frase en palabras
        promedio_palabras_por_frase = num_palabras / num_frases

        # Calcular promedio de sílabas por palabra
        promedio_silabas_por_palabra = num_silabas / \
            num_palabras if num_palabras > 0 else 0

        # Fórmula del índice Szigriszt-Pazos
        indice = 206.835 - (62.3 * promedio_silabas_por_palabra) - \
            (promedio_palabras_por_frase)

        # Limitar el resultado a un rango válido (0-100)
        indice = max(0, min(100, indice))

        return round(indice, 2)

    except Exception as e:
        logger.error(f"Error al calcular índice Szigriszt: {str(e)}")
        return 0


def interpretar_indice_szigriszt(indice):
    """
    Interpreta el valor del índice Szigriszt-Pazos.

    Args:
        indice: Valor del índice (0-100)

    Returns:
        tuple: (nivel_dificultad, descripcion, nivel_educativo)
    """
    if indice > 80:
        return ("Muy fácil", "Textos muy fáciles de leer", "Educación primaria")
    elif indice > 65:
        return ("Fácil", "Textos bastante fáciles de leer", "Educación secundaria")
    elif indice > 50:
        return ("Normal", "Textos de dificultad media", "Bachillerato/Formación profesional")
    elif indice > 35:
        return ("Difícil", "Textos bastante difíciles", "Universitario")
    else:
        return ("Muy difícil", "Textos muy difíciles de leer", "Especialización/Posgrado")


def calcular_nivel_complejidad(texto):
    """
    Calcula un índice de complejidad lingüística para el texto.
    Combina varios indicadores para estimar el nivel aproximado.

    Args:
        texto: Texto a analizar

    Returns:
        float: Valor de complejidad en escala 0-10
    """
    if not texto:
        return 0

    try:
        # Limpiar el texto
        texto_limpio = re.sub(r'\s+', ' ', texto).strip()

        # Extraer palabras
        palabras = re.findall(r'\b\w+\b', texto_limpio.lower())
        num_palabras = len(palabras)

        if num_palabras == 0:
            return 0

        # Métricas básicas
        palabras_unicas = len(set(palabras))
        longitud_media_palabra = sum(len(palabra)
                                     for palabra in palabras) / num_palabras

        # Extraer frases
        frases = re.split(r'[.!?;]', texto_limpio)
        frases = [f for f in frases if f.strip()]
        num_frases = len(frases) or 1  # Evitar división por cero

        # Longitud media de las frases
        palabras_por_frase = num_palabras / num_frases

        # Índice de diversidad léxica (Type-Token Ratio)
        diversidad_lexica = palabras_unicas / num_palabras

        # Identificar palabras complejas (más de 3 sílabas)
        def es_palabra_compleja(palabra):
            silabas = 0
            vocales = "aeiouáéíóúü"
            hay_vocal = False

            for i, letra in enumerate(palabra):
                if letra in vocales:
                    if i == 0 or palabra[i-1] not in vocales:
                        silabas += 1
                    hay_vocal = True

            return hay_vocal and silabas > 3

        palabras_complejas = sum(
            1 for palabra in palabras if es_palabra_compleja(palabra))
        porcentaje_palabras_complejas = (
            palabras_complejas / num_palabras) if num_palabras > 0 else 0

        # Detectar estructuras subordinadas (aproximación)
        conjunciones_subordinacion = [
            "que", "porque", "aunque", "cuando", "como", "si", "donde", "mientras", "pues"]
        subordinadas = sum(
            1 for palabra in palabras if palabra in conjunciones_subordinacion)
        subordinadas_por_frase = subordinadas / num_frases

        # Calcular un índice de complejidad ponderado
        # Pesos para diferentes factores
        pesos = {
            "diversidad_lexica": 0.25,
            "longitud_media_palabra": 0.15,
            "palabras_por_frase": 0.20,
            "palabras_complejas": 0.25,
            "subordinadas_por_frase": 0.15
        }

        # Normalizar cada factor a escala 0-1
        factores = {
            # 0.7 es un valor alto de referencia
            "diversidad_lexica": min(1, diversidad_lexica / 0.7),
            # 3-8 caracteres
            "longitud_media_palabra": min(1, (longitud_media_palabra - 3) / 5),
            # 25 palabras por frase es complejo
            "palabras_por_frase": min(1, palabras_por_frase / 25),
            # 20% palabras complejas
            "palabras_complejas": min(1, porcentaje_palabras_complejas / 0.2),
            # 2 subordinadas por frase
            "subordinadas_por_frase": min(1, subordinadas_por_frase / 2)
        }

        # Calcular complejidad ponderada
        complejidad = sum(factores[factor] * pesos[factor]
                          for factor in factores)

        # Escalar a 0-10
        complejidad = complejidad * 10

        return round(complejidad, 2)

    except Exception as e:
        logger.error(f"Error al calcular nivel de complejidad: {str(e)}")
        return 0


def determinar_nivel_cefr(texto):
    """
    Estima el nivel CEFR (A1-C2) del texto basado en su complejidad.

    Args:
        texto: Texto a analizar

    Returns:
        str: Nivel CEFR estimado (A1, A2, B1, B2, C1, C2)
    """
    # Calcular métricas
    complejidad = calcular_nivel_complejidad(texto)
    indice_szigriszt = calcular_indice_szigriszt(texto)

    # Convertir índice Szigriszt a escala inversa (más alto = más complejo)
    szigriszt_inverso = max(0, (100 - indice_szigriszt) / 10)

    # Ponderar ambas métricas (60% complejidad, 40% szigriszt)
    valor_combinado = (complejidad * 0.6) + (szigriszt_inverso * 0.4)

    # Mapear a niveles CEFR
    if valor_combinado < 2:
        return "A1"
    elif valor_combinado < 3.5:
        return "A2"
    elif valor_combinado < 5:
        return "B1"
    elif valor_combinado < 6.5:
        return "B2"
    elif valor_combinado < 8:
        return "C1"
    else:
        return "C2"


def analizar_lexica_basica(texto):
    """
    Realiza un análisis léxico básico del texto.

    Args:
        texto: Texto a analizar

    Returns:
        dict: Diccionario con estadísticas léxicas
    """
    if not texto:
        return {
            "num_palabras": 0,
            "palabras_unicas": 0,
            "longitud_media": 0,
            "palabras_frecuentes": []
        }

    try:
        # Limpiar el texto
        texto_limpio = re.sub(r'\s+', ' ', texto).strip().lower()

        # Extraer palabras
        palabras = re.findall(r'\b\w+\b', texto_limpio)
        num_palabras = len(palabras)

        if num_palabras == 0:
            return {
                "num_palabras": 0,
                "palabras_unicas": 0,
                "longitud_media": 0,
                "palabras_frecuentes": []
            }

        # Calcular estadísticas básicas
        palabras_unicas = set(palabras)
        num_palabras_unicas = len(palabras_unicas)
        longitud_media = sum(len(palabra)
                             for palabra in palabras) / num_palabras

        # Contar frecuencia de palabras
        from collections import Counter
        contador = Counter(palabras)

        # Excluir palabras muy comunes (stop words básicas en español)
        stop_words = set([
            "el", "la", "los", "las", "un", "una", "unos", "unas", "y", "o", "pero", "porque",
            "que", "de", "en", "a", "con", "por", "para", "del", "al", "mi", "tu", "su", "es",
            "no", "si", "sí", "ya", "le", "lo", "se", "me", "te", "nos", "muy", "sin", "sobre",
            "este", "esta", "estos", "estas", "ese", "esa", "esos", "esas", "como", "cuando",
            "donde", "quien", "quienes", "yo", "tú", "él", "ella", "nosotros", "vosotros", "ellos"
        ])

        # Filtrar stop words y obtener palabras más frecuentes
        palabras_frecuentes = [
            {"palabra": palabra, "frecuencia": freq}
            for palabra, freq in contador.most_common(10)
            # Filtrar palabras cortas
            if palabra not in stop_words and len(palabra) > 3
        ]

        return {
            "num_palabras": num_palabras,
            "palabras_unicas": num_palabras_unicas,
            "diversidad_lexica": round(num_palabras_unicas / num_palabras, 3),
            "longitud_media": round(longitud_media, 2),
            # Limitamos a 5 para no sobrecargar
            "palabras_frecuentes": palabras_frecuentes[:5]
        }

    except Exception as e:
        logger.error(f"Error en análisis léxico: {str(e)}")
        return {
            "error": str(e),
            "num_palabras": 0,
            "palabras_unicas": 0,
            "longitud_media": 0,
            "palabras_frecuentes": []
        }


def analizar_estructura_texto(texto):
    """
    Analiza la estructura básica del texto: párrafos, oraciones, etc.

    Args:
        texto: Texto a analizar

    Returns:
        dict: Estadísticas de la estructura
    """
    if not texto:
        return {
            "num_parrafos": 0,
            "num_oraciones": 0,
            "oraciones_por_parrafo": 0,
        }

    try:
        # Dividir en párrafos (por doble salto de línea)
        parrafos = re.split(r'\n\s*\n', texto)
        parrafos = [p.strip() for p in parrafos if p.strip()]
        num_parrafos = len(parrafos)

        # Dividir en oraciones
        patron_oracion = r'[.!?]+\s+[A-ZÁÉÍÓÚÑÜ]|[.!?]+$'
        oraciones = re.split(patron_oracion, texto)
        oraciones = [o.strip() for o in oraciones if o.strip()]
        num_oraciones = len(oraciones)

        # Calcular promedios
        oraciones_por_parrafo = num_oraciones / num_parrafos if num_parrafos > 0 else 0

        # Detectar patrones de estructura
        tiene_introduccion = len(parrafos) >= 3
        tiene_conclusion = len(parrafos) >= 3

        return {
            "num_parrafos": num_parrafos,
            "num_oraciones": num_oraciones,
            "oraciones_por_parrafo": round(oraciones_por_parrafo, 2),
            "estructura_detectada": {
                "tiene_introduccion": tiene_introduccion,
                "tiene_conclusion": tiene_conclusion,
                "estructura_clara": num_parrafos > 1 and oraciones_por_parrafo < 5  # Heurística simple
            }
        }

    except Exception as e:
        logger.error(f"Error en análisis de estructura: {str(e)}")
        return {
            "error": str(e),
            "num_parrafos": 0,
            "num_oraciones": 0,
            "oraciones_por_parrafo": 0,
        }


def analizar_texto_completo(texto, nivel="intermedio"):
    """
    Realiza un análisis completo del texto combinando todas las métricas.

    Args:
        texto: Texto a analizar
        nivel: Nivel esperado del estudiante

    Returns:
        dict: Análisis completo
    """
    if not texto:
        return {"error": "Texto vacío"}

    try:
        # Realizar todos los análisis
        lexica = analizar_lexica_basica(texto)
        estructura = analizar_estructura_texto(texto)
        complejidad = calcular_nivel_complejidad(texto)
        indice_szigriszt = calcular_indice_szigriszt(texto)
        nivel_estimado = determinar_nivel_cefr(texto)

        # Interpretar el índice Szigriszt
        interpretacion_szigriszt = interpretar_indice_szigriszt(
            indice_szigriszt)

        # Verificar adecuación al nivel declarado
        nivel_map = {
            "principiante": ["A1", "A2"],
            "intermedio": ["B1", "B2"],
            "avanzado": ["C1", "C2"]
        }

        nivel_esperado = nivel_map.get(nivel.lower(), ["B1", "B2"])
        adecuado_nivel = nivel_estimado in nivel_esperado

        # Determinar si el nivel es demasiado alto o bajo
        todos_niveles = ["A1", "A2", "B1", "B2", "C1", "C2"]
        idx_estimado = todos_niveles.index(nivel_estimado)
        idx_esperado_min = todos_niveles.index(nivel_esperado[0])
        idx_esperado_max = todos_niveles.index(nivel_esperado[-1])

        nivel_relacion = "adecuado"
        if idx_estimado < idx_esperado_min:
            nivel_relacion = "por debajo"
        elif idx_estimado > idx_esperado_max:
            nivel_relacion = "por encima"

        # Generar recomendaciones basadas en el análisis
        recomendaciones = []

        # Recomendaciones según complejidad
        if complejidad < 3 and nivel != "principiante":
            recomendaciones.append(
                "Intenta usar estructuras más complejas y vocabulario más variado")
        elif complejidad > 7 and nivel == "principiante":
            recomendaciones.append(
                "Tu texto puede ser demasiado complejo para tu nivel. Intenta simplificar las estructuras")

        # Recomendaciones según legibilidad
        if indice_szigriszt < 50 and nivel in ["principiante", "intermedio"]:
            recomendaciones.append(
                "Tu texto es bastante difícil de leer. Intenta usar frases más cortas y vocabulario más sencillo")

        # Recomendaciones según estructura
        if not estructura.get("estructura_detectada", {}).get("estructura_clara", True):
            recomendaciones.append(
                "Mejora la estructura de tu texto dividiendo ideas en párrafos más claros")

        # Recomendaciones según diversidad léxica
        diversidad = lexica.get("diversidad_lexica", 0)
        if diversidad < 0.4 and nivel != "principiante":
            recomendaciones.append(
                "Intenta utilizar un vocabulario más variado, evitando repeticiones")

        # Resultado completo
        return {
            "estadisticas_basicas": {
                "num_palabras": lexica.get("num_palabras", 0),
                "num_parrafos": estructura.get("num_parrafos", 0),
                "num_oraciones": estructura.get("num_oraciones", 0)
            },
            "analisis_lexico": lexica,
            "analisis_estructura": estructura,
            "complejidad": {
                "nivel_complejidad": complejidad,
                "indice_szigriszt": indice_szigriszt,
                "interpretacion_szigriszt": interpretacion_szigriszt,
                "nivel_estimado_cefr": nivel_estimado
            },
            "adecuacion_nivel": {
                "nivel_declarado": nivel,
                "nivel_estimado": nivel_estimado,
                "adecuado": adecuado_nivel,
                "relacion": nivel_relacion
            },
            "recomendaciones": recomendaciones
        }

    except Exception as e:
        logger.error(f"Error en análisis completo: {str(e)}")
        return {"error": f"Error en análisis: {str(e)}"}

# ==================================================================================
# SPANISH FACTORIA - TEXTOCORRECTOR ELE
# ==================================================================================
# Artefacto 12: Requisitos y Documentación de Instalación
# ==================================================================================
#
# Este artefacto contiene:
# 1. Requisitos de bibliotecas necesarias para ejecutar la aplicación
# 2. Instrucciones de configuración de APIs externas
# 3. Guía de instalación y despliegue
# 4. Estructura de directorios recomendada
#
# Este archivo es crucial para que cualquier persona pueda instalar, configurar
# y ejecutar correctamente la aplicación.
# ==================================================================================


def get_app_description():
    """
    Retorna la descripción básica de la aplicación.
    Esta función debe ser llamada explícitamente cuando se necesite la información.
    """
    return {
        "nombre": "Spanish FactorIA - Textocorrector ELE",
        "version": "3.2.0",
        "autor": "Spanish FactorIA",
        "descripcion": "Una aplicación Streamlit para la corrección de textos en español con análisis "
        "contextual avanzado, enfocada en estudiantes de español como lengua extranjera (ELE).",
        "requisitos_sistema": [
            "Python 3.8 o superior",
            "Acceso a Internet para APIs externas",
            "Mínimo 2GB de RAM",
            "Al menos 200MB de espacio en disco"
        ]
    }


def get_requirements():
    """
    Retorna la lista de dependencias para requirements.txt.
    """
    return """
# Bibliotecas principales
streamlit>=1.26.0
pandas>=1.5.0
numpy>=1.24.0
matplotlib>=3.7.0
altair>=5.0.0
pillow>=9.0.0

# Procesamiento de texto y documentos
python-docx>=0.8.11
markdown>=3.4.0
qrcode>=7.3.0

# APIs y servicios externos
requests>=2.28.0
firebase-admin>=6.2.0

# Utilidades adicionales
python-dateutil>=2.8.2
pytz>=2023.3
"""


def get_api_config_instructions():
    """
    Retorna las instrucciones de configuración de APIs externas.
    """
    return """
Configuración de APIs Externas
-----------------------------

La aplicación utiliza varias APIs externas que requieren configuración:

1. OpenAI API
   - Regístrate en https://platform.openai.com/
   - Obtén una API Key desde la sección de configuración
   - Añade la clave a los secretos de Streamlit como OPENAI_API_KEY

2. Firebase (opcional, para almacenamiento persistente)
   - Crea un proyecto en https://console.firebase.google.com/
   - Genera una nueva clave privada para tu cuenta de servicio
   - Convierte el archivo JSON de credenciales a formato string
   - Añade las credenciales a los secretos de Streamlit como FIREBASE_CREDENTIALS

3. ElevenLabs (opcional, para síntesis de voz)
   - Regístrate en https://elevenlabs.io/
   - Obtén una API Key desde la configuración de tu cuenta
   - Selecciona una voz y obtén su ID
   - Añade las claves a los secretos como ELEVENLABS_API_KEY y ELEVENLABS_VOICE_ID

Configuración de secretos en Streamlit
-------------------------------------

1. Para desarrollo local:
   - Crea un archivo .streamlit/secrets.toml en el directorio raíz con el siguiente formato:

   ```toml
   OPENAI_API_KEY = "sk-tu-clave-de-openai"
   
   [firebase]
   FIREBASE_CREDENTIALS = '''
   {
     "type": "service_account",
     "project_id": "tu-proyecto",
     ...resto de tu JSON de credenciales...
   }
   '''
   
   ELEVENLABS_API_KEY = "tu-clave-de-elevenlabs"
   ELEVENLABS_VOICE_ID = "id-de-voz"
   ```

2. Para Streamlit Cloud:
   - Ve a la configuración de tu aplicación en Streamlit Cloud
   - Navega a la sección "Secrets"
   - Añade las variables de entorno en el mismo formato
"""


def get_directory_structure():
    """
    Retorna la estructura de directorios recomendada.
    """
    return """
Estructura de Directorios Recomendada
-----------------------------------

```
textocorrector-ele/
│
├── .streamlit/                  # Configuración de Streamlit
│   ├── config.toml              # Configuración general
│   └── secrets.toml             # Secretos (local, no subir a Git)
│
├── assets/                      # Recursos estáticos
│   ├── Spanish_FactorIA_Logo.png # Logo principal
│   └── favicon.ico              # Favicon
│
├── data/                        # Datos estáticos o de caché
│   └── cached_resources/        # Recursos cacheados
│
├── docs/                        # Documentación
│   ├── MANUAL.md                # Manual de usuario
│   └── API.md                   # Documentación de API
│
├── app.py                       # Punto de entrada principal
├── requirements.txt             # Dependencias
├── README.md                    # Documentación principal
└── .gitignore                   # Configuración de Git
```
"""


def get_installation_instructions():
    """
    Retorna las instrucciones de instalación.
    """
    return """
Instrucciones de Instalación
---------------------------

1. Clonar el repositorio o descargar los archivos:
   ```
   git clone https://github.com/spanishfactoria/textocorrector-ele.git
   cd textocorrector-ele
   ```

2. Crear un entorno virtual (recomendado):
   ```
   python -m venv venv
   
   # En Windows
   venv\\Scripts\\activate
   
   # En macOS/Linux
   source venv/bin/activate
   ```

3. Instalar dependencias:
   ```
   pip install -r requirements.txt
   ```

4. Configurar secretos para APIs:
   - Crear archivo .streamlit/secrets.toml según las instrucciones anteriores

5. Ejecutar la aplicación:
   ```
   streamlit run app.py
   ```

6. La aplicación debería estar disponible en http://localhost:8501
"""


def get_deployment_instructions():
    """
    Retorna las instrucciones de despliegue.
    """
    return """
Instrucciones de Despliegue
--------------------------

Despliegue en Streamlit Cloud:
1. Sube el código a un repositorio de GitHub
2. Inicia sesión en https://share.streamlit.io/
3. Selecciona "New app" y elige tu repositorio
4. Configura la ruta del archivo principal como "app.py"
5. Configura los secretos en la sección de secretos
6. Haz clic en "Deploy!"

Despliegue con Docker:
1. Crea un Dockerfile:
   ```
   FROM python:3.9-slim
   
   WORKDIR /app
   
   COPY requirements.txt .
   RUN pip install -r requirements.txt
   
   COPY . .
   
   EXPOSE 8501
   
   CMD ["streamlit", "run", "app.py"]
   ```

2. Construye la imagen:
   ```
   docker build -t textocorrector-ele .
   ```

3. Ejecuta el contenedor:
   ```
   docker run -p 8501:8501 textocorrector-ele
   ```

Nota sobre secretos:
- Para Docker, puedes usar variables de entorno:
  ```
  docker run -p 8501:8501 -e OPENAI_API_KEY=sk-your-key textocorrector-ele
  ```
- O crear un archivo .env y usarlo:
  ```
  docker run -p 8501:8501 --env-file .env textocorrector-ele
  ```
"""


def get_troubleshooting():
    """
    Retorna la guía de solución de problemas comunes.
    """
    return """
Solución de Problemas
--------------------

1. Error "API key not configured":
   - Verifica que has configurado correctamente OPENAI_API_KEY en secrets.toml
   - Asegúrate de que la API key sea válida y tenga fondos suficientes

2. Error "Firebase connection failed":
   - Verifica el formato del JSON de credenciales
   - Comprueba que tu proyecto Firebase tiene Firestore habilitado
   - Verifica los permisos de la cuenta de servicio

3. Error "Module not found":
   - Ejecuta `pip install -r requirements.txt` para asegurarte de que todas las dependencias están instaladas
   - Si una biblioteca específica causa problemas, intenta instalarla individualmente

4. Error "Connection timeout":
   - La aplicación intenta reintentar automáticamente las conexiones
   - Verifica tu conexión a Internet
   - Las APIs de OpenAI pueden estar sobrecargadas, intenta más tarde

5. Problemas de rendimiento:
   - Aumenta la memoria disponible para la aplicación
   - Considera utilizar un servicio en la nube con más recursos
   - Optimiza el almacenamiento en caché para consultas frecuentes

Para obtener más ayuda, visita: https://github.com/spanishfactoria/textocorrector-ele/issues
"""


def generate_requirements_file():
    """
    Genera el contenido del archivo requirements.txt.
    """
    return get_requirements().strip()


def generate_readme():
    """
    Genera el contenido del archivo README.md.
    """
    return """# Spanish FactorIA - Textocorrector ELE

![Logo](assets/Spanish_FactorIA_Logo.png)

## Descripción

Textocorrector ELE es una aplicación avanzada para la corrección de textos en español, diseñada específicamente para estudiantes de español como lengua extranjera (ELE). Utiliza inteligencia artificial para proporcionar correcciones detalladas y análisis contextual del texto.

## Características

- **Corrección detallada** con análisis gramatical, léxico y contextual
- **Análisis contextual** de coherencia, cohesión, registro lingüístico y adecuación cultural
- **Seguimiento de progreso** con estadísticas y visualización de mejora
- **Recomendaciones personalizadas** adaptadas al nivel del estudiante
- **Exportación de informes** en diferentes formatos (Word, HTML, CSV)
- **Herramientas complementarias** como generador de imágenes para descripción, transcripción de textos manuscritos y generador de consignas

## Instalación

Consulta el archivo [INSTALLATION.md](docs/INSTALLATION.md) para instrucciones detalladas de instalación y configuración.

## Uso

1. Inicia la aplicación con `streamlit run app.py`
2. Regístrate o continúa como invitado
3. Selecciona la pestaña de corrección de texto
4. Escribe o pega tu texto y configura las opciones
5. Haz clic en "Corregir texto" para recibir el análisis
6. Explora las demás herramientas desde las diferentes pestañas

## Requisitos

- Python 3.8+
- Streamlit 1.26+
- Conexión a Internet para APIs externas

## Licencia

Este proyecto está bajo la licencia MIT. Consulta el archivo LICENSE para más detalles.

## Contacto

Spanish FactorIA - info@spanishfactoria.com

Sitio web: [www.spanishfactoria.com](https://www.spanishfactoria.com)
"""


def generate_config_toml():
    """
    Genera el contenido del archivo config.toml para Streamlit.
    """
    return """[theme]
primaryColor = "#3498db"
backgroundColor = "#ffffff"
secondaryBackgroundColor = "#f8f9fa"
textColor = "#333333"
font = "sans serif"

[server]
enableCORS = true
enableXsrfProtection = true
maxUploadSize = 10

[browser]
gatherUsageStats = false
"""


def generate_user_guide():
    """
    Genera el contenido de una guía completa para el usuario.
    """
    return """# Guía de Usuario - Textocorrector ELE

## Introducción

Bienvenido/a a Textocorrector ELE, tu asistente avanzado para mejorar tus habilidades de escritura en español. Esta guía te ayudará a aprovechar al máximo todas las funcionalidades de la aplicación.

## Primeros pasos

### Acceso a la aplicación

1. Al abrir la aplicación, encontrarás una pantalla de bienvenida.
2. Puedes registrarte proporcionando tu nombre, correo electrónico y nivel de español.
3. También puedes continuar como invitado si prefieres no registrarte.

### Navegación principal

La aplicación está organizada en pestañas:

- **Corrección de texto**: Función principal para corregir textos.
- **Perfil y Progreso**: Seguimiento de tu evolución y estadísticas.
- **Herramientas ELE**: Herramientas complementarias para aprendizaje.
- **Simulador de examen**: Practica para exámenes oficiales.
- **Plan de estudio**: Recomendaciones personalizadas de estudio.
- **Acerca de**: Información sobre la aplicación.

## Corrección de texto

1. Selecciona el idioma en el que quieres recibir las explicaciones.
2. Elige el tipo de texto y el contexto cultural relevante.
3. Escribe o pega tu texto en el área correspondiente.
4. Opcionalmente, añade información adicional que pueda ser útil.
5. Haz clic en "Corregir texto".
6. Recibirás un análisis completo con:
   - Texto corregido
   - Errores detectados por categoría
   - Análisis contextual (coherencia, cohesión, registro, adecuación)
   - Consejo final personalizado

## Perfil y Progreso

En esta sección podrás:

1. Ver estadísticas generales de tu progreso.
2. Explorar gráficos de evolución de tus errores y puntuaciones.
3. Identificar áreas específicas que necesitan mejora.
4. Consultar un historial completo de tus correcciones.
5. Recibir consejos personalizados basados en tus datos.

## Herramientas ELE

### Generador de imágenes para descripción

1. Introduce un tema para la imagen.
2. Selecciona tu nivel de español.
3. Haz clic en "Generar imagen".
4. Recibirás una imagen con descripción y preguntas.
5. Practica escribiendo tu propia descripción.

### Transcripción de textos manuscritos

1. Sube una imagen con texto manuscrito.
2. Selecciona el idioma del texto.
3. Haz clic en "Transcribir texto".
4. Obtendrás una transcripción digital que podrás corregir.

### Generador de consignas

1. Selecciona tu nivel y el tipo de texto.
2. Haz clic en "Generar consigna".
3. Recibirás una consigna de escritura adaptada.
4. Puedes escribir una respuesta y corregirla después.

## Simulador de examen

1. Selecciona el nivel del examen DELE que quieres practicar.
2. Elige la duración del simulacro.
3. Haz clic en "Iniciar simulacro".
4. Recibirás una tarea realista con temporizador.
5. Escribe tu respuesta y finaliza para recibir corrección.

## Plan de estudio

Aquí encontrarás un plan de estudio personalizado según tu nivel, con:

1. Objetivos específicos para tu nivel.
2. Actividades recomendadas.
3. Recursos específicos.
4. Cronograma sugerido.
5. Consejos para la práctica independiente.

## Exportación de resultados

Después de cada corrección, puedes:

1. Descargar un informe en formato Word (DOCX).
2. Descargar un informe en formato HTML.
3. Descargar los datos de análisis en formato CSV.

## Consejos adicionales

- **Registro regular**: Para un seguimiento óptimo, utiliza la misma cuenta siempre.
- **Variedad de textos**: Practica con diferentes tipos de texto para un aprendizaje más completo.
- **Análisis de errores**: Presta atención a los patrones de error para enfocar tu aprendizaje.
- **Guardar informes**: Exporta tus resultados para revisarlos posteriormente.

## Solución de problemas

Si encuentras algún problema:

1. Verifica tu conexión a Internet.
2. Recarga la página si la aplicación no responde.
3. Si persisten los problemas, contacta con el soporte.

¡Disfruta aprendiendo y mejorando tu español con Textocorrector ELE!
"""


def generate_documentation():
    """
    Genera toda la documentación necesaria para el proyecto.
    Esta función debe ser llamada explícitamente cuando se necesite generar la documentación.
    """
    return {
        "requirements.txt": generate_requirements_file(),
        "README.md": generate_readme(),
        ".streamlit/config.toml": generate_config_toml(),
        "docs/INSTALLATION.md": get_installation_instructions(),
        "docs/DEPLOYMENT.md": get_deployment_instructions(),
        "docs/TROUBLESHOOTING.md": get_troubleshooting(),
        "docs/API_CONFIG.md": get_api_config_instructions(),
        "docs/USER_GUIDE.md": generate_user_guide()
    }


# Función principal para mostrar documentación en Streamlit (solo se llamará si es necesario)
def show_documentation_in_ui():
    """
    Muestra la documentación en la interfaz de Streamlit.
    Esta función debería ser llamada solo cuando sea necesario, 
    por ejemplo desde la sección 'Acerca de'.
    """
    import streamlit as st

    app_info = get_app_description()

    st.title(app_info["nombre"])
    st.write(f"Versión: {app_info['version']} | Autor: {app_info['autor']}")

    st.header("Descripción")
    st.write(app_info["descripcion"])

    st.header("Requisitos de sistema")
    for req in app_info["requisitos_sistema"]:
        st.markdown(f"- {req}")

    # Más documentación disponible en pestañas
    tabs = st.tabs(["Instalación", "APIs", "Estructura",
                   "Despliegue", "Solución de problemas"])

    with tabs[0]:
        st.markdown(get_installation_instructions())

    with tabs[1]:
        st.markdown(get_api_config_instructions())

    with tabs[2]:
        st.markdown(get_directory_structure())

    with tabs[3]:
        st.markdown(get_deployment_instructions())

    with tabs[4]:
        st.markdown(get_troubleshooting())


# Este bloque solo se ejecutará si este script se ejecuta directamente,
# no cuando se importa como módulo en la aplicación principal
if __name__ == "__main__":
    docs = generate_documentation()
    print("Documentación generada con éxito.")
    print(f"Archivos generados: {', '.join(docs.keys())}")
