#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Módulo de exportación de informes
---------------------------------
Este módulo contiene funciones para exportar correcciones y reportes en diferentes formatos.
"""

import base64
import io
import json
import logging
import pandas as pd
import streamlit as st
import traceback
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.session_manager import get_session_var, get_user_info
from utils.text_processing import clean_html_tags

logger = logging.getLogger(__name__)

def generar_pdf_base64(html_content, filename="correccion_ele"):
    """
    Convierte contenido HTML a PDF y lo devuelve como string base64.
    
    Args:
        html_content (str): Contenido HTML a convertir
        filename (str): Nombre del archivo (sin extensión)
        
    Returns:
        str: String base64 del PDF generado
    """
    try:
        # Utilizamos pdfkit para la conversión de HTML a PDF
        import pdfkit
        
        logger.info("Comenzando la generación de PDF con pdfkit")
        
        # Configuración para pdfkit
        options = {
            'page-size': 'A4',
            'margin-top': '1cm',
            'margin-right': '1cm',
            'margin-bottom': '1cm',
            'margin-left': '1cm',
            'encoding': 'UTF-8',
            'no-outline': None,
            'quiet': ''
        }
        
        # Intentar encontrar y usar la ruta a wkhtmltopdf
        try:
            import os
            wkhtmltopdf_path = os.environ.get('WKHTMLTOPDF_PATH')
            if wkhtmltopdf_path:
                logger.info(f"Usando ruta de wkhtmltopdf: {wkhtmltopdf_path}")
                # Crear configuración con la ruta
                config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
                # Crear el PDF en memoria con configuración específica
                pdf_bytes = pdfkit.from_string(html_content, False, options=options, configuration=config)
            else:
                # Crear el PDF en memoria
                logger.info("Usando wkhtmltopdf del sistema")
                pdf_bytes = pdfkit.from_string(html_content, False, options=options)
        except Exception as e:
            logger.warning(f"Error al intentar configurar wkhtmltopdf: {str(e)}")
            # Intentar sin configuración específica
            pdf_bytes = pdfkit.from_string(html_content, False, options=options)
        
        # Convertir a base64
        b64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
        logger.info(f"PDF generado correctamente, tamaño en bytes: {len(pdf_bytes)}")
        return b64_pdf
    except Exception as e:
        logger.error(f"Error generando PDF: {str(e)}")
        logger.debug(traceback.format_exc())
        return None

def genera_html_informe_correccion(data):
    """
    Genera el contenido HTML para un informe de corrección.
    
    Args:
        data (dict): Datos de la corrección
        
    Returns:
        str: Contenido HTML del informe
    """
    try:
        # Obtener información del usuario
        info_usuario = get_user_info()
        nombre_usuario = f"{info_usuario.get('nombre', 'Estudiante')} {info_usuario.get('apellido', '')}" if info_usuario else "Estudiante"
        nivel = info_usuario.get('nivel', 'No especificado') if info_usuario else "No especificado"
        
        # Generar HTML para errores
        html_errores = ""
        errores = data.get("errores", {})
        if errores:
            for categoria, lista_errores in errores.items():
                if lista_errores:
                    html_errores += f"""
                    <div class="error-categoria">
                        <h3>{categoria} ({len(lista_errores)})</h3>
                        <table class="error-tabla">
                            <thead>
                                <tr>
                                    <th>Error</th>
                                    <th>Corrección</th>
                                    <th>Explicación</th>
                                </tr>
                            </thead>
                            <tbody>
                    """
                    
                    for error in lista_errores:
                        html_errores += f"""
                                <tr>
                                    <td class="error-fragmento">{error.get('fragmento_erroneo', '')}</td>
                                    <td class="error-correccion">{error.get('correccion', '')}</td>
                                    <td class="error-explicacion">{error.get('explicacion', '')}</td>
                                </tr>
                        """
                    
                    html_errores += """
                            </tbody>
                        </table>
                    </div>
                    """
        else:
            html_errores = "<p>No se encontraron errores en el texto.</p>"
        
        # Generar HTML para análisis contextual
        html_analisis = ""
        analisis = data.get("analisis_contextual", {})
        if analisis:
            # Coherencia
            if "coherencia" in analisis:
                coherencia = analisis["coherencia"]
                html_analisis += f"""
                <div class="analisis-seccion">
                    <h3>Coherencia ({coherencia.get('puntuacion', 0)}/10)</h3>
                    <p>{coherencia.get('comentario', '')}</p>
                """
                
                if "sugerencias" in coherencia and coherencia["sugerencias"]:
                    html_analisis += "<h4>Sugerencias:</h4><ul>"
                    for sugerencia in coherencia["sugerencias"]:
                        html_analisis += f"<li>{sugerencia}</li>"
                    html_analisis += "</ul>"
                
                html_analisis += "</div>"
            
            # Cohesión
            if "cohesion" in analisis:
                cohesion = analisis["cohesion"]
                html_analisis += f"""
                <div class="analisis-seccion">
                    <h3>Cohesión ({cohesion.get('puntuacion', 0)}/10)</h3>
                    <p>{cohesion.get('comentario', '')}</p>
                """
                
                if "sugerencias" in cohesion and cohesion["sugerencias"]:
                    html_analisis += "<h4>Sugerencias:</h4><ul>"
                    for sugerencia in cohesion["sugerencias"]:
                        html_analisis += f"<li>{sugerencia}</li>"
                    html_analisis += "</ul>"
                
                html_analisis += "</div>"
            
            # Registro lingüístico
            if "registro_linguistico" in analisis:
                registro = analisis["registro_linguistico"]
                html_analisis += f"""
                <div class="analisis-seccion">
                    <h3>Registro lingüístico ({registro.get('puntuacion', 0)}/10)</h3>
                    <p><strong>Tipo detectado:</strong> {registro.get('tipo_detectado', '')}</p>
                    <p>{registro.get('adecuacion', '')}</p>
                """
                
                if "sugerencias" in registro and registro["sugerencias"]:
                    html_analisis += "<h4>Sugerencias:</h4><ul>"
                    for sugerencia in registro["sugerencias"]:
                        html_analisis += f"<li>{sugerencia}</li>"
                    html_analisis += "</ul>"
                
                html_analisis += "</div>"
            
            # Adecuación cultural
            if "adecuacion_cultural" in analisis:
                adecuacion = analisis["adecuacion_cultural"]
                html_analisis += f"""
                <div class="analisis-seccion">
                    <h3>Adecuación cultural ({adecuacion.get('puntuacion', 0)}/10)</h3>
                    <p>{adecuacion.get('comentario', '')}</p>
                """
                
                if "elementos_destacables" in adecuacion and adecuacion["elementos_destacables"]:
                    html_analisis += "<h4>Elementos destacables:</h4><ul>"
                    for elemento in adecuacion["elementos_destacables"]:
                        html_analisis += f"<li>{elemento}</li>"
                    html_analisis += "</ul>"
                
                if "sugerencias" in adecuacion and adecuacion["sugerencias"]:
                    html_analisis += "<h4>Sugerencias:</h4><ul>"
                    for sugerencia in adecuacion["sugerencias"]:
                        html_analisis += f"<li>{sugerencia}</li>"
                    html_analisis += "</ul>"
                
                html_analisis += "</div>"
        else:
            html_analisis = "<p>No hay análisis contextual disponible.</p>"
        
        # Generar HTML completo del informe
        html_content = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Informe de Corrección ELE</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    margin: 0;
                    padding: 20px;
                    color: #333;
                }}
                h1 {{
                    color: #1E88E5;
                    border-bottom: 2px solid #1E88E5;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #1976D2;
                    margin-top: 20px;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 5px;
                }}
                h3 {{
                    color: #0D47A1;
                    margin-top: 15px;
                }}
                .info-seccion {{
                    background-color: #f5f5f5;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 20px;
                }}
                .texto-seccion {{
                    margin-bottom: 20px;
                }}
                .error-categoria {{
                    margin-bottom: 25px;
                }}
                .error-tabla {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-bottom: 15px;
                }}
                .error-tabla th {{
                    background-color: #e0e0e0;
                    padding: 8px;
                    text-align: left;
                }}
                .error-tabla td {{
                    padding: 8px;
                    border-bottom: 1px solid #ddd;
                }}
                .error-fragmento {{
                    color: #D32F2F;
                    text-decoration: line-through;
                }}
                .error-correccion {{
                    color: #388E3C;
                    font-weight: bold;
                }}
                .error-explicacion {{
                    color: #616161;
                    font-style: italic;
                }}
                .analisis-seccion {{
                    background-color: #f5f5f5;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 15px;
                }}
                .consejo-final {{
                    background-color: #E3F2FD;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 20px;
                    border-left: 5px solid #2196F3;
                }}
                .pie-pagina {{
                    text-align: center;
                    font-size: 0.8em;
                    color: #757575;
                    margin-top: 40px;
                    border-top: 1px solid #ddd;
                    padding-top: 10px;
                }}
            </style>
        </head>
        <body>
            <h1>Informe de Corrección de Texto - ELE</h1>
            
            <div class="info-seccion">
                <p><strong>Estudiante:</strong> {nombre_usuario}</p>
                <p><strong>Nivel:</strong> {nivel}</p>
                <p><strong>Fecha:</strong> {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
            </div>
            
            <h2>Texto Original</h2>
            <div class="texto-seccion">
                <p>{data.get('texto_original', '')}</p>
            </div>
            
            <h2>Texto Corregido</h2>
            <div class="texto-seccion">
                <p>{data.get('texto_corregido', '')}</p>
            </div>
            
            <h2>Análisis de Errores</h2>
            {html_errores}
            
            <h2>Análisis Contextual</h2>
            {html_analisis}
            
            <h2>Consejo Final</h2>
            <div class="consejo-final">
                <p>{data.get('consejo_final', '')}</p>
            </div>
            
            <div class="pie-pagina">
                <p>Generado por Textocorrector ELE - {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
            </div>
        </body>
        </html>
        """
        
        return html_content
    except Exception as e:
        logger.error(f"Error generando HTML del informe: {str(e)}")
        logger.debug(traceback.format_exc())
        return f"<html><body><h1>Error generando informe</h1><p>{str(e)}</p></body></html>"

def exportar_correccion_pdf(data):
    """
    Exporta una corrección a PDF y genera un enlace de descarga.
    
    Args:
        data (dict): Datos de la corrección
        
    Returns:
        bool: True si se exportó correctamente, False en caso contrario
    """
    try:
        logger.info("Iniciando exportación a PDF")
        
        # Generar HTML del informe
        html_content = genera_html_informe_correccion(data)
        
        # Generar un ID único para el archivo
        unique_id = str(uuid.uuid4())[:8]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"correccion_{timestamp}_{unique_id}"
        
        # Convertir a PDF
        b64_pdf = generar_pdf_base64(html_content, filename)
        
        if not b64_pdf:
            logger.error("No se pudo generar el PDF.")
            st.error("Error al generar el PDF. Asegúrate de tener wkhtmltopdf instalado.")
            return False
        
        # Método 1: Enlace directo de descarga
        href = f'''
        <a id="pdf_download_{unique_id}" 
           href="data:application/pdf;base64,{b64_pdf}" 
           download="{filename}.pdf" 
           style="display: inline-block; padding: 0.5rem 1rem; 
                  background-color: #4CAF50; color: white; 
                  text-decoration: none; border-radius: 4px;">
           📥 Descargar PDF
        </a>
        '''
        st.markdown(href, unsafe_allow_html=True)
        
        # Método 2: Forzar la descarga con JavaScript
        auto_download_js = f'''
        <script>
            (function() {{
                // Crear un enlace invisible
                var link = document.createElement('a');
                link.href = 'data:application/pdf;base64,{b64_pdf}';
                link.download = '{filename}.pdf';
                link.style.display = 'none';
                
                // Añadir al DOM, hacer clic y eliminar
                document.body.appendChild(link);
                link.click();
                
                // Esperar un poco antes de eliminarlo
                setTimeout(function() {{
                    document.body.removeChild(link);
                }}, 100);
            }})();
        </script>
        '''
        st.components.v1.html(auto_download_js, height=0)
        
        logger.info(f"PDF generado correctamente: {filename}.pdf")
        return True
    except Exception as e:
        logger.error(f"Error exportando a PDF: {str(e)}")
        logger.debug(traceback.format_exc())
        st.error(f"Error al exportar a PDF: {str(e)}")
        return False

def exportar_correccion_word(data):
    """
    Exporta una corrección a documento Word y genera un enlace de descarga.
    
    Args:
        data (dict): Datos de la corrección
        
    Returns:
        bool: True si se exportó correctamente, False en caso contrario
    """
    try:
        logger.info("Iniciando exportación a Word")
        
        # Crear un nuevo documento Word
        doc = Document()
        
        # Configurar márgenes del documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Título
        titulo = doc.add_heading("Informe de Corrección de Texto - ELE", level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del estudiante
        info_alumno = get_user_info()
        if info_alumno:
            doc.add_heading("Información del Estudiante", level=2)
            p = doc.add_paragraph()
            p.add_run("Nombre: ").bold = True
            p.add_run(f"{info_alumno.get('nombre', 'No especificado')} {info_alumno.get('apellido', '')}")
            
            p = doc.add_paragraph()
            p.add_run("Nivel: ").bold = True
            p.add_run(f"{info_alumno.get('nivel', 'No especificado')}")
        
        # Fecha y hora
        p = doc.add_paragraph()
        p.add_run("Fecha: ").bold = True
        p.add_run(datetime.now().strftime("%d/%m/%Y %H:%M"))
        
        # Texto original
        doc.add_heading("Texto Original", level=2)
        doc.add_paragraph(data.get("texto_original", ""))
        
        # Texto corregido
        doc.add_heading("Texto Corregido", level=2)
        corregido_limpio = clean_html_tags(data.get("texto_corregido", ""))
        doc.add_paragraph(corregido_limpio)
        
        # Análisis de errores
        doc.add_heading("Análisis de Errores", level=2)
        
        errores = data.get("errores", {})
        if errores:
            for categoria, lista_errores in errores.items():
                if lista_errores:
                    doc.add_heading(f"{categoria} ({len(lista_errores)})", level=3)
                    
                    # Crear tabla para errores
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    
                    # Encabezados
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Error"
                    hdr_cells[1].text = "Corrección"
                    hdr_cells[2].text = "Explicación"
                    
                    # Contenido
                    for error in lista_errores:
                        row_cells = table.add_row().cells
                        row_cells[0].text = error.get("fragmento_erroneo", "")
                        row_cells[1].text = error.get("correccion", "")
                        row_cells[2].text = error.get("explicacion", "")
        else:
            doc.add_paragraph("No se detectaron errores.")
        
        # Análisis contextual
        doc.add_heading("Análisis Contextual", level=2)
        
        analisis = data.get("analisis_contextual", {})
        if analisis:
            # Coherencia
            if "coherencia" in analisis:
                coherencia = analisis["coherencia"]
                doc.add_heading(f"Coherencia ({coherencia.get('puntuacion', 0)}/10)", level=3)
                doc.add_paragraph(coherencia.get('comentario', ''))
                
                if "sugerencias" in coherencia and coherencia["sugerencias"]:
                    p = doc.add_paragraph()
                    p.add_run("Sugerencias:").bold = True
                    
                    for sugerencia in coherencia["sugerencias"]:
                        doc.add_paragraph(sugerencia, style='List Bullet')
            
            # Cohesión
            if "cohesion" in analisis:
                cohesion = analisis["cohesion"]
                doc.add_heading(f"Cohesión ({cohesion.get('puntuacion', 0)}/10)", level=3)
                doc.add_paragraph(cohesion.get('comentario', ''))
                
                if "sugerencias" in cohesion and cohesion["sugerencias"]:
                    p = doc.add_paragraph()
                    p.add_run("Sugerencias:").bold = True
                    
                    for sugerencia in cohesion["sugerencias"]:
                        doc.add_paragraph(sugerencia, style='List Bullet')
            
            # Registro lingüístico
            if "registro_linguistico" in analisis:
                registro = analisis["registro_linguistico"]
                doc.add_heading(f"Registro lingüístico ({registro.get('puntuacion', 0)}/10)", level=3)
                
                p = doc.add_paragraph()
                p.add_run("Tipo detectado: ").bold = True
                p.add_run(registro.get('tipo_detectado', ''))
                
                doc.add_paragraph(registro.get('adecuacion', ''))
                
                if "sugerencias" in registro and registro["sugerencias"]:
                    p = doc.add_paragraph()
                    p.add_run("Sugerencias:").bold = True
                    
                    for sugerencia in registro["sugerencias"]:
                        doc.add_paragraph(sugerencia, style='List Bullet')
            
            # Adecuación cultural
            if "adecuacion_cultural" in analisis:
                adecuacion = analisis["adecuacion_cultural"]
                doc.add_heading(f"Adecuación cultural ({adecuacion.get('puntuacion', 0)}/10)", level=3)
                doc.add_paragraph(adecuacion.get('comentario', ''))
                
                if "elementos_destacables" in adecuacion and adecuacion["elementos_destacables"]:
                    p = doc.add_paragraph()
                    p.add_run("Elementos destacables:").bold = True
                    
                    for elemento in adecuacion["elementos_destacables"]:
                        doc.add_paragraph(elemento, style='List Bullet')
                
                if "sugerencias" in adecuacion and adecuacion["sugerencias"]:
                    p = doc.add_paragraph()
                    p.add_run("Sugerencias:").bold = True
                    
                    for sugerencia in adecuacion["sugerencias"]:
                        doc.add_paragraph(sugerencia, style='List Bullet')
        else:
            doc.add_paragraph("No hay análisis contextual disponible.")
        
        # Consejo final
        doc.add_heading("Consejo Final", level=2)
        doc.add_paragraph(data.get("consejo_final", ""))
        
        # Pie de página
        doc.add_paragraph(f"Generado por Textocorrector ELE - {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Guardar el documento en memoria
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        doc_bytes = f.getvalue()
        
        # Generar un ID único para el archivo
        unique_id = str(uuid.uuid4())[:8]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"correccion_{timestamp}_{unique_id}.docx"
        
        # Método 1: Enlace directo de descarga
        b64_doc = base64.b64encode(doc_bytes).decode()
        href = f'''
        <a id="word_download_{unique_id}" 
           href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}" 
           download="{filename}" 
           style="display: inline-block; padding: 0.5rem 1rem; 
                  background-color: #2196F3; color: white; 
                  text-decoration: none; border-radius: 4px;">
           📥 Descargar Word
        </a>
        '''
        st.markdown(href, unsafe_allow_html=True)
        
        # Método 2: Forzar la descarga con JavaScript
        auto_download_js = f'''
        <script>
            (function() {{
                // Crear un enlace invisible
                var link = document.createElement('a');
                link.href = 'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}';
                link.download = '{filename}';
                link.style.display = 'none';
                
                // Añadir al DOM, hacer clic y eliminar
                document.body.appendChild(link);
                link.click();
                
                // Esperar un poco antes de eliminarlo
                setTimeout(function() {{
                    document.body.removeChild(link);
                }}, 100);
            }})();
        </script>
        '''
        st.components.v1.html(auto_download_js, height=0)
        
        logger.info(f"Documento Word generado correctamente: {filename}")
        return True
    except Exception as e:
        logger.error(f"Error exportando a Word: {str(e)}")
        logger.debug(traceback.format_exc())
        st.error(f"Error al exportar a Word: {str(e)}")
        return False
